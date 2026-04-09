import os
import io
import re
from flask import Flask, render_template, redirect, url_for, request, flash, jsonify, session, send_from_directory, make_response
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import json
from config import Config
from models import db, User, KnowledgeBase, KnowledgeFile, SpecialReport, AssignmentTask, TaskSubmission, AIConversation, AIModelConfig, ChatSession, ChatMessage, KnowledgeFavorite, KnowledgeBrowseLog
from models import BriefingSource, BriefingKeyword, Briefing, BriefingArticle, BriefingScheduledTask, BriefingStatistics, BriefingSystemLog
from models import Role, SystemOperationLog, DocTemplate, SystemUsageStat
from models import OfficialDoc, DocFlow, DocReadRecord
from models import MeetingRoom, Meeting, MeetingAttendance, SupervisionTask, SupervisionProgress
from models import PerformancePeriod, PerformanceAssessment, WorkLog, WorkLogReview

from models import Organization, Department, Position
from models import CrawlerTask, CrawlerPage, CrawlerAttachment
from models import UrlLibrary, UrlItem, MonitorResult, MonitorLog, MonitorScheduledTask, MonitorSystemLog

from forms import *
from utils import *
from smart_knowledge import smart_kb, SmartKnowledgeBase
from config_manager import config_manager
# 档案管理模块 - 延迟导入避免循环依赖
# from archive_models import ...  # 在需要时导入
# from archive_digitizer import ...
# from archive_routes import ...
from flask_caching import Cache
from flask_wtf.csrf import CSRFProtect, generate_csrf
import time
import uuid
import threading

app = Flask(__name__)
app.config.from_object(Config)
db.init_app(app)

# 初始化 CSRF 保护
csrf = CSRFProtect(app)

# 初始化缓存
cache = Cache(app)

# 注册档案管理蓝图（包含质检/档号等新路由）
# 冲突路由已在下方注释排除
from archive_routes import archive_bp
app.register_blueprint(archive_bp)

# 档案模块：API路由免CSRF（前端AJAX无token context）
for rule in app.url_map.iter_rules():
    if rule.endpoint in ('archive.api_re_extract_file', 'archive.api_analyze_file'):
        csrf.exempt(app.view_functions[rule.endpoint])

# 以下直接路由已禁用，由蓝图统一提供（避免重复定义）
# @app.route('/archive/fonds') ... archive_fonds_list
# @app.route('/archive/files') ... archive_file_list
# @app.route('/archive/search') ... archive_search
# @app.route('/archive/my_borrows') ... archive_my_borrows
# @app.route('/archive/statistics') ... archive_statistics
# @app.route('/archive/file/<int:file_id>') ... archive_file_detail

# ==================== 档案管理首页 ====================
@app.route('/archive')
@login_required
def archive_index():
    """档案管理首页 - 重定向到蓝图版本"""
    return redirect(url_for('archive.index'))

@app.route('/archive/batch_upload')
@login_required
def archive_batch_upload():
    """批量上传档案页面 - 转发到蓝图"""
    return redirect(url_for('archive.batch_upload'))

@app.route('/archive/api/catalogs')
@login_required
def api_catalogs():
    """获取目录列表 API"""
    try:
        from archive_models import ArchiveCatalog
        fonds_id = request.args.get('fonds_id', type=int)
        if not fonds_id:
            return jsonify([])
        
        catalogs = ArchiveCatalog.query.filter_by(fonds_id=fonds_id).all()
        return jsonify([{
            'id': c.id,
            'catalog_code': c.catalog_code,
            'catalog_name': c.catalog_name,
            'retention_period': c.retention_period
        } for c in catalogs])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/archive/api/volumes')
@login_required
def api_volumes():
    """获取案卷列表 API"""
    try:
        from archive_models import ArchiveVolume
        catalog_id = request.args.get('catalog_id', type=int)
        if not catalog_id:
            return jsonify([])
        
        volumes = ArchiveVolume.query.filter_by(catalog_id=catalog_id).all()
        return jsonify([{
            'id': v.id,
            'volume_code': v.volume_code,
            'volume_title': v.volume_title
        } for v in volumes])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/archive/api/task/progress/<int:task_id>')
@login_required
def archive_task_progress(task_id):
    """查询档案批量上传任务进度"""
    try:
        from archive_models import ArchiveDigitizationTask
        
        task = ArchiveDigitizationTask.query.get_or_404(task_id)
        
        return jsonify({
            'task_id': task.id,
            'status': task.status,
            'total': task.total_files or 0,
            'processed': task.completed_files or 0,
            'success': task.completed_files or 0,
            'failed': task.failed_files or 0,
            'current_file': ''
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ===== P1: 自定义模板过滤器 =====
@app.template_filter('filesizeformat')
def filesizeformat_filter(size):
    """将字节数格式化为人类可读的大小"""
    if not size:
        return '0 B'
    if size >= 1073741824:
        return f'{size / 1073741824:.1f} GB'
    elif size >= 1048576:
        return f'{size / 1048576:.1f} MB'
    elif size >= 1024:
        return f'{size / 1024:.1f} KB'
    else:
        return f'{size} B'

@app.template_filter('highlight')
def highlight_filter(text, keyword):
    """在文本中高亮关键词（用于搜索结果）"""
    if not text or not keyword:
        return text
    try:
        escaped = re.escape(keyword)
        return re.sub(f'({escaped})', r'<mark class="keyword-highlight">\1</mark>', text, flags=re.IGNORECASE)
    except Exception:
        return text

@app.template_filter('highlight_excerpt')
def highlight_excerpt_filter(text, keyword, context_len=80):
    """提取包含关键词的片段并高亮"""
    if not text or not keyword:
        return text[:200] if text else ''
    try:
        escaped = re.escape(keyword)
        match = re.search(escaped, text, re.IGNORECASE)
        if not match:
            return text[:200] + ('...' if len(text) > 200 else '')
        start = max(0, match.start() - context_len)
        end = min(len(text), match.end() + context_len)
        excerpt = text[start:end]
        # 高亮
        excerpt = re.sub(f'({escaped})', r'<mark class="keyword-highlight">\1</mark>', excerpt, flags=re.IGNORECASE)
        prefix = '...' if start > 0 else ''
        suffix = '...' if end < len(text) else ''
        return prefix + excerpt + suffix
    except Exception:
        return text[:200]

@app.template_filter('from_json')
def from_json_filter(value):
    """解析JSON字符串"""
    if not value:
        return []
    try:
        return json.loads(value)
    except Exception:
        return []

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.unauthorized_handler
def unauthorized():
    """对AJAX请求返回JSON 401，而不是重定向到登录页"""
    from flask import request
    # 检测AJAX请求：fetch默认会带 X-Requested-With: XMLHttpRequest
    is_ajax = request.headers.get('X-Requested-With', '').lower() == 'xmlhttprequest'
    if is_ajax:
        return jsonify({'error': '登录已失效，请刷新页面重新登录'}), 401
    return redirect(url_for('login'))

@login_manager.user_loader
def load_user(user_id):
    try:
        return db.session.get(User, int(user_id))
    except (ValueError, TypeError):
        return None

# 初始化数据库和默认管理员
with app.app_context():
    db.create_all()

    # 轻量迁移：补齐 crawler_tasks.max_depth 字段（SQLite 不会因 create_all 自动补列）
    try:
        crawler_task_columns = {
            row[1] for row in db.session.execute(db.text("PRAGMA table_info(crawler_tasks)")).fetchall()
        }
        if 'max_depth' not in crawler_task_columns:
            db.session.execute(db.text("ALTER TABLE crawler_tasks ADD COLUMN max_depth INTEGER DEFAULT 3"))
            db.session.commit()
    except Exception:
        db.session.rollback()

    admin = User.query.filter_by(username='admin').first()

    if not admin:
        admin = User(
            username='admin',
            password=generate_password_hash('admin123'),
            name='管理员',
            department='办公室',
            role='admin',
            is_reporter=True,
            is_receiver=True
        )
        db.session.add(admin)
        db.session.commit()
        # 为管理员创建默认个人知识库
        personal_kb = KnowledgeBase.query.filter_by(owner_id=admin.id, type='personal').first()
        if not personal_kb:
            personal_kb = KnowledgeBase(name=f"{admin.name}的个人知识库", type='personal', owner_id=admin.id)
            db.session.add(personal_kb)
            db.session.commit()

    # 初始化简报系统默认数据
    if not BriefingSource.query.first():
        default_sources = [
            BriefingSource(name='人民日报', url='http://paper.people.com.cn/rmrb/', source_type='website', category='中央媒体'),
            BriefingSource(name='新华网', url='http://www.xinhuanet.com/', source_type='website', category='中央媒体'),
            BriefingSource(name='央视新闻', url='http://www.cctv.com/', source_type='website', category='中央媒体'),
            BriefingSource(name='中国政府网', url='http://www.gov.cn/', source_type='website', category='政府网站'),
            BriefingSource(name='光明日报', url='https://www.gmw.cn/', source_type='website', category='中央媒体'),
        ]
        for s in default_sources:
            db.session.add(s)
        db.session.commit()

    if not BriefingKeyword.query.first():
        default_keywords = [
            BriefingKeyword(text='经济', category='经济', color='#e74c3c'),
            BriefingKeyword(text='科技', category='科技', color='#3498db'),
            BriefingKeyword(text='民生', category='民生', color='#2ecc71'),
            BriefingKeyword(text='教育', category='教育', color='#f39c12'),
            BriefingKeyword(text='乡村振兴', category='农业', color='#27ae60'),
            BriefingKeyword(text='改革', category='政策', color='#8e44ad'),
            BriefingKeyword(text='创新', category='科技', color='#2980b9'),
        ]
        for k in default_keywords:
            db.session.add(k)
        db.session.commit()

    # 初始化默认角色
    if not Role.query.first():
        default_roles = [
            Role(name='admin', display_name='系统管理员', description='拥有所有权限', is_system=True,
                 permissions=json.dumps(['user_manage','role_manage','template_manage',
                                         'report_view','report_manage','task_manage',
                                         'knowledge_manage','ai_config','stats_view',
                                         'operation_log','briefing_manage'], ensure_ascii=False)),
            Role(name='manager', display_name='部门经理', description='管理本部门事务，可查看统计', is_system=True,
                 permissions=json.dumps(['report_view','report_manage','task_manage',
                                         'knowledge_manage','stats_view','briefing_manage'], ensure_ascii=False)),
            Role(name='reporter', display_name='信息报送员', description='负责信息上报和任务完成', is_system=False,
                 permissions=json.dumps(['report_view','report_submit','task_view',
                                         'knowledge_view'], ensure_ascii=False)),
            Role(name='employee', display_name='普通员工', description='基础功能访问', is_system=False,
                 permissions=json.dumps(['report_view','knowledge_view'], ensure_ascii=False)),
        ]
        for r in default_roles:
            db.session.add(r)
        db.session.commit()

    # 初始化默认公文模板
    if not DocTemplate.query.first():
        default_templates = [
            DocTemplate(name='请示（标准格式）', category='请示', file_type='txt',
                       description='标准行政请示文件格式',
                       tags='请示,行政,标准',
                       content="""【发文机关】
    XXX单位

    关于XXX的请示

    XXX（上级机关）：

    【正文】
    一、事项说明
    （说明请示事项的背景、原因及必要性）

    二、请示内容
    （具体的请示内容，要求明确、具体）

    三、相关情况
    （说明相关准备工作、方案等）

    妥否，请批示。

                                        XXX单位（盖章）
                                        XXXX年XX月XX日"""),
            DocTemplate(name='报告（工作总结）', category='报告', file_type='txt',
                       description='工作总结报告格式',
                       tags='报告,总结,工作',
                       content="""关于XXXX工作总结报告

    XXX（上级机关）：

    现将我单位XXXX工作情况报告如下：

    一、主要工作完成情况
    （说明本阶段主要工作的完成情况，包括具体数字、成效等）

    二、主要做法和经验
    （总结工作中的主要做法、好的经验和做法）

    三、存在的问题和不足
    （客观反映工作中存在的问题和短板）

    四、下一步工作打算
    （提出下一步工作思路和具体措施）

    以上报告，请审阅。

                                        XXX单位（盖章）
                                        XXXX年XX月XX日"""),
            DocTemplate(name='通知（标准格式）', category='通知', file_type='txt',
                       description='行政通知标准格式',
                       tags='通知,行政',
                       content="""关于XXXX的通知

    各有关单位：

    【正文】
    根据XXX要求，现将有关事项通知如下：

    一、XXXX
    （通知内容第一条）

    二、XXXX
    （通知内容第二条）

    三、其他事项
    （其他需要说明的事项）

    请各单位遵照执行，如有问题请及时联系。

    联系人：XXX，联系电话：XXXXXXXX

                                        XXX单位（盖章）
                                        XXXX年XX月XX日"""),
            DocTemplate(name='函（商洽事项）', category='函', file_type='txt',
                       description='行政函件格式',
                       tags='函,商洽',
                       content="""关于XXXX的函

    XXX（对方单位）：

    【正文】
    为（目的/原因），现就XXX事项函告（或：商洽）如下：

    一、XXXX

    二、XXXX

    请贵单位研究处理，并将处理结果函复。

                                        XXX单位（盖章）
                                        XXXX年XX月XX日"""),
            DocTemplate(name='会议纪要', category='纪要', file_type='txt',
                       description='会议纪要标准格式',
                       tags='会议,纪要',
                       content="""XXX会议纪要

    会议时间：XXXX年XX月XX日
    会议地点：XXXXXX
    主持人：XXX
    参会人员：XXX、XXX、XXX
    记录人：XXX

    会议主要内容如下：

    一、XXX情况通报
    （通报相关情况）

    二、研究讨论XXX事项
    （会议讨论的主要事项及结论）

    三、会议决定
    1. XXX（具体决定事项）
    2. XXX（具体决定事项）

    四、其他事项
    （其他讨论内容）

    本纪要经与会人员审阅，如无异议，自发出之日起生效。"""),
        ]
        admin_user = User.query.filter_by(username='admin').first()
        for t in default_templates:
            if admin_user:
                t.created_by = admin_user.id
            db.session.add(t)
        db.session.commit()

    # 初始化默认组织架构数据
    if not Organization.query.first():
        # 创建根机构
        root_org = Organization(
            name='智能服务办公平台',
            short_name='本单位',
            code='ROOT',
            org_type='unit',
            level=1,
            sort_order=0,
            description='系统默认根机构'
        )
        db.session.add(root_org)
        db.session.flush()  # 获取 root_org.id

        # 创建默认部门
        default_depts = [
            Department(name='办公室', code='OFFICE', org_id=root_org.id, dept_type='functional', sort_order=1, description='综合协调、文件收发、日常行政管理'),
            Department(name='人事部门', code='HR', org_id=root_org.id, dept_type='functional', sort_order=2, description='人员招聘、考核、培训及薪资管理'),
            Department(name='财务部门', code='FINANCE', org_id=root_org.id, dept_type='functional', sort_order=3, description='财务管理、预算执行、资产管理'),
            Department(name='业务部门', code='BUSINESS', org_id=root_org.id, dept_type='business', sort_order=4, description='核心业务开展与管理'),
            Department(name='信息技术部', code='IT', org_id=root_org.id, dept_type='support', sort_order=5, description='信息化建设与运维'),
        ]
        for d in default_depts:
            db.session.add(d)
        db.session.flush()

        # 从部门列表取办公室id
        office_dept = next((d for d in default_depts if d.code == 'OFFICE'), None)
        hr_dept = next((d for d in default_depts if d.code == 'HR'), None)
        biz_dept = next((d for d in default_depts if d.code == 'BUSINESS'), None)

        # 创建默认岗位
        default_positions = []
        if office_dept:
            default_positions += [
                Position(name='系统管理员', code='ADMIN', dept_id=office_dept.id, role_name='admin', level='manager', headcount=1, description='负责系统运维和管理'),
                Position(name='办公室主任', code='OFFICE_MGR', dept_id=office_dept.id, role_name='manager', level='manager', headcount=1, description='主持办公室全面工作'),
                Position(name='文秘', code='SECRETARY', dept_id=office_dept.id, role_name='reporter', level='staff', headcount=2, description='负责公文起草、信息报送'),
            ]
        if hr_dept:
            default_positions += [
                Position(name='人事主管', code='HR_MGR', dept_id=hr_dept.id, role_name='manager', level='supervisor', headcount=1, description='负责人事管理工作'),
                Position(name='人事专员', code='HR_STAFF', dept_id=hr_dept.id, role_name='employee', level='staff', headcount=2, description='日常人事事务处理'),
            ]
        if biz_dept:
            default_positions += [
                Position(name='业务经理', code='BIZ_MGR', dept_id=biz_dept.id, role_name='manager', level='manager', headcount=1, description='负责业务管理'),
                Position(name='业务员', code='BIZ_STAFF', dept_id=biz_dept.id, role_name='reporter', level='staff', headcount=5, description='负责具体业务'),
            ]
        for p in default_positions:
            db.session.add(p)
        db.session.flush()

        # 更新 admin 用户关联到办公室 + 系统管理员岗位
        admin_u = User.query.filter_by(username='admin').first()
        if admin_u and office_dept:
            admin_u.org_id = root_org.id
            admin_u.dept_id = office_dept.id
            admin_pos = next((p for p in default_positions if p.code == 'ADMIN'), None)
            if admin_pos:
                admin_u.position_id = admin_pos.id

        db.session.commit()


# ==================== 登录与主页 ====================
@app.route('/login', methods=['GET', 'POST'])
@csrf.exempt
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')

        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            session.permanent = True
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        flash('用户名或密码错误')

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

#@app.route('/')
#@login_required
#def index():
#    return render_template('index.html')

# ===== 系统配置上下文处理器 =====
@app.context_processor
def inject_system_config():
    """向所有模板注入系统配置"""
    return {
        'sys_config': config_manager,
    }

@app.route('/')
@login_required
@cache.cached(timeout=300, key_prefix=lambda: f'index_{current_user.id}')
def index():
    style = current_user.home_style or 'default'
    
    # 统计数据
    stats = {
        'total_reports': SpecialReport.query.filter_by(reporter_id=current_user.id).count(),
        'pending_tasks': TaskSubmission.query.filter_by(user_id=current_user.id, status='submitted').count(),
        'total_knowledge': KnowledgeFile.query.filter_by(uploaded_by=current_user.id).count(),
        'ai_chats': ChatSession.query.filter_by(user_id=current_user.id).count(),
    }

    # 注意：recent_docs/common_tools/templates/todo_tasks 不再传递给模板
    # 所有首页已改为静态展示，无需动态数据
    return render_template(
        f'index_{style}.html',
        stats=stats
    )

@app.route('/set_home_style/<style>', methods=['GET'])
@app.route('/set_home_style', methods=['POST'])
@login_required
def set_home_style(style=None):
    if request.method == 'POST':
        style = request.form.get('style')
    
    if style and style in ['default', 'dark', 'anime', 'fresh', 'tech']:
        current_user.home_style = style
        db.session.commit()
        # 清除首页缓存
        cache.delete(f'index_{current_user.id}')
        flash('首页风格已更新')
    else:
        flash('无效的风格')
    return redirect(url_for('index'))

# ==================== 个人中心 ====================
@app.route('/personal_center')
@login_required
def personal_center():
    return render_template('personal_center.html', user=current_user)

@app.route('/personal_center/update', methods=['POST'])
@login_required
def update_personal_center():
    current_user.name = request.form.get('name')
    current_user.department = request.form.get('department')
    current_user.phone = request.form.get('phone')
    current_user.email = request.form.get('email')   # 新增邮箱
    db.session.commit()
    flash('信息更新成功')
    return redirect(url_for('personal_center'))

@app.route('/column_settings', methods=['GET', 'POST'])
@login_required
def column_settings():
    if request.method == 'POST':
        selected = request.form.getlist('columns')
        session['user_columns'] = selected
        flash('栏目设置已保存')
        return redirect(url_for('index'))
    return render_template('column_settings.html')

@app.route('/update_notification_settings', methods=['POST'])
@login_required
def update_notification_settings():
    current_user.email_notify = request.form.get('email_notify') == 'on'
    current_user.task_notify = request.form.get('task_notify') == 'on'
    current_user.system_notify = request.form.get('system_notify') == 'on'
    db.session.commit()
    flash('通知设置已保存')
    return redirect(url_for('personal_center'))

@app.route('/change_password', methods=['POST'])
@login_required
def change_password():
    old_password = request.form.get('old_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    if not check_password_hash(current_user.password, old_password):
        flash('原密码错误')
        return redirect(url_for('personal_center'))
    
    if new_password != confirm_password:
        flash('两次输入的新密码不一致')
        return redirect(url_for('personal_center'))
    
    if len(new_password) < 6:
        flash('密码长度至少6位')
        return redirect(url_for('personal_center'))
    
    current_user.password = generate_password_hash(new_password)
    db.session.commit()
    flash('密码修改成功，请重新登录')
    return redirect(url_for('logout'))

# ==================== 智能办公模块 ====================

@app.route('/smart_office/document_writing', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def document_writing():
    if request.method == 'POST':
        # 处理AJAX请求
        data = request.get_json()
        template = data.get('template')
        keywords = data.get('keywords')
        model_id = data.get('model_id')
        selected_kb_ids = data.get('knowledge_bases', [])

        # 获取模型配置
        model_config = AIModelConfig.query.get(model_id) if model_id else None
        if not model_config:
            return jsonify({'error': '请选择有效的AI模型'}), 400

        # ===== P0: RAG 智能检索 =====
        kb_refs = [f"kb_{kid}" for kid in selected_kb_ids]
        knowledge_context, _ = get_knowledge_context_for_ai(
            kb_refs, keywords or template, max_snippets=3, max_chars=2000
        )

        prompt = f"请根据模板【{template}】和关键词【{keywords}】撰写一份规范的公文。"
        answer = generate_ai_response(model_config, prompt, knowledge_context)
        return jsonify({'result': answer})

    # GET请求渲染页面
    models = AIModelConfig.query.filter_by(is_active=True).all()
    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    policy_kbs = KnowledgeBase.query.filter_by(type='policy').all()
    return render_template('smart_office/document_writing.html', models=models, personal_kb=personal_kb, shared_kbs=shared_kbs, policy_kbs=policy_kbs)

@app.route('/smart_office/document_polish', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def document_polish():
    if request.method == 'POST':
        data = request.get_json()
        content = data.get('content')
        model_id = data.get('model_id')
        selected_kb_ids = data.get('knowledge_bases', [])

        model_config = AIModelConfig.query.get(model_id) if model_id else None
        if not model_config:
            return jsonify({'error': '请选择有效的AI模型'}), 400

        # ===== P0: RAG 智能检索 =====
        kb_refs = [f"kb_{kid}" for kid in selected_kb_ids]
        knowledge_context, _ = get_knowledge_context_for_ai(
            kb_refs, content or '', max_snippets=3, max_chars=2000
        )

        prompt = f"请润色以下文本，使其表达更流畅、语言更精炼：\n{content}"
        answer = generate_ai_response(model_config, prompt, knowledge_context)
        return jsonify({'result': answer})

    models = AIModelConfig.query.filter_by(is_active=True).all()
    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    policy_kbs = KnowledgeBase.query.filter_by(type='policy').all()
    return render_template('smart_office/document_polish.html', models=models, personal_kb=personal_kb, shared_kbs=shared_kbs, policy_kbs=policy_kbs)


@app.route('/smart_office/document_proofread', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def document_proofread():
    if request.method == 'POST':
        # 文件上传处理
        file = request.files.get('file')
        text_content = ''
        if file:
            filepath = save_upload_file(file, 'proofread')
            if filepath:
                if filepath.endswith('.txt'):
                    with open(filepath, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                elif filepath.endswith('.pdf'):
                    text_content = pdf_to_text(filepath)
                elif filepath.endswith('.docx'):
                    import docx
                    doc = docx.Document(filepath)
                    text_content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    text_content = '不支持的文件类型'
        else:
            text_content = request.form.get('content', '')

        model_id = request.form.get('model_id')
        selected_kb_ids = request.form.getlist('knowledge_bases')
        scheme = request.form.get('scheme', 'standard')

        model_config = AIModelConfig.query.get(model_id) if model_id else None
        if not model_config:
            return jsonify({'error': '请选择有效的AI模型'}), 400

        knowledge_context = ''
        if selected_kb_ids:
            kb_refs = [f"kb_{kid}" for kid in selected_kb_ids]
            knowledge_context, _ = get_knowledge_context_for_ai(
                kb_refs, text_content or '', max_snippets=3, max_chars=2000
            )

        prompt = f"请校对以下文本，找出错别字、语法错误、搭配错误，并给出修改建议。校对方案：{scheme}\n{text_content}"
        answer = generate_ai_response(model_config, prompt, knowledge_context)
        return jsonify({'result': answer})

    models = AIModelConfig.query.filter_by(is_active=True).all()
    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    policy_kbs = KnowledgeBase.query.filter_by(type='policy').all()
    return render_template('smart_office/document_proofread.html', models=models, personal_kb=personal_kb, shared_kbs=shared_kbs, policy_kbs=policy_kbs)

@app.route('/smart_office/suggestion', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def suggestion():
    if request.method == 'POST':
        file = request.files.get('file')
        text_content = ''
        if file:
            filepath = save_upload_file(file, 'suggestion')
            if filepath:
                if filepath.endswith('.txt'):
                    with open(filepath, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                elif filepath.endswith('.pdf'):
                    text_content = pdf_to_text(filepath)
                elif filepath.endswith('.docx'):
                    import docx
                    doc = docx.Document(filepath)
                    text_content = '\n'.join([p.text for p in doc.paragraphs])
        else:
            text_content = request.form.get('content', '')

        department_duty = request.form.get('department_duty')
        model_id = request.form.get('model_id')
        selected_kb_ids = request.form.getlist('knowledge_bases')

        model_config = AIModelConfig.query.get(model_id) if model_id else None
        if not model_config:
            return jsonify({'error': '请选择有效的AI模型'}), 400

        knowledge_context = ''
        if selected_kb_ids:
            kb_refs = [f"kb_{kid}" for kid in selected_kb_ids]
            knowledge_context, _ = get_knowledge_context_for_ai(
                kb_refs, text_content or '', max_snippets=3, max_chars=2000
            )

        prompt = f"请根据以下内容，生成公文摘要和拟办意见。部门职责：{department_duty}\n{text_content}"
        answer = generate_ai_response(model_config, prompt, knowledge_context)
        return jsonify({'result': answer})

    models = AIModelConfig.query.filter_by(is_active=True).all()
    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    policy_kbs = KnowledgeBase.query.filter_by(type='policy').all()
    return render_template('smart_office/suggestion.html', models=models, personal_kb=personal_kb, shared_kbs=shared_kbs, policy_kbs=policy_kbs)

@app.route('/smart_office/template_library')
@login_required
def template_library():
    templates = ['通知模板', '报告模板', '请示模板', '会议纪要模板', '决定模板', '批复模板']
    return render_template('smart_office/template_library.html', templates=templates)

@app.route('/smart_office/meeting_minutes', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def meeting_minutes():
    if request.method == 'POST':
        audio = request.files.get('audio')
        if audio:
            filepath = save_upload_file(audio, 'audio')
            text = audio_to_text(filepath)  # 模拟转写
        else:
            text = request.form.get('content', '')

        template = request.form.get('template')
        model_id = request.form.get('model_id')
        selected_kb_ids = request.form.getlist('knowledge_bases')

        model_config = AIModelConfig.query.get(model_id) if model_id else None
        if not model_config:
            return jsonify({'error': '请选择有效的AI模型'}), 400

        knowledge_context = ''
        if selected_kb_ids:
            files = KnowledgeFile.query.filter(
                KnowledgeFile.knowledge_base_id.in_(selected_kb_ids),
                KnowledgeFile.status == 'approved'
            ).all()
            for f in files:
                if f.file_path.endswith('.txt'):
                    try:
                        with open(f.file_path, 'r', encoding='utf-8') as txt:
                            content = txt.read()
                            # 使用配置的长度限制
                            max_length = config_manager.get("document.max_ai_sample_length", 5000)
                            if max_length == -1:
                                max_length = len(content)  # -1表示无限制
                            if max_length > 0 and len(content) > max_length:
                                knowledge_context += f"\n【文件：{f.original_name}】\n{content[:max_length]}...\n"
                            else:
                                knowledge_context += f"\n【文件：{f.original_name}】\n{content}\n"
                    except:
                        pass

        prompt = f"请根据以下录音转写内容，使用{template}模板生成会议纪要：\n{text}"
        answer = generate_ai_response(model_config, prompt, knowledge_context)
        return jsonify({'result': answer})

    models = AIModelConfig.query.filter_by(is_active=True).all()
    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    policy_kbs = KnowledgeBase.query.filter_by(type='policy').all()
    return render_template('smart_office/meeting_minutes.html', models=models, personal_kb=personal_kb, shared_kbs=shared_kbs, policy_kbs=policy_kbs)

# ============================================================
# 文档转换全局状态
# ============================================================
doc_convert_tasks = {}   # task_id -> {status, progress, message, result, error, started_at}
# ============================================================
# 文档转换 - 基于 transdoc9.py 的 LibreOffice 方案
# ============================================================

def _find_soffice():
    """查找 LibreOffice 可执行文件路径"""
    paths = [
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/opt/libreoffice/program/soffice",
    ]
    for p in paths:
        if os.path.exists(p):
            return p
    raise FileNotFoundError("未检测到 LibreOffice，请先安装")


def _do_convert(input_file, output_format, output_dir):
    """
    执行文件转换，参考 transdoc9.py 的 LibreOfficeController.convert()
    完全重写，解决所有权限问题
    """
    import subprocess, shutil, tempfile, sys, time

    soffice_path = _find_soffice()

    # 1. 创建安全的工作环境
    work_dir = tempfile.mkdtemp(prefix='docconv_')
    temp_input = os.path.join(work_dir, os.path.basename(input_file))

    try:
        # 2. 复制文件到工作目录（解决原始文件权限问题）
        shutil.copy2(input_file, temp_input)

        # 3. 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        # 4. 构造 LibreOffice 命令
        cmd = [
            soffice_path,
            "--headless",
            "--norestore",
            "--nodefault",
            "--nologo",
            "--convert-to", output_format,
            "--outdir", work_dir,
            temp_input,
        ]

        # 5. Windows 特殊处理
        if sys.platform == "win32":
            CREATE_NO_WINDOW = 0x08000000
            kwargs = {
                'creationflags': CREATE_NO_WINDOW,
                'shell': True,
            }
        else:
            kwargs = {}

        # 6. 执行转换（5分钟超时）
        result = subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=300,
            **kwargs,
        )

        # 7. 查找生成的文件
        base_name = os.path.splitext(os.path.basename(temp_input))[0]
        output_files = [
            f for f in os.listdir(work_dir)
            if f.startswith(base_name) and f != os.path.basename(temp_input)
        ]

        if not output_files:
            raise RuntimeError("未生成输出文件，可能是格式不支持")

        # 8. 移动结果文件到输出目录
        output_file = output_files[0]
        dst = os.path.join(output_dir, output_file)
        if os.path.exists(dst):
            try:
                os.remove(dst)
            except PermissionError:
                # 如果无法删除，尝试重命名
                name_parts = os.path.splitext(dst)
                dst = f"{name_parts[0]}_{int(time.time())}{name_parts[1]}"
        shutil.move(os.path.join(work_dir, output_file), dst)

        return dst, output_file

    except subprocess.TimeoutExpired:
        raise RuntimeError("转换超时，请检查文件是否损坏")
    except subprocess.CalledProcessError as e:
        stderr_str = e.stderr.decode('utf-8', errors='ignore') if e.stderr else ''
        stdout_str = e.stdout.decode('utf-8', errors='ignore') if e.stdout else ''
        err_detail = stderr_str or stdout_str or f"LibreOffice 转换失败（返回码 {e.returncode}）"
        if 'too large' in err_detail.lower():
            err_detail = '文件过大，建议分拆后重试'
        elif 'password' in err_detail.lower() or 'encrypt' in err_detail.lower():
            err_detail = '文件已加密，请先解除密码保护'
        elif e.returncode == 77:
            err_detail = '格式不支持或文件已损坏'
        raise RuntimeError(err_detail)
    except Exception as e:
        raise RuntimeError(f"转换失败: {str(e)}")
    finally:
        # 9. 清理工作目录
        try:
            shutil.rmtree(work_dir, ignore_errors=True)
        except Exception:
            pass


@app.route('/smart_office/pdf_convert', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def pdf_convert():
    """
    文档全格式转换 - 基于 LibreOffice headless
    参考 transdoc9.py 的 DocumentProcessor 实现
    """
    import subprocess, os, tempfile, shutil, sys
    from datetime import datetime

    SUPPORTED_FORMATS = {
        'doc': 'Word 97-2003',
        'docx': 'Word 2007+',
        'pdf': 'PDF',
        'xls': 'Excel 97-2003',
        'xlsx': 'Excel 2007+',
        'ppt': 'PowerPoint 97-2003',
        'pptx': 'PowerPoint 2007+',
        'odt': 'OpenDocument Text',
        'ods': 'OpenDocument Spreadsheet',
        'odp': 'OpenDocument Presentation',
    }

    # 支持的转换类型映射 (源格式, 目标格式) -> 实际转换格式
    CONVERSION_MAP = {
        ('doc', 'docx'): 'docx',
        ('doc', 'pdf'): 'pdf',
        ('doc', 'odt'): 'odt',
        ('docx', 'pdf'): 'pdf',
        ('docx', 'doc'): 'doc',
        ('docx', 'odt'): 'odt',
        ('odt', 'docx'): 'docx',
        ('odt', 'pdf'): 'pdf',
        ('xls', 'xlsx'): 'xlsx',
        ('xls', 'pdf'): 'pdf',
        ('xls', 'ods'): 'ods',
        ('xlsx', 'pdf'): 'pdf',
        ('xlsx', 'xls'): 'xls',
        ('xlsx', 'ods'): 'ods',
        ('ods', 'xlsx'): 'xlsx',
        ('ods', 'pdf'): 'pdf',
        ('ppt', 'pptx'): 'pptx',
        ('ppt', 'pdf'): 'pdf',
        ('ppt', 'odp'): 'odp',
        ('pptx', 'pdf'): 'pdf',
        ('pptx', 'ppt'): 'ppt',
        ('pptx', 'odp'): 'odp',
        ('odp', 'pptx'): 'pptx',
        ('odp', 'pdf'): 'pdf',
    }

    error_msg = None
    success_file = None
    download_url = None

    if request.method == 'POST':
        file = request.files.get('doc_file')
        target_format = request.form.get('target_format', '').strip()

        if not file:
            error_msg = '请选择要转换的文件'
        elif not target_format:
            error_msg = '请选择目标格式'
        else:
            ext = os.path.splitext(file.filename)[1][1:].lower()
            if (ext, target_format) not in CONVERSION_MAP:
                error_msg = f'不支持从 {ext.upper()} 转换为 {target_format.upper()}'
            else:
                actual_format = CONVERSION_MAP[(ext, target_format)]
                try:
                    # 保存上传文件
                    upload_dir = os.path.join('uploads', 'doc_convert', datetime.now().strftime('%Y%m%d'))
                    os.makedirs(upload_dir, exist_ok=True)
                    input_path = os.path.join(upload_dir, file.filename)
                    file.save(input_path)

                    # 输出目录
                    download_dir = os.path.join('static', 'downloads', 'doc_convert')
                    os.makedirs(download_dir, exist_ok=True)

                    # 执行转换（参考 transdoc9.py）
                    final_path, output_filename = _do_convert(input_path, actual_format, download_dir)

                    # 清理上传文件
                    try:
                        os.remove(input_path)
                    except Exception:
                        pass

                    success_file = output_filename
                    download_url = '/' + final_path.replace('\\', '/')

                except FileNotFoundError as e:
                    error_msg = str(e)
                except Exception as e:
                    error_msg = str(e)

    # 构建 JavaScript 可用的转换映射：源格式 -> [目标格式列表]
    js_conversion_map = {}
    for (src, dst), actual_fmt in CONVERSION_MAP.items():
        if src not in js_conversion_map:
            js_conversion_map[src] = []
        if actual_fmt not in js_conversion_map[src]:
            js_conversion_map[src].append(actual_fmt)

    return render_template(
        'smart_office/pdf_convert.html',
        formats=SUPPORTED_FORMATS,
        conversion_map=js_conversion_map,
        error_msg=error_msg,
        success_file=success_file,
        download_url=download_url,
    )


# ==================== 知识管理模块 ====================

# ===== P0: SQLite FTS5 全文搜索 =====

def _create_fts_index_table():
    db.session.execute(db.text('''
        CREATE VIRTUAL TABLE IF NOT EXISTS knowledge_files_fts USING fts5(
            file_id UNINDEXED,
            title,
            content,
            kb_name,
            kb_type,
            file_type,
            tags
        )
    '''))


def init_fts_index(force_rebuild=False):
    """初始化并校验 FTS5 虚拟表，发现旧版错误结构时自动重建。"""
    try:
        if force_rebuild:
            db.session.execute(db.text("DROP TABLE IF EXISTS knowledge_files_fts"))
        _create_fts_index_table()
        db.session.commit()
        db.session.execute(db.text("SELECT COUNT(*) FROM knowledge_files_fts")).scalar()
        return True
    except Exception as e:
        db.session.rollback()
        logger.error(f"FTS初始化失败: {e}")
        if force_rebuild:
            return False

    try:
        db.session.execute(db.text("DROP TABLE IF EXISTS knowledge_files_fts"))
        _create_fts_index_table()
        db.session.commit()
        db.session.execute(db.text("SELECT COUNT(*) FROM knowledge_files_fts")).scalar()
        logger.warning("检测到旧版损坏的 knowledge_files_fts 结构，已自动重建空索引表")
        return True
    except Exception as rebuild_error:
        db.session.rollback()
        logger.error(f"FTS重建失败: {rebuild_error}")
        return False


# 全局标志：FTS是否已初始化
_fts_initialized = False


def _upsert_fts_index_record(kf):
    kb = KnowledgeBase.query.get(kf.knowledge_base_id)
    db.session.execute(db.text(
        "DELETE FROM knowledge_files_fts WHERE file_id = :fid"
    ), {'fid': kf.id})
    db.session.execute(db.text('''
        INSERT INTO knowledge_files_fts(file_id, title, content, kb_name, kb_type, file_type, tags)
        VALUES (:fid, :title, :content, :kb_name, :kb_type, :ftype, :tags)
    '''), {
        'fid': kf.id,
        'title': kf.original_name or kf.filename,
        'content': kf.content_text or '',
        'kb_name': kb.name if kb else '',
        'kb_type': kb.type if kb else '',
        'ftype': kf.file_type or '',
        'tags': kf.tags or ''
    })


def rebuild_fts_index():
    """按现有 KnowledgeFile 数据重建全文索引。"""
    global _fts_initialized
    if not init_fts_index(force_rebuild=True):
        return False

    try:
        rows = KnowledgeFile.query.all()
        for kf in rows:
            kb = KnowledgeBase.query.get(kf.knowledge_base_id)
            db.session.execute(db.text('''
                INSERT INTO knowledge_files_fts(file_id, title, content, kb_name, kb_type, file_type, tags)
                VALUES (:fid, :title, :content, :kb_name, :kb_type, :ftype, :tags)
            '''), {
                'fid': kf.id,
                'title': kf.original_name or kf.filename,
                'content': kf.content_text or '',
                'kb_name': kb.name if kb else '',
                'kb_type': kb.type if kb else '',
                'ftype': kf.file_type or '',
                'tags': kf.tags or ''
            })
        db.session.commit()
        _fts_initialized = True
        logger.info(f"知识库 FTS 索引已重建，共同步 {len(rows)} 条记录")
        return True
    except Exception as e:
        db.session.rollback()
        _fts_initialized = False
        logger.error(f"FTS索引重建失败: {e}")
        return False


def _ensure_fts():
    """延迟初始化 FTS 表，并在旧索引损坏或缺失时自动重建。"""
    global _fts_initialized
    if _fts_initialized:
        return True

    if not init_fts_index():
        return rebuild_fts_index()

    try:
        index_count = db.session.execute(db.text("SELECT COUNT(*) FROM knowledge_files_fts")).scalar() or 0
        source_count = KnowledgeFile.query.count()
        if source_count > 0 and index_count < source_count:
            return rebuild_fts_index()
        _fts_initialized = True
        return True
    except Exception as e:
        db.session.rollback()
        logger.error(f"FTS健康检查失败: {e}")
        return rebuild_fts_index()


def update_fts_index(kf):
    """将 KnowledgeFile 记录同步到 FTS 索引"""
    if not _ensure_fts():
        return
    try:
        _upsert_fts_index_record(kf)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.error(f"FTS索引更新失败: {e}")
        rebuild_fts_index()


def delete_fts_index(file_id):
    """从 FTS 索引中删除记录"""
    if not _ensure_fts():
        return
    try:
        db.session.execute(db.text(
            "DELETE FROM knowledge_files_fts WHERE file_id = :fid"
        ), {'fid': file_id})
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.error(f"FTS索引删除失败: {e}")
        if rebuild_fts_index():
            try:
                db.session.execute(db.text(
                    "DELETE FROM knowledge_files_fts WHERE file_id = :fid"
                ), {'fid': file_id})
                db.session.commit()
            except Exception as retry_error:
                db.session.rollback()
                logger.error(f"FTS索引删除重试失败: {retry_error}")


def _search_knowledge_like(keyword, kb_type=None, file_type=None, kb_ids=None,
                           user_id=None, page=1, page_size=20):
    """当 FTS 不可用时，回退到普通 LIKE 检索，保证搜索结果可用。"""
    terms = [t for t in re.split(r'\s+', keyword.strip()) if t]
    pattern = f"%{keyword.strip()}%"

    query = KnowledgeFile.query.join(KnowledgeBase, KnowledgeBase.id == KnowledgeFile.knowledge_base_id)
    query = query.filter(KnowledgeFile.status == 'approved')

    if kb_type and kb_type != 'all':
        query = query.filter(KnowledgeBase.type == kb_type)
    if file_type and file_type != 'all':
        query = query.filter(KnowledgeFile.file_type == file_type)
    if kb_ids:
        query = query.filter(KnowledgeFile.knowledge_base_id.in_(kb_ids))
    if user_id:
        query = query.filter(db.or_(KnowledgeBase.type != 'personal', KnowledgeBase.owner_id == user_id))

    query = query.filter(db.or_(
        KnowledgeFile.original_name.like(pattern),
        KnowledgeFile.filename.like(pattern),
        KnowledgeFile.content_text.like(pattern),
        KnowledgeFile.summary.like(pattern),
        KnowledgeFile.tags.like(pattern)
    ))

    total = query.count()
    files = query.order_by(KnowledgeFile.upload_time.desc()).offset((page - 1) * page_size).limit(page_size).all()

    results = []
    for kf in files:
        kb = KnowledgeBase.query.get(kf.knowledge_base_id)
        title = kf.original_name or kf.filename or '未命名文档'
        
        # 清理标题
        if title and '_' in title:
            # 尝试移除时间戳和UUID
            clean_title = re.sub(r'^\d+_', '', title)  # 移除数字前缀
            clean_title = re.sub(r'^\d{14}_', '', clean_title)  # 移除14位时间戳
            clean_title = re.sub(r'_[a-f0-9]{8,}', '', clean_title)  # 移除UUID
            if clean_title and len(clean_title) > 2:
                title = clean_title
        
        snippet_source = kf.content_text or kf.summary or ''
        snippet = (snippet_source[:300] + '...') if len(snippet_source) > 300 else snippet_source
        
        # 高亮搜索词
        for term in terms:
            if term and len(term) > 1:
                title = title.replace(term, f'<mark>{term}</mark>')
                snippet = snippet.replace(term, f'<mark>{term}</mark>')
        
        results.append({
            'file_id': kf.id,
            'title': title,
            'snippet': snippet or '（无预览内容）',
            'kb_name': kb.name if kb else '',
            'kb_type': kb.type if kb else '',
            'file_type': kf.file_type or '',
            'tags': kf.tags or '',
            'rank': 1.0,
            'match_type': 'like'
        })

    return {'results': results, 'total': total, 'mode': 'like_fallback'}


def search_knowledge_fts(keyword, kb_type=None, file_type=None, kb_ids=None,
                          user_id=None, page=1, page_size=20):
    """
    FTS5 全文搜索
    返回: {'results': [...], 'total': int}
    """
    if not keyword or not keyword.strip():
        return {'results': [], 'total': 0}

    keyword = keyword.strip()
    terms = [t for t in re.split(r'\s+', keyword) if t]
    if not terms:
        return {'results': [], 'total': 0}

    if not _ensure_fts():
        logger.warning('知识库 FTS 不可用，已回退到 LIKE 检索')
        return _search_knowledge_like(keyword, kb_type, file_type, kb_ids, user_id, page, page_size)

    fts_terms = []
    for term in terms:
        cleaned = re.sub("[\\\"':*]", ' ', term).strip()

        if cleaned:
            fts_terms.append(f'"{cleaned}"')

    if not fts_terms:
        return _search_knowledge_like(keyword, kb_type, file_type, kb_ids, user_id, page, page_size)

    fts_query = ' AND '.join(fts_terms)
    params = {'query': fts_query}
    sql_body = """
        SELECT knowledge_files_fts.file_id,
               knowledge_files_fts.title,
               substr(knowledge_files_fts.content, 1, 300) as snippet,
               knowledge_files_fts.kb_name,
               knowledge_files_fts.kb_type,
               knowledge_files_fts.file_type,
               knowledge_files_fts.tags,
               bm25(knowledge_files_fts) as rank
        FROM knowledge_files_fts
        JOIN knowledge_files kf ON kf.id = knowledge_files_fts.file_id
        JOIN knowledge_bases kb ON kb.id = kf.knowledge_base_id
        WHERE knowledge_files_fts MATCH :query
    """
    count_sql = """
        SELECT COUNT(*)
        FROM knowledge_files_fts
        JOIN knowledge_files kf ON kf.id = knowledge_files_fts.file_id
        JOIN knowledge_bases kb ON kb.id = kf.knowledge_base_id
        WHERE knowledge_files_fts MATCH :query
    """

    where_parts = []
    if kb_type and kb_type != 'all':
        where_parts.append("kb.type = :kb_type")
        params['kb_type'] = kb_type
    if file_type and file_type != 'all':
        where_parts.append("knowledge_files_fts.file_type = :file_type")
        params['file_type'] = file_type
    if kb_ids:
        placeholders = ','.join(f':kb{i}' for i in range(len(kb_ids)))
        where_parts.append(f"kf.knowledge_base_id IN ({placeholders})")
        for i, kid in enumerate(kb_ids):
            params[f'kb{i}'] = kid
    if user_id:
        where_parts.append("(kb.type != 'personal' OR kb.owner_id = :uid)")
        params['uid'] = user_id
    where_parts.append("kf.status = 'approved'")

    if where_parts:
        where_sql = ' AND ' + ' AND '.join(where_parts)
        sql_body += where_sql
        count_sql += where_sql

    sql_body += " ORDER BY rank LIMIT :limit OFFSET :offset"
    params['limit'] = page_size
    params['offset'] = (page - 1) * page_size

    try:
        rows = db.session.execute(db.text(sql_body), params).fetchall()
        total = db.session.execute(db.text(count_sql), params).scalar() or 0

        results = []
        for row in rows:
            title = row.title or ''
            snippet = row.snippet or ''
            
            # 如果标题为空或太短，尝试从数据库中获取更好的标题
            if not title or len(title.strip()) < 3:
                try:
                    kf = KnowledgeFile.query.get(row.file_id)
                    if kf and kf.original_name:
                        title = kf.original_name
                    elif kf and kf.filename:
                        # 从文件名中提取一个更好的标题
                        filename = kf.filename
                        # 移除时间戳前缀
                        clean_name = re.sub(r'^\d+_', '', filename)
                        clean_name = re.sub(r'^\d{14}_', '', clean_name)
                        # 移除UUID部分
                        clean_name = re.sub(r'_[a-f0-9]{8,}', '', clean_name)
                        # 移除扩展名
                        clean_name = re.sub(r'\.[a-z]{2,4}$', '', clean_name, flags=re.IGNORECASE)
                        if clean_name and len(clean_name) > 2:
                            title = clean_name
                except Exception:
                    pass
            
            # 高亮搜索词
            for term in terms:
                if term and len(term) > 1:
                    title = title.replace(term, f'<mark>{term}</mark>')
                    snippet = snippet.replace(term, f'<mark>{term}</mark>')
            
            results.append({
                'file_id': row.file_id,
                'title': title or '未命名文档',
                'snippet': snippet or '（无预览内容）',
                'kb_name': row.kb_name or '',
                'kb_type': row.kb_type or '',
                'file_type': row.file_type or '',
                'tags': row.tags or '',
                'rank': float(row.rank) if row.rank else 0,
                'match_type': 'fts'
            })

        if total == 0:
            fallback_result = _search_knowledge_like(keyword, kb_type, file_type, kb_ids, user_id, page, page_size)
            if fallback_result.get('total', 0) > 0:
                return fallback_result

        return {'results': results, 'total': total, 'mode': 'fts'}

    except Exception as e:
        db.session.rollback()
        logger.error(f"FTS搜索失败: {e}")
        rebuild_fts_index()
        return _search_knowledge_like(keyword, kb_type, file_type, kb_ids, user_id, page, page_size)



# ===== P0: 知识库搜索 API =====
@app.route('/knowledge/api/search', methods=['POST'])
@login_required
@csrf.exempt
def knowledge_api_search():
    """全文搜索 API"""
    data = request.get_json(silent=True) or {}
    keyword = (data.get('keyword') or data.get('query') or '').strip()
    kb_type = data.get('kb_type')  # personal / shared / policy / all
    file_type = data.get('file_type')  # pdf / docx / txt / all
    kb_ids = data.get('kb_ids')  # 指定知识库ID列表
    page = data.get('page', 1)

    if not keyword:
        return jsonify({'results': [], 'total': 0})

    result = search_knowledge_fts(
        keyword=keyword,
        kb_type=kb_type,
        file_type=file_type,
        kb_ids=kb_ids,
        user_id=current_user.id,
        page=page
    )
    return jsonify(result)


@app.route('/knowledge/api/search_page')
@login_required
def knowledge_search_page():
    """知识库全文搜索结果页"""
    keyword = request.args.get('keyword', '').strip()
    kb_type = request.args.get('kb_type', 'all')
    file_type = request.args.get('file_type', 'all')
    page = request.args.get('page', 1, type=int)

    results = []
    total = 0
    if keyword:
        result = search_knowledge_fts(
            keyword=keyword, kb_type=kb_type, file_type=file_type,
            user_id=current_user.id, page=page
        )
        results = result.get('results', [])
        total = result.get('total', 0)

    # 渲染搜索结果页面
    return render_template('knowledge/search_results.html',
                          keyword=keyword, results=results, total=total,
                          kb_type=kb_type, file_type=file_type, page=page)


# ===== P1: 知识库文件下载（按ID） =====
@app.route('/knowledge/download/<int:file_id>')
@login_required
def knowledge_download(file_id):
    """按文件ID下载"""
    kf = KnowledgeFile.query.get_or_404(file_id)
    kb = KnowledgeBase.query.get(kf.knowledge_base_id)
    if kb.type == 'personal' and kb.owner_id != current_user.id and current_user.role != 'admin':
        flash('无权下载')
        return redirect(url_for('personal_knowledge_base'))
    return redirect(url_for('download_file', filepath=kf.file_path.replace('\\', '/')))


# ===== P1: 知识库文件预览 =====
@app.route('/knowledge/preview/<int:file_id>')
@login_required
def knowledge_preview(file_id):
    """文件预览"""
    kf = KnowledgeFile.query.get_or_404(file_id)
    kb = KnowledgeBase.query.get(kf.knowledge_base_id)

    # 权限检查
    if kb.type == 'personal' and kb.owner_id != current_user.id and current_user.role != 'admin':
        flash('无权访问')
        return redirect(url_for('personal_knowledge_base'))

    # 记录浏览
    log = KnowledgeBrowseLog(user_id=current_user.id, file_id=file_id)
    db.session.add(log)
    db.session.commit()

    # 使用 config_manager 获取预览长度限制
    max_length = config_manager.get("document.max_preview_length", 100000)
    content = kf.content_text or ''

    if kf.file_type in ('txt', 'md'):
        # 纯文本直接渲染
        if max_length > 0 and len(content) > max_length:
            content = content[:max_length] + '\n\n...（内容过长，已截断）'
        return render_template('knowledge/preview_text.html', file=kf, content=content)

    elif kf.file_type == 'docx':
        # Word 转为 HTML 预览
        try:
            html = docx_to_html(kf.file_path)
            preview_len = config_manager.get("document.max_file_preview_length", 5000)
            if len(html) > preview_len:
                html = html[:preview_len] + '<p>...（内容过长，已截断）</p>'
            return render_template('knowledge/preview_docx.html', file=kf, html=html)
        except Exception as e:
            flash(f'预览失败: {e}')
            return redirect(url_for('knowledge_download', file_id=file_id))

    elif kf.file_type == 'pdf':
        # PDF 尝试显示提取的文本内容
        if content:
            truncated = False
            if max_length > 0 and len(content) > max_length:
                content = content[:max_length]
                truncated = True
            return render_template('knowledge/preview_pdf.html',
                                  file=kf, content=content, truncated=truncated,
                                  max_length=max_length)
        else:
            flash('该 PDF 未提取到文本内容，请下载查看')
            return redirect(url_for('download_file', filepath=kf.file_path.replace('\\', '/')))

    elif kf.file_type in ('xlsx', 'xls'):
        # Excel 显示提取的文本内容
        if content:
            truncated = False
            if max_length > 0 and len(content) > max_length:
                content = content[:max_length]
                truncated = True
            return render_template('knowledge/preview_pdf.html',
                                  file=kf, content=content, truncated=truncated,
                                  max_length=max_length)
        else:
            flash('该 Excel 未提取到文本内容，请下载查看')
            return redirect(url_for('download_file', filepath=kf.file_path.replace('\\', '/')))

    else:
        flash('该文件类型不支持预览，请下载查看')
        return redirect(url_for('download_file', filepath=kf.file_path.replace('\\', '/')))


# ===== P1: 收藏功能 =====
@app.route('/knowledge/favorite/<int:file_id>', methods=['POST'])
@login_required
@csrf.exempt
def toggle_favorite(file_id):
    """收藏 / 取消收藏"""
    kf = KnowledgeFile.query.get_or_404(file_id)
    existing = KnowledgeFavorite.query.filter_by(
        user_id=current_user.id, file_id=file_id
    ).first()
    if existing:
        db.session.delete(existing)
        db.session.commit()
        return jsonify({'status': 'removed', 'message': '已取消收藏'})
    else:
        fav = KnowledgeFavorite(user_id=current_user.id, file_id=file_id)
        db.session.add(fav)
        db.session.commit()
        return jsonify({'status': 'added', 'message': '已收藏'})


@app.route('/knowledge/favorite/by_fav_id/<int:fav_id>', methods=['POST'])
@login_required
@csrf.exempt
def remove_favorite_by_id(fav_id):
    """根据收藏ID删除收藏记录"""
    fav = KnowledgeFavorite.query.get_or_404(fav_id)
    if fav.user_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': '无权删除'}), 403
    db.session.delete(fav)
    db.session.commit()
    return jsonify({'status': 'removed', 'message': '已取消收藏'})


@app.route('/knowledge/my_favorites')
@login_required
def my_favorites():
    """我的收藏"""
    favorites = KnowledgeFavorite.query.filter_by(user_id=current_user.id).order_by(
        KnowledgeFavorite.created_at.desc()
    ).all()
    # 构建知识库ID到名称的映射
    kb_ids = set(f.file.knowledge_base_id for f in favorites if f.file)
    kb_map = {kb.id: kb for kb in KnowledgeBase.query.filter(KnowledgeBase.id.in_(kb_ids)).all()}
    return render_template('knowledge/my_favorites.html', favorites=favorites, knowledge_bases_map=kb_map)


@app.route('/knowledge/recent')
@login_required
def knowledge_recent():
    """最近浏览"""
    recents = KnowledgeBrowseLog.query.filter_by(user_id=current_user.id).order_by(
        KnowledgeBrowseLog.browsed_at.desc()
    ).limit(20).all()
    # 为每条记录补充 kb_name
    for rec in recents:
        if rec.file:
            kb = KnowledgeBase.query.get(rec.file.knowledge_base_id)
            rec.file.kb_name = kb.name if kb else '未知'
    return render_template('knowledge/recent_browse.html', recents=recents)


# ===== P1: AI 自动生成摘要和标签 =====
@app.route('/knowledge/api/auto_tags', methods=['POST'])
@login_required
@csrf.exempt
def auto_generate_tags():
    """AI 自动为文件生成摘要和标签"""
    data = request.get_json() or {}
    file_id = data.get('file_id')
    kf = KnowledgeFile.query.get_or_404(file_id)

    # 权限检查
    kb = KnowledgeBase.query.get(kf.knowledge_base_id)
    if kb.type == 'personal' and kb.owner_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': '无权操作'}), 403

    if not kf.content_text:
        return jsonify({'error': '文件无文本内容'}), 400

    # 使用配置的长度限制获取文本样本供AI生成摘要
    content = kf.content_text
    max_length = config_manager.get("document.max_ai_sample_length", 5000)
    if max_length == -1:
        max_length = len(content)  # -1表示无限制
    text_sample = content[:max_length] if max_length > 0 else content
    model_cfg = AIModelConfig.query.filter_by(is_active=True).first()
    if not model_cfg:
        return jsonify({'error': '无AI模型配置'}), 400

    prompt = f"""请为以下文档自动生成：
1. 一句话摘要（50字以内）
2. 3-5个标签（用逗号分隔）

文档内容：
{text_sample}...

请按以下JSON格式返回（不要加代码块）：
{{"summary": "摘要内容", "tags": "标签1,标签2,标签3"}}
"""
    answer = generate_ai_response(model_cfg, prompt, '')
    try:
        m = re.search(r'\{[^{}]*\}', answer, re.DOTALL)
        if m:
            import json
            info = json.loads(m.group())
            summary = info.get('summary', '')
            tags = info.get('tags', '')
        else:
            summary = answer[:100]
            tags = ''
    except Exception:
        summary = answer[:100]
        tags = ''

    kf.summary = summary
    kf.tags = tags
    db.session.commit()

    return jsonify({'summary': summary, 'tags': tags})


# ===== P0: 全文搜索入口（供 AI 模块调用）=====
def get_knowledge_context_for_ai(selected_kbs, question, max_snippets=5, max_chars=3000):
    """
    从选中的知识库中检索与问题最相关的片段（用于 RAG）
    selected_kbs: ['personal_1', 'shared_2', ...]
    question: 用户问题
    返回: (context文本, 引用列表)
    """
    context_parts = []
    citations = []

    for kb_ref in selected_kbs:
        parts = kb_ref.split('_', 1)
        if len(parts) < 2:
            continue
        kb_type, kb_id_str = parts[0], parts[1]
        try:
            kb_id = int(kb_id_str)
        except ValueError:
            continue

        kb = KnowledgeBase.query.get(kb_id)
        if not kb:
            continue

        # 权限判断
        if kb.type == 'personal' and kb.owner_id != current_user.id and current_user.role != 'admin':
            continue

        # FTS 搜索
        result = search_knowledge_fts(
            keyword=question,
            kb_type=kb_type,
            kb_ids=[kb_id],
            user_id=current_user.id,
            page=1,
            page_size=max_snippets
        )
        snippets = result.get('results', [])
        for snip in snippets:
            # 截取片段（取前max_chars字符）
            text = snip.get('snippet', '')[:max_chars]
            context_parts.append(f"【来源：{snip['title']}（{kb.name}）】\n{text}")
            citations.append({
                'file_id': snip['file_id'],
                'title': snip['title'],
                'kb_name': kb.name,
                'kb_type': kb_type
            })

    # 限制总长度
    full_context = '\n\n---\n\n'.join(context_parts)
    if len(full_context) > max_chars * max_snippets:
        full_context = full_context[:max_chars * max_snippets]

    return full_context, citations


# 个人知识库 - 兼容新旧模式
@app.route('/knowledge/personal')
@app.route('/knowledge/personal/<int:kb_id>')
@login_required
def personal_knowledge_base(kb_id=None):
    # 如果指定了kb_id，加载该知识库
    if kb_id:
        kb = KnowledgeBase.query.filter_by(id=kb_id, owner_id=current_user.id, type='personal').first()
        if not kb:
            flash('未找到指定的个人知识库')
            return redirect(url_for('personal_kb_management'))
    else:
        # 如果没有指定kb_id，使用第一个个人知识库（向后兼容）
        kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
        if not kb:
            # 如果用户没有任何个人知识库，创建一个默认的
            kb = KnowledgeBase(name=f"{current_user.name}的个人知识库", type='personal', owner_id=current_user.id)
            db.session.add(kb)
            db.session.commit()
    
    # 获取用户的所有个人知识库（用于侧边栏展示）
    user_all_personal_kbs = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').all()
    
    # 获取当前知识库的文件
    files = KnowledgeFile.query.filter_by(knowledge_base_id=kb.id).all()
    return render_template('knowledge/personal_knowledge_base_new.html', 
                          kb=kb, 
                          files=files,
                          user_personal_kbs=user_all_personal_kbs)

# ===== 个人知识库管理功能 =====

# 个人知识库管理 - 列表页
@app.route('/knowledge/personal/management')
@login_required
def personal_kb_management():
    """个人知识库管理页面 - 显示用户所有个人知识库"""
    knowledge_bases = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').all()
    return render_template('knowledge/personal_kb_management.html', knowledge_bases=knowledge_bases)

# 创建个人知识库
@app.route('/knowledge/personal/create', methods=['GET', 'POST'])
@login_required
def create_personal_kb():
    """创建个人知识库"""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash('请输入知识库名称')
            return render_template('knowledge/create_edit_personal_kb.html', mode='create')
        
        category = request.form.get('category', '')
        description = request.form.get('description', '')
        
        # 检查是否已存在同名知识库
        existing = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal', name=name).first()
        if existing:
            flash('已存在同名的个人知识库')
            return render_template('knowledge/create_edit_personal_kb.html', mode='create')
        
        # 创建新知识库
        kb = KnowledgeBase(
            name=name,
            type='personal',
            owner_id=current_user.id,
            is_public=False,  # 个人知识库默认不公开
            category=category,
            description=description
        )
        db.session.add(kb)
        db.session.commit()
        
        flash(f'个人知识库 "{name}" 创建成功')
        return redirect(url_for('personal_kb_management'))
    
    return render_template('knowledge/create_edit_personal_kb.html', mode='create')

# 编辑个人知识库
@app.route('/knowledge/personal/edit/<int:kb_id>', methods=['GET', 'POST'])
@login_required 
def edit_personal_kb(kb_id):
    """编辑个人知识库"""
    kb = KnowledgeBase.query.filter_by(id=kb_id, owner_id=current_user.id, type='personal').first()
    if not kb:
        flash('未找到指定的个人知识库')
        return redirect(url_for('personal_kb_management'))
    
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash('请输入知识库名称')
            return render_template('knowledge/create_edit_personal_kb.html', mode='edit', kb=kb)
        
        # 检查除当前知识库外是否已存在同名知识库
        existing = KnowledgeBase.query.filter(
            KnowledgeBase.owner_id == current_user.id,
            KnowledgeBase.type == 'personal',
            KnowledgeBase.name == name,
            KnowledgeBase.id != kb_id
        ).first()
        if existing:
            flash('已存在同名的个人知识库')
            return render_template('knowledge/create_edit_personal_kb.html', mode='edit', kb=kb)
        
        # 更新知识库信息
        kb.name = name
        kb.category = request.form.get('category', '')
        kb.description = request.form.get('description', '')
        db.session.commit()
        
        flash('知识库信息更新成功')
        return redirect(url_for('personal_kb_management'))
    
    return render_template('knowledge/create_edit_personal_kb.html', mode='edit', kb=kb)

# 删除个人知识库
@app.route('/knowledge/personal/delete/<int:kb_id>', methods=['POST'])
@login_required
@csrf.exempt
def delete_personal_kb(kb_id):
    """删除个人知识库"""
    kb = KnowledgeBase.query.filter_by(id=kb_id, owner_id=current_user.id, type='personal').first()
    if not kb:
        return jsonify({'error': '未找到指定的个人知识库'}), 404
    
    # 检查知识库是否为空
    file_count = KnowledgeFile.query.filter_by(knowledge_base_id=kb.id).count()
    if file_count > 0:
        return jsonify({'error': '知识库非空，请先删除或移动所有文件后再删除知识库'}), 400
    
    try:
        # 删除知识库
        db.session.delete(kb)
        db.session.commit()
        return jsonify({'status': 'ok', 'message': f'个人知识库 "{kb.name}" 已删除'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'删除失败: {str(e)}'}), 500

# 个人知识库分享/取消分享
@app.route('/knowledge/personal/share/<int:kb_id>', methods=['POST'])
@login_required
@csrf.exempt
def share_personal_kb(kb_id):
    """个人知识库转为共享知识库"""
    kb = KnowledgeBase.query.filter_by(id=kb_id, owner_id=current_user.id, type='personal').first()
    if not kb:
        return jsonify({'error': '未找到指定的个人知识库'}), 404
    
    try:
        # 将知识库设置为公开并转为共享类型
        kb.is_public = True
        kb.type = 'shared'
        db.session.commit()
        return jsonify({'status': 'ok', 'message': f'知识库 "{kb.name}" 已转为共享知识库'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'共享失败: {str(e)}'}), 500

# 取消分享共享知识库（转为个人知识库）
@app.route('/knowledge/shared/unshare/<int:kb_id>', methods=['POST'])
@login_required
@csrf.exempt
def unshare_kb(kb_id):
    """共享知识库转为个人知识库"""
    # 检查权限：只有创建者或管理员可以取消分享
    kb = KnowledgeBase.query.filter_by(id=kb_id, type='shared').first()
    if not kb:
        return jsonify({'error': '未找到指定的共享知识库'}), 404
    
    # 检查用户权限：创建者或管理员
    if kb.owner_id != current_user.id and not current_user.is_admin:
        return jsonify({'error': '权限不足，只有创建者或管理员可以取消分享'}), 403
    
    try:
        # 将知识库设置为私有并转为个人类型
        kb.is_public = False
        kb.type = 'personal'
        db.session.commit()
        return jsonify({'status': 'ok', 'message': f'知识库 "{kb.name}" 已转为个人知识库'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'取消分享失败: {str(e)}'}), 500




@app.route('/knowledge/personal/upload', methods=['POST'])
@login_required
def personal_upload():
    kb_id = request.form.get('kb_id')
    file = request.files.get('file')
    if file:
        filepath = save_upload_file(file, f'knowledge/{kb_id}')
        if filepath:
            # ===== P0: 自动提取文件内容 =====
            content_text = extract_file_content(filepath)
            file_size, file_type, word_count = get_file_metadata(filepath)
            category = request.form.get('category', '')
            kf = KnowledgeFile(
                filename=os.path.basename(filepath),
                original_name=file.filename,
                file_path=filepath,
                knowledge_base_id=kb_id,
                uploaded_by=current_user.id,
                status='approved',
                content_text=content_text,
                file_size=file_size,
                file_type=file_type,
                word_count=word_count,
                category=category
            )
            db.session.add(kf)
            db.session.commit()
            # ===== P0: 更新 FTS 索引 =====
            update_fts_index(kf)
            flash('上传成功，文件内容已自动提取')
    # 重定向回当前知识库
    return redirect(url_for('personal_knowledge_base', kb_id=kb_id) if kb_id else url_for('personal_knowledge_base'))

# ==================== 智能知识库 - 批量上传 API ====================

@app.route('/knowledge/api/batch_upload', methods=['POST'])
@login_required
@csrf.exempt
def knowledge_batch_upload():
    """
    智能批量上传 API
    支持多文件上传，自动提取关键词、标签、摘要
    """
    kb_id = request.form.get('kb_id')
    kb_type = request.form.get('kb_type', 'personal')  # personal, shared, policy
    
    if not kb_id:
        return jsonify({'error': '未指定知识库'}), 400
    
    # 权限检查
    kb = KnowledgeBase.query.get_or_404(kb_id)
    if kb.type == 'personal' and kb.owner_id != current_user.id:
        return jsonify({'error': '无权访问'}), 403
    
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': '未选择文件'}), 400
    
    results = {'success': [], 'failed': []}
    
    for file in files:
        try:
            # 保存文件
            filepath = save_upload_file(file, f'knowledge/{kb_id}')
            if not filepath:
                results['failed'].append({'name': file.filename, 'error': '保存失败'})
                continue
            
            # 智能提取
            content_text = smart_kb.extract_content(filepath)
            file_size, file_type, word_count = get_file_metadata(filepath)
            
            # 生成智能元数据
            keywords = smart_kb.generate_keywords(content_text)
            summary = smart_kb.generate_summary(content_text)
            tags = smart_kb.auto_tag(content_text)
            suggested_title = smart_kb.suggest_title(content_text, file.filename)
            
            # 创建记录
            kf = KnowledgeFile(
                filename=os.path.basename(filepath),
                original_name=file.filename,
                file_path=filepath,
                knowledge_base_id=kb_id,
                uploaded_by=current_user.id,
                status='approved' if kb.type == 'personal' or current_user.role == 'admin' else 'pending',
                content_text=content_text,
                file_size=file_size,
                file_type=file_type,
                word_count=word_count,
                tags=','.join(tags),
                summary=summary
            )
            db.session.add(kf)
            db.session.commit()
            
            # 更新索引
            update_fts_index(kf)
            
            results['success'].append({
                'id': kf.id,
                'name': file.filename,
                'title': suggested_title,
                'keywords': keywords,
                'tags': tags,
                'summary': summary
            })
            
        except Exception as e:
            results['failed'].append({'name': file.filename, 'error': str(e)})
    
    return jsonify({
        'success_count': len(results['success']),
        'failed_count': len(results['failed']),
        'results': results
    })


@app.route('/knowledge/api/analyze', methods=['POST'])
@login_required
@csrf.exempt
def knowledge_analyze_file():
    """
    分析文件内容，返回智能元数据（不上传）
    用于预览和编辑
    """
    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    
    try:
        # 临时保存
        temp_dir = os.path.join('uploads', 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, file.filename)
        file.save(temp_path)
        
        # 分析
        content = smart_kb.extract_content(temp_path)
        keywords = smart_kb.generate_keywords(content)
        summary = smart_kb.generate_summary(content)
        tags = smart_kb.auto_tag(content)
        suggested_title = smart_kb.suggest_title(content, file.filename)
        
        # 清理临时文件
        os.remove(temp_path)
        
        return jsonify({
            'title': suggested_title,
            'keywords': keywords,
            'tags': tags,
            'summary': summary,
            'word_count': len(content) if content else 0
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/knowledge/api/smart_search', methods=['POST'])
@login_required
@csrf.exempt
def knowledge_smart_search():
    """
    智能检索 API
    支持全文检索、相似度检索和混合检索。
    """
    data = request.get_json(silent=True) or {}
    query = (data.get('query') or data.get('keyword') or '').strip()
    kb_ids = data.get('kb_ids', [])
    search_type = data.get('type', 'hybrid')  # fulltext, keyword, vector, hybrid
    file_type = data.get('file_type')
    kb_type = data.get('kb_type')

    if not query:
        return jsonify({'results': [], 'total': 0})

    try:
        fulltext_result = search_knowledge_fts(
            keyword=query,
            kb_type=kb_type,
            file_type=file_type,
            kb_ids=kb_ids,
            user_id=current_user.id,
            page=1,
            page_size=10
        )

        if search_type in ['fulltext', 'keyword']:
            return jsonify(fulltext_result)

        vector_results = smart_kb.search_similar(query, kb_ids, top_k=10)
        vector_payload = [{
            'file_id': r['file'].id,
            'title': r['file'].original_name or r['file'].filename,
            'snippet': r['snippet'],
            'kb_name': KnowledgeBase.query.get(r['file'].knowledge_base_id).name if KnowledgeBase.query.get(r['file'].knowledge_base_id) else '',
            'kb_type': KnowledgeBase.query.get(r['file'].knowledge_base_id).type if KnowledgeBase.query.get(r['file'].knowledge_base_id) else '',
            'file_type': r['file'].file_type,
            'tags': r['file'].tags,
            'score': round(r['score'], 3),
            'match_type': 'vector'
        } for r in vector_results]

        if search_type == 'vector':
            return jsonify({'results': vector_payload, 'total': len(vector_payload)})

        merged = []
        seen = set()

        for item in fulltext_result.get('results', []):
            file_id = item.get('file_id')
            item['match_type'] = item.get('match_type', 'fulltext')
            if file_id not in seen:
                merged.append(item)
                seen.add(file_id)

        for item in vector_payload:
            file_id = item.get('file_id')
            if file_id not in seen:
                merged.append(item)
                seen.add(file_id)

        return jsonify({'results': merged, 'total': len(merged)})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/knowledge/personal/delete/<int:file_id>')
@login_required
def personal_delete_file(file_id):
    kf = KnowledgeFile.query.get_or_404(file_id)
    kb = KnowledgeBase.query.get(kf.knowledge_base_id)
    if kb.type == 'personal' and kb.owner_id == current_user.id:
        # ===== P0: 删除 FTS 索引 =====
        delete_fts_index(file_id)
        if os.path.exists(kf.file_path):
            os.remove(kf.file_path)
        db.session.delete(kf)
        db.session.commit()
        flash('删除成功')
    return redirect(url_for('personal_knowledge_base'))

# 共享知识库
@app.route('/knowledge/shared')
@login_required
def shared_knowledge_base():
    # 普通用户只能看到已审核通过的共享知识库
    if current_user.role in ['manager', 'admin']:
        knowledge_bases = KnowledgeBase.query.filter_by(type='shared').all()
    else:
        knowledge_bases = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    return render_template('knowledge/shared_knowledge_base.html', knowledge_bases=knowledge_bases)


@app.route('/knowledge/shared/delete/<int:kb_id>', methods=['POST'])
@login_required
@csrf.exempt
def delete_shared_kb(kb_id):
    """删除共享知识库（管理员或创建者）"""
    kb = KnowledgeBase.query.get_or_404(kb_id)
    if current_user.role not in ['manager', 'admin'] and kb.owner_id != current_user.id:
        return jsonify({'error': '无权删除'}), 403
    # 删除该知识库下所有文件的FTS索引
    for f in kb.files:
        delete_fts_index(f.id)
    # 删除知识库（级联删除文件）
    db.session.delete(kb)
    db.session.commit()
    return jsonify({'status': 'ok'})

@app.route('/knowledge/shared/create', methods=['GET', 'POST'])
@login_required
def create_shared_kb():
    if current_user.role not in ['manager', 'admin']:
        flash('权限不足')
        return redirect(url_for('shared_knowledge_base'))
    if request.method == 'POST':
        name = request.form.get('name')
        is_public = request.form.get('is_public') == 'on'
        category = request.form.get('category', '')
        description = request.form.get('description', '')
        kb = KnowledgeBase(name=name, type='shared', owner_id=current_user.id, is_public=is_public, category=category, description=description)
        db.session.add(kb)
        db.session.commit()
        flash('创建成功')
        return redirect(url_for('shared_knowledge_base'))
    return render_template('knowledge/create_shared_kb.html')

@app.route('/knowledge/shared/<int:kb_id>')
@login_required
def shared_knowledge_detail(kb_id):
    kb = KnowledgeBase.query.get_or_404(kb_id)
    if not (kb.is_public or current_user.role in ['manager', 'admin'] or kb.owner_id == current_user.id):
        flash('无权访问该知识库')
        return redirect(url_for('shared_knowledge_base'))
    files = KnowledgeFile.query.filter_by(knowledge_base_id=kb_id).all()
    return render_template('knowledge/shared_knowledge_base_new.html', kb=kb, files=files)

@app.route('/knowledge/shared/upload/<int:kb_id>', methods=['POST'])
@login_required
def shared_upload(kb_id):
    kb = KnowledgeBase.query.get_or_404(kb_id)
    if not (kb.is_public or current_user.role in ['manager', 'admin'] or kb.owner_id == current_user.id):
        flash('无权上传')
        return redirect(url_for('shared_knowledge_detail', kb_id=kb_id))
    file = request.files.get('file')
    if file:
        filepath = save_upload_file(file, f'shared/{kb_id}')
        if filepath:
            # ===== P0: 自动提取文件内容 =====
            content_text = extract_file_content(filepath)
            file_size, file_type, word_count = get_file_metadata(filepath)
            category = request.form.get('category', '')
            kf = KnowledgeFile(
                filename=os.path.basename(filepath),
                original_name=file.filename,
                file_path=filepath,
                knowledge_base_id=kb_id,
                uploaded_by=current_user.id,
                status='pending' if current_user.role != 'admin' else 'approved',
                content_text=content_text,
                file_size=file_size,
                file_type=file_type,
                word_count=word_count,
                category=category
            )
            db.session.add(kf)
            db.session.commit()
            # ===== P0: 更新 FTS 索引 =====
            update_fts_index(kf)
            flash('上传成功，待审核')
    return redirect(url_for('shared_knowledge_detail', kb_id=kb_id))

@app.route('/knowledge/shared/delete_file/<int:file_id>')
@login_required
def shared_delete_file(file_id):
    kf = KnowledgeFile.query.get_or_404(file_id)
    kb = KnowledgeBase.query.get(kf.knowledge_base_id)
    if current_user.role in ['manager', 'admin'] or kb.owner_id == current_user.id:
        # ===== P0: 删除 FTS 索引 =====
        delete_fts_index(file_id)
        if os.path.exists(kf.file_path):
            os.remove(kf.file_path)
        db.session.delete(kf)
        db.session.commit()
        flash('删除成功')
    return redirect(url_for('shared_knowledge_detail', kb_id=kb.id))

# 政策文件库（类似共享知识库，仅管理员可管理）
@app.route('/knowledge/policy')
@login_required
def policy_library():
    kbs = KnowledgeBase.query.filter_by(type='policy').all()
    return render_template('knowledge/policy_library.html', kbs=kbs)

@app.route('/knowledge/policy/create', methods=['GET', 'POST'])
@login_required
def create_policy():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('policy_library'))
    if request.method == 'POST':
        name = request.form.get('name')
        kb = KnowledgeBase(name=name, type='policy', owner_id=current_user.id, is_public=True)
        db.session.add(kb)
        db.session.commit()
        flash('创建成功')
        return redirect(url_for('policy_library'))
    return render_template('knowledge/create_policy.html')

@app.route('/knowledge/policy/<int:kb_id>')
@login_required
def policy_detail(kb_id):
    kb = KnowledgeBase.query.get_or_404(kb_id)
    if kb.type != 'policy':
        flash('无效的政策库')
        return redirect(url_for('policy_library'))
    files = KnowledgeFile.query.filter_by(knowledge_base_id=kb_id).all()
    return render_template('knowledge/policy_detail.html', kb=kb, files=files)


@app.route('/knowledge/policy/upload/<int:kb_id>', methods=['POST'])
@login_required
def policy_upload(kb_id):
    """政策库批量上传（支持多文件）"""
    kb = KnowledgeBase.query.get_or_404(kb_id)
    if kb.type != 'policy':
        flash('无效的政策库')
        return redirect(url_for('policy_library'))

    # 权限：仅管理员可上传
    if current_user.role != 'admin':
        flash('权限不足，仅管理员可上传政策文件')
        return redirect(url_for('policy_detail', kb_id=kb_id))

    files = request.files.getlist('files')
    if not files or all(f.filename == '' for f in files):
        flash('请选择要上传的文件')
        return redirect(url_for('policy_detail', kb_id=kb_id))

    uploaded_count = 0
    for file in files:
        if file.filename == '':
            continue
        filepath = save_upload_file(file, f'policy/{kb_id}')
        if not filepath:
            continue
        # 自动提取内容
        content_text = extract_file_content(filepath)
        file_size, file_type, word_count = get_file_metadata(filepath)
        category = request.form.get('category', '')

        kf = KnowledgeFile(
            filename=os.path.basename(filepath),
            original_name=file.filename,
            file_path=filepath,
            knowledge_base_id=kb_id,
            uploaded_by=current_user.id,
            status='approved',   # 政策库文件直接审核通过
            content_text=content_text,
            file_size=file_size,
            file_type=file_type,
            word_count=word_count,
            category=category,
        )
        db.session.add(kf)
        db.session.commit()
        # 更新 FTS 索引
        update_fts_index(kf)
        uploaded_count += 1

    if uploaded_count > 0:
        flash(f'成功上传 {uploaded_count} 个文件')
    else:
        flash('上传失败，请检查文件格式（支持 PDF、Word、TXT、Markdown）')

    return redirect(url_for('policy_detail', kb_id=kb_id))

# 文件下载（通用）
@app.route('/download/<path:filepath>')
@login_required
def download_file(filepath):
    # 安全：验证文件路径，防止路径遍历攻击
    from werkzeug.utils import secure_filename
    secure_name = secure_filename(filepath)
    if '..' in secure_name or secure_name.startswith('/'):
        flash('非法的文件路径')
        return redirect(url_for('index'))
    
    full_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_name)
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('文件不存在')
        return redirect(url_for('index'))
    
    return send_from_directory(app.config['UPLOAD_FOLDER'], secure_name, as_attachment=True)



@app.route('/qa/knowledge_search', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def knowledge_search():
    if request.method == 'POST':
        selected_kbs = [int(k) for k in request.form.getlist('knowledge_bases') if str(k).isdigit()]
        question = (request.form.get('question') or '').strip()

        if not question:
            return jsonify({'answer': '请输入搜索问题。', 'results': []})

        result = search_knowledge_fts(
            keyword=question,
            kb_ids=selected_kbs or None,
            user_id=current_user.id,
            page=1,
            page_size=5
        )

        if result.get('results'):
            answer_lines = [f"共找到 {result.get('total', 0)} 条相关结果，前 5 条如下："]
            for idx, item in enumerate(result['results'][:5], 1):
                answer_lines.append(
                    f"{idx}. {item.get('title', '未命名文件')}（{item.get('kb_name', '未知知识库')}）"
                )
            answer = '\n'.join(answer_lines)
        else:
            answer = f'在选中的知识库中未找到与"{question}"相关的记录。'

        return jsonify({'answer': answer, 'results': result.get('results', []), 'total': result.get('total', 0)})

    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    return render_template('intelligent_qa/knowledge_search.html', personal_kb=personal_kb, shared_kbs=shared_kbs)

@app.route('/qa/policy_search', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def policy_search():
    if request.method == 'POST':
        question = request.form.get('question')
        answer = f'在政策文件库中查询"{question}"的结果：暂无相关政策。'
        return jsonify({'answer': answer})
    return render_template('intelligent_qa/policy_search.html')

# ==================== 专项信息上报 - 报送方 ====================
@app.route('/special_report/reporter/report_list')
@login_required
def report_list():
    reports = SpecialReport.query.filter_by(reporter_id=current_user.id).order_by(SpecialReport.created_at.desc()).all()
    return render_template('special_report/reporter/report_list.html', reports=reports)

@app.route('/special_report/reporter/report_form', methods=['GET', 'POST'])
@login_required
def report_form():
    if request.method == 'POST':
        title = request.form.get('title')
        category = request.form.get('category')
        content = request.form.get('content')
        target = request.form.get('target_department')
        attachments = request.form.get('attachments')  # 简化，实际应处理文件
        report = SpecialReport(
            title=title, category=category, content=content, attachments=attachments,
            reporter_id=current_user.id, reporter_name=current_user.name,
            reporter_phone=current_user.phone, target_department=target,
            status='pending'
        )
        db.session.add(report)
        db.session.commit()
        flash('报送成功')
        return redirect(url_for('report_list'))
    # 从数据库动态读取启用中的部门，兼容无部门时的降级处理
    from models import Department as DeptModel
    db_depts = DeptModel.query.filter_by(is_active=True).order_by(DeptModel.sort_order, DeptModel.name).all()
    departments = [d.name for d in db_depts] if db_depts else ['办公室', '人事部', '财务部', '技术部']
    return render_template('special_report/reporter/report_form.html', departments=departments)

@app.route('/special_report/reporter/edit_report/<int:report_id>', methods=['GET', 'POST'])
@login_required
def edit_report(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    if report.reporter_id != current_user.id or report.status not in ['draft', 'rejected']:
        flash('无权编辑或状态不允许编辑')
        return redirect(url_for('report_list'))
    if request.method == 'POST':
        report.title = request.form.get('title')
        report.category = request.form.get('category')
        report.content = request.form.get('content')
        report.target_department = request.form.get('target_department')
        report.attachments = request.form.get('attachments')
        report.status = 'pending' if request.form.get('submit_action') == 'submit' else 'draft'
        db.session.commit()
        flash('保存成功')
        return redirect(url_for('report_list'))
    from models import Department as DeptModel
    db_depts = DeptModel.query.filter_by(is_active=True).order_by(DeptModel.sort_order, DeptModel.name).all()
    departments = [d.name for d in db_depts] if db_depts else ['办公室', '人事部', '财务部', '技术部']
    return render_template('special_report/reporter/report_form.html', report=report, departments=departments)

@app.route('/special_report/reporter/delete_report/<int:report_id>')
@login_required
def delete_report(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    if report.reporter_id == current_user.id and report.status in ['draft', 'rejected']:
        db.session.delete(report)
        db.session.commit()
        flash('删除成功')
    else:
        flash('无法删除')
    return redirect(url_for('report_list'))

@app.route('/special_report/reporter/withdraw_report/<int:report_id>')
@login_required
def withdraw_report(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    if report.reporter_id == current_user.id and report.status == 'pending':
        report.status = 'draft'
        db.session.commit()
        flash('已撤回')
    return redirect(url_for('report_list'))

# 约稿任务待办
@app.route('/special_report/reporter/task_todo')
@login_required
def task_todo():
    # 获取分配给当前用户的未完成（未报送）任务
    tasks_with_status = []
    all_tasks = AssignmentTask.query.filter_by(status='active').all()
    for t in all_tasks:
        if t.assigned_to and str(current_user.id) in t.assigned_to.split(','):
            submission = TaskSubmission.query.filter_by(task_id=t.id, user_id=current_user.id).first()
            if not submission or submission.status != 'submitted':
                # 携带签收状态
                signed = submission is not None and submission.status == 'signed'
                tasks_with_status.append({'task': t, 'signed': signed})
    return render_template('special_report/reporter/task_todo.html', tasks=tasks_with_status)

@app.route('/special_report/reporter/task_sign/<int:task_id>')
@login_required
def task_sign(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    # 检查是否已签收（存在提交记录）
    existing = TaskSubmission.query.filter_by(task_id=task.id, user_id=current_user.id).first()
    if existing:
        flash('您已签收该任务')
        return redirect(url_for('task_todo'))
    # 创建签收记录（status='signed' 表示已签收未报送）
    submission = TaskSubmission(
        task_id=task.id,
        user_id=current_user.id,
        report_id=None,
        status='signed'
    )
    db.session.add(submission)
    db.session.commit()
    flash('签收成功')
    return redirect(url_for('task_todo'))

@app.route('/special_report/reporter/task_submit/<int:task_id>', methods=['GET', 'POST'])
@login_required
def task_submit(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    # 检查是否已提交（已报送状态）
    existing_submission = TaskSubmission.query.filter_by(
        task_id=task.id, user_id=current_user.id, status='submitted').first()
    if existing_submission:
        flash('您已报送过该任务，如需追报请使用追报功能')
        return redirect(url_for('task_todo'))
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        content = request.form.get('content', '').strip()
        if not title or not content:
            flash('标题和内容不能为空')
            return render_template('special_report/reporter/task_submit.html', task=task)
        # 获取约稿创建者的部门名称
        creator = User.query.get(task.created_by)
        target_dept = creator.department if creator else str(task.created_by)
        report = SpecialReport(
            title=title, content=content, category=task.category,
            reporter_id=current_user.id, reporter_name=current_user.name,
            reporter_phone=current_user.phone, target_department=target_dept,
            status='pending'
        )
        db.session.add(report)
        db.session.commit()
        # 若已签收则更新状态，否则新建
        signed_sub = TaskSubmission.query.filter_by(
            task_id=task.id, user_id=current_user.id, status='signed').first()
        if signed_sub:
            signed_sub.report_id = report.id
            signed_sub.status = 'submitted'
        else:
            submission = TaskSubmission(task_id=task.id, user_id=current_user.id,
                                        report_id=report.id, status='submitted')
            db.session.add(submission)
        db.session.commit()
        flash('报送成功！等待审核')
        return redirect(url_for('task_todo'))
    return render_template('special_report/reporter/task_submit.html', task=task, now=datetime.now())

@app.route('/special_report/reporter/task_done')
@login_required
def task_done():
    submissions = TaskSubmission.query.filter_by(user_id=current_user.id).all()
    tasks = []
    for sub in submissions:
        task = AssignmentTask.query.get(sub.task_id)
        if task:
            tasks.append({'task': task, 'submission': sub})
    return render_template('special_report/reporter/task_done.html', tasks=tasks)

@app.route('/special_report/reporter/task_additional/<int:task_id>', methods=['GET', 'POST'])
@login_required
def task_additional(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')
        report = SpecialReport(
            title=title, content=content, category=task.category,
            reporter_id=current_user.id, reporter_name=current_user.name,
            reporter_phone=current_user.phone, target_department=str(task.created_by),
            status='pending'
        )
        db.session.add(report)
        db.session.commit()
        # 追加报送，不改变原任务状态
        flash('追报成功')
        return redirect(url_for('task_done'))
    return render_template('special_report/reporter/task_submit.html', task=task, is_additional=True, now=datetime.now())

# ==================== 专项信息上报 - 接收方 ====================
@app.route('/special_report/receiver/info_pool')
@login_required
def info_pool():
    # 获取发送给当前用户所在部门的所有信息（支持按状态筛选）
    status_filter = request.args.get('status', 'all')
    query = SpecialReport.query.filter_by(target_department=current_user.department)
    if status_filter != 'all':
        query = query.filter_by(status=status_filter)
    reports = query.order_by(SpecialReport.created_at.desc()).all()
    # 统计各状态数量
    total = SpecialReport.query.filter_by(target_department=current_user.department).count()
    pending_count = SpecialReport.query.filter_by(target_department=current_user.department, status='pending').count()
    adopted_count = SpecialReport.query.filter_by(target_department=current_user.department, status='adopted').count()
    rejected_count = SpecialReport.query.filter_by(target_department=current_user.department, status='rejected').count()
    return render_template('special_report/receiver/info_pool.html', reports=reports,
                           status_filter=status_filter, total=total,
                           pending_count=pending_count, adopted_count=adopted_count,
                           rejected_count=rejected_count)

@app.route('/special_report/receiver/info_detail/<int:report_id>')
@login_required
def info_detail(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    return render_template('special_report/receiver/info_detail.html', report=report)

@app.route('/special_report/receiver/edit_info/<int:report_id>', methods=['GET', 'POST'])
@login_required
def edit_info(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    if request.method == 'POST':
        report.title = request.form.get('title')
        report.content = request.form.get('content')
        report.category = request.form.get('category')
        report.attachments = request.form.get('attachments')
        db.session.commit()
        flash('修改成功')
        return redirect(url_for('info_pool'))
    return render_template('special_report/receiver/edit_info.html', report=report)

@app.route('/special_report/receiver/refine_report/<int:report_id>', methods=['GET', 'POST'])
@login_required
def refine_report(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    if request.method == 'POST':
        # 模拟AI精编，实际可调用AI接口
        new_content = f"【AI精编】\n{report.content}\n\n补充原因：{request.form.get('reason')}"
        # 创建新报告，标记为AI生成
        new_report = SpecialReport(
            title=f"【精编】{report.title}",
            content=new_content,
            category=report.category,
            reporter_id=current_user.id,
            reporter_name=current_user.name,
            reporter_phone=current_user.phone,
            target_department=report.target_department,
            status='pending'
        )
        db.session.add(new_report)
        db.session.commit()
        flash('精编保存成功')
        return redirect(url_for('info_pool'))
    return render_template('special_report/receiver/refine_report.html', report=report)

@app.route('/special_report/receiver/adopt_report/<int:report_id>')
@login_required
def adopt_report(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    report.status = 'adopted'
    db.session.commit()
    flash('已采用')
    # 若来自任务详情页则返回任务详情，否则返回信息池
    next_url = request.args.get('next', url_for('info_pool'))
    return redirect(next_url)

@app.route('/special_report/receiver/cancel_adopt/<int:report_id>')
@login_required
def cancel_adopt(report_id):
    """取消采用 - 将报告状态从 adopted 退回 pending"""
    report = SpecialReport.query.get_or_404(report_id)
    if report.status == 'adopted':
        report.status = 'pending'
        db.session.commit()
        flash('已取消采用，信息已退回待审核状态')
    else:
        flash('当前状态无法取消采用')
    next_url = request.args.get('next', url_for('info_pool'))
    return redirect(next_url)

@app.route('/special_report/receiver/return_report/<int:report_id>', methods=['GET', 'POST'])
@login_required
def return_report(report_id):
    report = SpecialReport.query.get_or_404(report_id)
    if request.method == 'POST':
        feedback = request.form.get('feedback')
        report.feedback = feedback
        report.status = 'rejected'
        db.session.commit()
        flash('已退回')
        next_url = request.form.get('next', url_for('info_pool'))
        return redirect(next_url)
    next_url = request.args.get('next', url_for('info_pool'))
    return render_template('special_report/receiver/return_report.html', report=report, next_url=next_url)

@app.route('/special_report/receiver/cancel_return/<int:report_id>')
@login_required
def cancel_return(report_id):
    """取消退回 - 将报告状态从 rejected 重置为 pending"""
    report = SpecialReport.query.get_or_404(report_id)
    if report.status == 'rejected':
        report.status = 'pending'
        report.feedback = None
        db.session.commit()
        flash('已取消退回，信息重新进入待审核队列')
    else:
        flash('当前状态无法取消退回')
    next_url = request.args.get('next', url_for('info_pool'))
    return redirect(next_url)

@app.route('/special_report/reporter/view_report/<int:report_id>')
@login_required
def view_report(report_id):
    """查看原稿（只读查看上报信息详情）"""
    report = SpecialReport.query.get_or_404(report_id)
    if report.reporter_id != current_user.id:
        flash('无权查看该信息')
        return redirect(url_for('report_list'))
    return render_template('special_report/reporter/view_report.html', report=report)

# 专项约稿管理
@app.route('/special_report/receiver/task_manage')
@login_required
def task_manage():
    tasks = AssignmentTask.query.filter_by(created_by=current_user.id).order_by(AssignmentTask.created_at.desc()).all()
    # 为每个任务附加统计信息
    tasks_info = []
    for t in tasks:
        # 被分配总人数
        assigned_count = len(t.assigned_to.split(',')) if t.assigned_to else 0
        # 已签收数（signed 或 submitted）
        signed_count = TaskSubmission.query.filter(
            TaskSubmission.task_id == t.id,
            TaskSubmission.status.in_(['signed', 'submitted', 'approved', 'rejected'])
        ).count()
        # 已报送数
        submitted_count = TaskSubmission.query.filter(
            TaskSubmission.task_id == t.id,
            TaskSubmission.status.in_(['submitted', 'approved'])
        ).count()
        tasks_info.append({
            'task': t,
            'assigned_count': assigned_count,
            'signed_count': signed_count,
            'submitted_count': submitted_count
        })
    return render_template('special_report/receiver/task_manage.html', tasks=tasks_info, now=datetime.now())

@app.route('/special_report/receiver/create_task', methods=['GET', 'POST'])
@login_required
def create_task():
    # 仅接收方可以创建约稿任务（可根据实际业务调整）
    if not current_user.is_receiver:
        flash('您没有权限创建约稿任务')
        return redirect(url_for('index'))

    if request.method == 'POST':
        try:
            title = request.form.get('title')
            category = request.form.get('category')
            description = request.form.get('description')
            start_time_str = request.form.get('start_time')
            end_time_str = request.form.get('end_time')
            urgency = request.form.get('urgency')
            assigned_to = ','.join(request.form.getlist('assigned_to'))

            # 检查必要字段
            if not all([title, category, description, start_time_str, end_time_str, urgency, assigned_to]):
                flash('请填写所有必填项')
                return redirect(url_for('create_task'))

            start_time = datetime.strptime(start_time_str, '%Y-%m-%dT%H:%M')
            end_time = datetime.strptime(end_time_str, '%Y-%m-%dT%H:%M')

            task = AssignmentTask(
                title=title, category=category, description=description,
                start_time=start_time, end_time=end_time, urgency=urgency,
                assigned_to=assigned_to, created_by=current_user.id, status='active'
            )
            db.session.add(task)
            db.session.commit()
            flash('约稿任务创建成功')
            return redirect(url_for('task_manage'))

        except ValueError as e:
            flash(f'日期时间格式错误：{e}')
            return redirect(url_for('create_task'))
        except Exception as e:
            flash(f'创建失败：{str(e)}')
            return redirect(url_for('create_task'))

    users = User.query.filter(
        db.or_(User.is_active == True, User.is_active == None)
    ).order_by(User.dept_id, User.name).all()
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    return render_template('special_report/receiver/create_task.html', users=users, depts=depts)

@app.route('/special_report/receiver/task_detail/<int:task_id>')
@login_required
def task_detail(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    submissions = TaskSubmission.query.filter_by(task_id=task.id).all()
    # 获取每个报送的详细报告
    reports = []
    for sub in submissions:
        report = SpecialReport.query.get(sub.report_id)
        if report:
            reports.append({'report': report, 'submission': sub})
    return render_template('special_report/receiver/task_detail.html', task=task, reports=reports)

@app.route('/special_report/receiver/close_task/<int:task_id>')
@login_required
def close_task(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    if task.created_by == current_user.id:
        task.status = 'closed'
        db.session.commit()
        flash('任务已关闭')
    return redirect(url_for('task_manage'))

@app.route('/special_report/receiver/delete_task/<int:task_id>')
@login_required
def delete_task(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    if task.created_by == current_user.id and task.status == 'closed':
        # 检查是否有上报信息
        submissions = TaskSubmission.query.filter_by(task_id=task.id).count()
        if submissions == 0:
            db.session.delete(task)
            db.session.commit()
            flash('删除成功')
        else:
            flash('该任务已有报送信息，无法删除')
    return redirect(url_for('task_manage'))

@app.route('/special_report/receiver/approve_submission/<int:submission_id>')
@login_required
def approve_submission(submission_id):
    """约稿审核 - 采用报送内容"""
    submission = TaskSubmission.query.get_or_404(submission_id)
    task = AssignmentTask.query.get(submission.task_id)
    if not task or task.created_by != current_user.id:
        flash('无权审核该报送')
        return redirect(url_for('task_manage'))
    # 更新提交状态
    submission.status = 'approved'
    # 同步更新关联报告状态为已采用
    if submission.report_id:
        report = SpecialReport.query.get(submission.report_id)
        if report:
            report.status = 'adopted'
    db.session.commit()
    flash('审核通过，已采用该信息')
    return redirect(url_for('task_detail', task_id=submission.task_id))

@app.route('/special_report/receiver/reject_submission/<int:submission_id>', methods=['GET', 'POST'])
@login_required
def reject_submission(submission_id):
    """约稿审核 - 退回报送内容"""
    submission = TaskSubmission.query.get_or_404(submission_id)
    task = AssignmentTask.query.get(submission.task_id)
    if not task or task.created_by != current_user.id:
        flash('无权审核该报送')
        return redirect(url_for('task_manage'))
    if request.method == 'POST':
        feedback = request.form.get('feedback', '')
        submission.status = 'rejected'
        if submission.report_id:
            report = SpecialReport.query.get(submission.report_id)
            if report:
                report.status = 'rejected'
                report.feedback = feedback
        db.session.commit()
        flash('已退回该报送，退回意见已通知报送方')
        return redirect(url_for('task_detail', task_id=submission.task_id))
    return render_template('special_report/receiver/reject_submission.html',
                           submission=submission, task=task)

@app.route('/special_report/receiver/ai_compile/<int:task_id>', methods=['GET', 'POST'])
@login_required
def ai_compile(task_id):
    task = AssignmentTask.query.get_or_404(task_id)
    if request.method == 'POST':
        # 获取选中的报告ID列表
        report_ids = request.form.getlist('report_ids')
        if report_ids:
            # 模拟AI汇编
            compiled_content = "AI汇编结果：\n"
            for rid in report_ids:
                report = SpecialReport.query.get(rid)
                compiled_content += f"- {report.title}\n  {report.content[:100]}...\n"
            # 保存为新的AI报告（可以新建一个报告）
            new_report = SpecialReport(
                title=f"【AI汇编】{task.title}",
                content=compiled_content,
                category=task.category,
                reporter_id=current_user.id,
                reporter_name=current_user.name,
                reporter_phone=current_user.phone,
                target_department='',
                status='pending'
            )
            db.session.add(new_report)
            db.session.commit()
            flash('AI汇编完成，已保存为新报告')
            return redirect(url_for('info_pool'))
    # 获取已报送且审核通过的报告
    submissions = TaskSubmission.query.filter_by(task_id=task.id).all()
    reports = []
    for sub in submissions:
        report = SpecialReport.query.get(sub.report_id)
        if report and report.status == 'adopted':
            reports.append(report)
    return render_template('special_report/receiver/ai_compile.html', task=task, reports=reports)

# 专项报送统计
@app.route('/special_report/receiver/statistics')
@login_required
def statistics():
    users = User.query.all()
    stats = []
    for u in users:
        total = SpecialReport.query.filter_by(reporter_id=u.id).count()
        if total > 0:  # 只显示有报送的用户
            pending = SpecialReport.query.filter_by(reporter_id=u.id, status='pending').count()
            adopted = SpecialReport.query.filter_by(reporter_id=u.id, status='adopted').count()
            rejected = SpecialReport.query.filter_by(reporter_id=u.id, status='rejected').count()
            stats.append({'user': u, 'count': total, 'pending': pending,
                          'adopted': adopted, 'rejected': rejected})
    # 按总数降序
    stats.sort(key=lambda x: x['count'], reverse=True)
    # 总体统计
    total_reports = SpecialReport.query.count()
    total_adopted = SpecialReport.query.filter_by(status='adopted').count()
    total_pending = SpecialReport.query.filter_by(status='pending').count()
    total_rejected = SpecialReport.query.filter_by(status='rejected').count()
    return render_template('special_report/receiver/statistics.html', stats=stats,
                           total_reports=total_reports, total_adopted=total_adopted,
                           total_pending=total_pending, total_rejected=total_rejected)


# ==================== 系统管理：角色权限 ====================

def _log_op(module, action, target='', detail=''):
    """记录操作日志辅助函数"""
    try:
        log = SystemOperationLog(
            user_id=current_user.id,
            username=current_user.username,
            module=module,
            action=action,
            target=target,
            detail=detail,
            ip_addr=request.remote_addr
        )
        db.session.add(log)
        db.session.commit()
    except Exception:
        pass

def _track_usage(module, action):
    """记录功能使用统计"""
    try:
        today = datetime.now().strftime('%Y-%m-%d')
        stat = SystemUsageStat.query.filter_by(stat_date=today, module=module, action=action).first()
        if stat:
            stat.count += 1
        else:
            stat = SystemUsageStat(stat_date=today, module=module, action=action, count=1, user_count=1)
            db.session.add(stat)
        db.session.commit()
    except Exception:
        try:
            db.session.rollback()
        except Exception:
            pass


@app.route('/admin/roles')
@login_required
def admin_roles():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    roles = Role.query.order_by(Role.is_system.desc(), Role.id).all()
    # 统计每个角色的用户数
    user_counts = {}
    for r in roles:
        user_counts[r.name] = User.query.filter_by(role=r.name).count()
    return render_template('admin/roles.html', roles=roles, user_counts=user_counts)

@app.route('/admin/role/add', methods=['GET', 'POST'])
@login_required
def admin_role_add():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        display_name = request.form.get('display_name', '').strip()
        description = request.form.get('description', '').strip()
        permissions = request.form.getlist('permissions')
        if not name or not display_name:
            flash('角色标识和名称不能为空')
            return redirect(url_for('admin_role_add'))
        if Role.query.filter_by(name=name).first():
            flash('角色标识已存在')
            return redirect(url_for('admin_role_add'))
        role = Role(name=name, display_name=display_name, description=description)
        role.set_permissions(permissions)
        db.session.add(role)
        db.session.commit()
        _log_op('角色管理', 'create', name, f'创建角色 {display_name}')
        flash('角色创建成功')
        return redirect(url_for('admin_roles'))
    all_perms = _get_all_permissions()
    return render_template('admin/role_form.html', role=None, all_perms=all_perms)

@app.route('/admin/role/edit/<int:role_id>', methods=['GET', 'POST'])
@login_required
def admin_role_edit(role_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    role = Role.query.get_or_404(role_id)
    if request.method == 'POST':
        if role.is_system and role.name == 'admin':
            flash('系统内置管理员角色不允许修改权限')
            return redirect(url_for('admin_roles'))
        role.display_name = request.form.get('display_name', role.display_name).strip()
        role.description = request.form.get('description', '').strip()
        permissions = request.form.getlist('permissions')
        role.set_permissions(permissions)
        role.updated_at = datetime.now()
        db.session.commit()
        _log_op('角色管理', 'update', role.name, f'更新角色 {role.display_name} 权限')
        flash('角色更新成功')
        return redirect(url_for('admin_roles'))
    all_perms = _get_all_permissions()
    return render_template('admin/role_form.html', role=role, all_perms=all_perms)

@app.route('/admin/role/delete/<int:role_id>')
@login_required
def admin_role_delete(role_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    role = Role.query.get_or_404(role_id)
    if role.is_system:
        flash('系统内置角色不能删除')
        return redirect(url_for('admin_roles'))
    # 检查是否有用户使用此角色
    user_count = User.query.filter_by(role=role.name).count()
    if user_count > 0:
        flash(f'该角色下有 {user_count} 名用户，请先修改这些用户的角色')
        return redirect(url_for('admin_roles'))
    _log_op('角色管理', 'delete', role.name, f'删除角色 {role.display_name}')
    db.session.delete(role)
    db.session.commit()
    flash('角色删除成功')
    return redirect(url_for('admin_roles'))

@app.route('/admin/operation_logs')
@login_required
def admin_operation_logs():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    page = request.args.get('page', 1, type=int)
    module_filter = request.args.get('module', '')
    action_filter = request.args.get('action', '')
    user_filter = request.args.get('user', '')
    q = SystemOperationLog.query
    if module_filter:
        q = q.filter_by(module=module_filter)
    if action_filter:
        q = q.filter_by(action=action_filter)
    if user_filter:
        q = q.filter(SystemOperationLog.username.like(f'%{user_filter}%'))
    logs = q.order_by(SystemOperationLog.created_at.desc()).paginate(page=page, per_page=30, error_out=False)
    modules = db.session.query(SystemOperationLog.module).distinct().all()
    return render_template('admin/operation_logs.html', logs=logs, modules=modules,
                           module_filter=module_filter, action_filter=action_filter, user_filter=user_filter)

def _get_all_permissions():
    """返回所有可分配权限清单"""
    return [
        {'code': 'user_manage',      'name': '用户管理',      'group': '系统管理'},
        {'code': 'role_manage',      'name': '角色权限管理',   'group': '系统管理'},
        {'code': 'template_manage',  'name': '公文模板管理',   'group': '系统管理'},
        {'code': 'operation_log',    'name': '操作日志查看',   'group': '系统管理'},
        {'code': 'ai_config',        'name': 'AI模型配置',     'group': '系统管理'},
        {'code': 'report_view',      'name': '信息池查看',     'group': '专报管理'},
        {'code': 'report_submit',    'name': '信息报送',       'group': '专报管理'},
        {'code': 'report_manage',    'name': '信息审核/管理',  'group': '专报管理'},
        {'code': 'task_view',        'name': '任务查看',       'group': '任务管理'},
        {'code': 'task_manage',      'name': '任务创建/管理',  'group': '任务管理'},
        {'code': 'knowledge_view',   'name': '知识库查看',     'group': '知识管理'},
        {'code': 'knowledge_manage', 'name': '知识库管理',     'group': '知识管理'},
        {'code': 'stats_view',       'name': '统计报表查看',   'group': '统计分析'},
        {'code': 'briefing_manage',  'name': '简报系统',       'group': '简报管理'},
    ]


# ==================== 公文模板动态管理 ====================

@app.route('/admin/doc_templates')
@login_required
def admin_doc_templates():
    if current_user.role not in ('admin', 'manager'):
        flash('权限不足')
        return redirect(url_for('index'))
    category = request.args.get('category', '')
    q = DocTemplate.query.filter_by(is_active=True)
    if category:
        q = q.filter_by(category=category)
    templates = q.order_by(DocTemplate.sort_order.desc(), DocTemplate.updated_at.desc()).all()
    categories = db.session.query(DocTemplate.category).filter(DocTemplate.is_active == True).distinct().all()
    categories = [c[0] for c in categories if c[0]]
    _track_usage('公文模板', 'view')
    return render_template('admin/doc_templates.html', templates=templates,
                           categories=categories, current_category=category)

@app.route('/admin/doc_template/add', methods=['GET', 'POST'])
@login_required
def admin_doc_template_add():
    if current_user.role not in ('admin', 'manager'):
        flash('权限不足')
        return redirect(url_for('index'))
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        category = request.form.get('category', '通用').strip()
        description = request.form.get('description', '').strip()
        content = request.form.get('content', '').strip()
        tags = request.form.get('tags', '').strip()
        sort_order = int(request.form.get('sort_order', 0))
        if not name:
            flash('模板名称不能为空')
            return redirect(url_for('admin_doc_template_add'))
        tpl = DocTemplate(
            name=name, category=category, description=description,
            content=content, tags=tags, sort_order=sort_order,
            created_by=current_user.id, updated_by=current_user.id
        )
        # 处理文件上传
        f = request.files.get('template_file')
        if f and f.filename:
            import uuid
            ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else 'txt'
            filename = f'template_{uuid.uuid4().hex}.{ext}'
            save_dir = os.path.join(app.root_path, 'uploads', 'templates')
            os.makedirs(save_dir, exist_ok=True)
            f.save(os.path.join(save_dir, filename))
            tpl.file_path = f'uploads/templates/{filename}'
            tpl.file_type = ext
        else:
            tpl.file_type = 'txt'
        db.session.add(tpl)
        db.session.commit()
        _log_op('公文模板', 'create', name, f'新建公文模板 {name}')
        flash('模板创建成功')
        return redirect(url_for('admin_doc_templates'))
    return render_template('admin/doc_template_form.html', template=None)

@app.route('/admin/doc_template/edit/<int:tpl_id>', methods=['GET', 'POST'])
@login_required
def admin_doc_template_edit(tpl_id):
    if current_user.role not in ('admin', 'manager'):
        flash('权限不足')
        return redirect(url_for('index'))
    tpl = DocTemplate.query.get_or_404(tpl_id)
    if request.method == 'POST':
        tpl.name = request.form.get('name', tpl.name).strip()
        tpl.category = request.form.get('category', tpl.category).strip()
        tpl.description = request.form.get('description', '').strip()
        tpl.content = request.form.get('content', '').strip()
        tpl.tags = request.form.get('tags', '').strip()
        tpl.sort_order = int(request.form.get('sort_order', 0))
        tpl.updated_by = current_user.id
        tpl.updated_at = datetime.now()
        # 处理文件上传
        f = request.files.get('template_file')
        if f and f.filename:
            import uuid
            ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else 'txt'
            filename = f'template_{uuid.uuid4().hex}.{ext}'
            save_dir = os.path.join(app.root_path, 'uploads', 'templates')
            os.makedirs(save_dir, exist_ok=True)
            f.save(os.path.join(save_dir, filename))
            tpl.file_path = f'uploads/templates/{filename}'
            tpl.file_type = ext
        db.session.commit()
        _log_op('公文模板', 'update', tpl.name, f'更新公文模板 {tpl.name}')
        flash('模板更新成功')
        return redirect(url_for('admin_doc_templates'))
    return render_template('admin/doc_template_form.html', template=tpl)

@app.route('/admin/doc_template/delete/<int:tpl_id>')
@login_required
def admin_doc_template_delete(tpl_id):
    if current_user.role not in ('admin', 'manager'):
        flash('权限不足')
        return redirect(url_for('index'))
    tpl = DocTemplate.query.get_or_404(tpl_id)
    tpl.is_active = False  # 软删除
    db.session.commit()
    _log_op('公文模板', 'delete', tpl.name, f'删除公文模板 {tpl.name}')
    flash('模板已删除')
    return redirect(url_for('admin_doc_templates'))

@app.route('/admin/doc_template/view/<int:tpl_id>')
@login_required
def admin_doc_template_view(tpl_id):
    tpl = DocTemplate.query.get_or_404(tpl_id)
    if not tpl.is_active:
        flash('模板不存在')
        return redirect(url_for('admin_doc_templates'))
    tpl.use_count = (tpl.use_count or 0) + 1
    db.session.commit()
    _track_usage('公文模板', 'use')
    return jsonify({'id': tpl.id, 'name': tpl.name, 'content': tpl.content or '',
                    'category': tpl.category, 'description': tpl.description or ''})

@app.route('/admin/doc_template/download/<int:tpl_id>')
@login_required
def admin_doc_template_download(tpl_id):
    tpl = DocTemplate.query.get_or_404(tpl_id)
    if tpl.file_path and os.path.exists(os.path.join(app.root_path, tpl.file_path)):
        directory = os.path.join(app.root_path, os.path.dirname(tpl.file_path))
        filename = os.path.basename(tpl.file_path)
        tpl.use_count = (tpl.use_count or 0) + 1
        db.session.commit()
        _track_usage('公文模板', 'download')
        return send_from_directory(directory, filename, as_attachment=True,
                                   download_name=f"{tpl.name}.{tpl.file_type}")
    # 如果没有文件，则动态生成 txt
    from flask import Response
    tpl.use_count = (tpl.use_count or 0) + 1
    db.session.commit()
    _track_usage('公文模板', 'download')
    content = tpl.content or f'【{tpl.name}】\n\n{tpl.description or ""}'
    resp = make_response(content.encode('utf-8'))
    resp.headers['Content-Type'] = 'text/plain; charset=utf-8'
    resp.headers['Content-Disposition'] = f'attachment; filename="{tpl.name}.txt"'
    return resp


# ==================== 公文模板（用户端展示） ====================
# 替换原 template_library 路由（保留原路由，同时新增管理员路由）
@app.route('/smart_office/template_library_v2')
@login_required
def template_library_v2():
    """新版公文模板库（从数据库加载）"""
    category = request.args.get('category', '')
    keyword = request.args.get('keyword', '')
    q = DocTemplate.query.filter_by(is_active=True)
    if category:
        q = q.filter_by(category=category)
    if keyword:
        q = q.filter(DocTemplate.name.like(f'%{keyword}%') | DocTemplate.tags.like(f'%{keyword}%'))
    templates = q.order_by(DocTemplate.sort_order.desc(), DocTemplate.use_count.desc()).all()
    categories = db.session.query(DocTemplate.category).filter(DocTemplate.is_active == True).distinct().all()
    categories = [c[0] for c in categories if c[0]]
    _track_usage('公文模板', 'view')
    return render_template('smart_office/template_library_v2.html', templates=templates,
                           categories=categories, current_category=category, keyword=keyword)


# ==================== 报表统计 ====================

@app.route('/admin/stats')
@login_required
def admin_stats():
    """综合报表统计 Dashboard"""
    if current_user.role not in ('admin', 'manager'):
        flash('权限不足')
        return redirect(url_for('index'))
    _track_usage('报表统计', 'view')

    # 用户统计
    total_users = User.query.count()
    role_dist = {}
    for role_name in ('admin', 'manager', 'reporter', 'employee'):
        role_dist[role_name] = User.query.filter_by(role=role_name).count()

    # 专报统计（近30天）
    from datetime import timedelta
    now = datetime.now()
    thirty_days_ago = now - timedelta(days=30)
    report_total = SpecialReport.query.count()
    report_month = SpecialReport.query.filter(SpecialReport.created_at >= thirty_days_ago).count()
    report_status_dist = {
        'draft': SpecialReport.query.filter_by(status='draft').count(),
        'pending': SpecialReport.query.filter_by(status='pending').count(),
        'approved': SpecialReport.query.filter_by(status='approved').count(),
        'adopted': SpecialReport.query.filter_by(status='adopted').count(),
        'rejected': SpecialReport.query.filter_by(status='rejected').count(),
    }

    # 任务统计
    task_total = AssignmentTask.query.count()
    task_active = AssignmentTask.query.filter_by(status='active').count()

    # 知识库统计
    kb_total = KnowledgeBase.query.count()
    file_total = KnowledgeFile.query.count()

    # AI使用统计
    ai_msg_total = ChatMessage.query.count()
    ai_session_total = ChatSession.query.count()

    # 功能使用频率（近7天）
    today_str = now.strftime('%Y-%m-%d')
    week_start = (now - timedelta(days=6)).strftime('%Y-%m-%d')
    _usage_rows = db.session.query(
        SystemUsageStat.module,
        db.func.sum(SystemUsageStat.count).label('total')
    ).filter(
        SystemUsageStat.stat_date >= week_start
    ).group_by(SystemUsageStat.module).order_by(db.func.sum(SystemUsageStat.count).desc()).all()
    # 转换为可JSON序列化的字典列表
    usage_data = [{'module': row.module, 'total': int(row.total)} for row in _usage_rows]

    # 近7天每日报送量
    daily_reports = []
    for i in range(6, -1, -1):
        d = now - timedelta(days=i)
        d_start = d.replace(hour=0, minute=0, second=0, microsecond=0)
        d_end = d.replace(hour=23, minute=59, second=59, microsecond=999999)
        cnt = SpecialReport.query.filter(
            SpecialReport.created_at >= d_start,
            SpecialReport.created_at <= d_end
        ).count()
        daily_reports.append({'date': d.strftime('%m-%d'), 'count': cnt})

    # 近7天AI对话量
    daily_ai = []
    for i in range(6, -1, -1):
        d = now - timedelta(days=i)
        d_start = d.replace(hour=0, minute=0, second=0, microsecond=0)
        d_end = d.replace(hour=23, minute=59, second=59, microsecond=999999)
        cnt = ChatMessage.query.filter(
            ChatMessage.created_at >= d_start,
            ChatMessage.created_at <= d_end,
            ChatMessage.role == 'user'
        ).count()
        daily_ai.append({'date': d.strftime('%m-%d'), 'count': cnt})

    # 模板使用排行
    top_templates = DocTemplate.query.filter_by(is_active=True).order_by(
        DocTemplate.use_count.desc()).limit(10).all()

    # 活跃用户（报送量Top10）
    from sqlalchemy import func
    active_reporters = db.session.query(
        User, func.count(SpecialReport.id).label('cnt')
    ).join(SpecialReport, SpecialReport.reporter_id == User.id
    ).filter(SpecialReport.created_at >= thirty_days_ago
    ).group_by(User.id).order_by(func.count(SpecialReport.id).desc()).limit(10).all()

    return render_template('admin/stats.html',
        total_users=total_users, role_dist=role_dist,
        report_total=report_total, report_month=report_month,
        report_status_dist=report_status_dist,
        task_total=task_total, task_active=task_active,
        kb_total=kb_total, file_total=file_total,
        ai_msg_total=ai_msg_total, ai_session_total=ai_session_total,
        usage_data=usage_data, daily_reports=daily_reports, daily_ai=daily_ai,
        top_templates=top_templates, active_reporters=active_reporters
    )

@app.route('/admin/stats/api/usage_trend')
@login_required
def admin_stats_usage_trend():
    """使用趋势 API（返回 JSON）"""
    if current_user.role not in ('admin', 'manager'):
        return jsonify({'error': '权限不足'}), 403
    days = request.args.get('days', 30, type=int)
    from datetime import timedelta
    now = datetime.now()
    start_date = (now - timedelta(days=days-1)).strftime('%Y-%m-%d')
    data = db.session.query(
        SystemUsageStat.stat_date,
        SystemUsageStat.module,
        db.func.sum(SystemUsageStat.count).label('total')
    ).filter(SystemUsageStat.stat_date >= start_date
    ).group_by(SystemUsageStat.stat_date, SystemUsageStat.module
    ).order_by(SystemUsageStat.stat_date).all()
    result = {}
    for row in data:
        d = row.stat_date
        if d not in result:
            result[d] = {}
        result[d][row.module] = int(row.total)
    return jsonify(result)

@app.route('/admin/stats/api/report_trend')
@login_required  
def admin_stats_report_trend():
    """专报趋势 API"""
    if current_user.role not in ('admin', 'manager'):
        return jsonify({'error': '权限不足'}), 403
    from datetime import timedelta
    days = request.args.get('days', 30, type=int)
    now = datetime.now()
    result = []
    for i in range(days-1, -1, -1):
        d = now - timedelta(days=i)
        d_start = d.replace(hour=0, minute=0, second=0, microsecond=0)
        d_end = d.replace(hour=23, minute=59, second=59, microsecond=999999)
        total = SpecialReport.query.filter(
            SpecialReport.created_at >= d_start,
            SpecialReport.created_at <= d_end
        ).count()
        adopted = SpecialReport.query.filter(
            SpecialReport.created_at >= d_start,
            SpecialReport.created_at <= d_end,
            SpecialReport.status == 'adopted'
        ).count()
        result.append({'date': d.strftime('%Y-%m-%d'), 'total': total, 'adopted': adopted})
    return jsonify(result)


# ==================== 组织机构管理 ====================

def _require_admin():
    """检查是否管理员，非管理员返回重定向"""
    if current_user.role != 'admin':
        flash('权限不足，需要管理员权限', 'warning')
        return redirect(url_for('index'))
    return None

# ---------- 组织机构 ----------
@app.route('/admin/org')
@login_required
def admin_org_list():
    err = _require_admin()
    if err: return err
    orgs = Organization.query.order_by(Organization.level, Organization.sort_order).all()
    return render_template('admin/org_list.html', orgs=orgs)

@app.route('/admin/org/add', methods=['GET', 'POST'])
@login_required
def admin_org_add():
    err = _require_admin()
    if err: return err
    orgs = Organization.query.filter_by(is_active=True).order_by(Organization.level, Organization.sort_order).all()
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash('机构名称不能为空', 'danger')
        else:
            parent_id = request.form.get('parent_id') or None
            org = Organization(
                name=name,
                short_name=request.form.get('short_name', '').strip(),
                code=request.form.get('code', '').strip() or None,
                org_type=request.form.get('org_type', 'unit'),
                parent_id=int(parent_id) if parent_id else None,
                level=(Organization.query.get(int(parent_id)).level + 1) if parent_id else 1,
                sort_order=int(request.form.get('sort_order', 0)),
                address=request.form.get('address', '').strip(),
                phone=request.form.get('phone', '').strip(),
                description=request.form.get('description', '').strip(),
            )
            db.session.add(org)
            db.session.commit()
            _log_op('org', 'create', f'新建机构: {name}')
            flash('机构创建成功', 'success')
            return redirect(url_for('admin_org_list'))
    return render_template('admin/org_form.html', org=None, orgs=orgs, action='add')

@app.route('/admin/org/edit/<int:org_id>', methods=['GET', 'POST'])
@login_required
def admin_org_edit(org_id):
    err = _require_admin()
    if err: return err
    org = Organization.query.get_or_404(org_id)
    orgs = Organization.query.filter(Organization.id != org_id, Organization.is_active==True).order_by(Organization.level, Organization.sort_order).all()
    if request.method == 'POST':
        org.name = request.form.get('name', '').strip()
        org.short_name = request.form.get('short_name', '').strip()
        org.code = request.form.get('code', '').strip() or None
        org.org_type = request.form.get('org_type', 'unit')
        parent_id = request.form.get('parent_id') or None
        org.parent_id = int(parent_id) if parent_id else None
        org.sort_order = int(request.form.get('sort_order', 0))
        org.address = request.form.get('address', '').strip()
        org.phone = request.form.get('phone', '').strip()
        org.description = request.form.get('description', '').strip()
        org.is_active = request.form.get('is_active') == '1'
        db.session.commit()
        _log_op('org', 'update', f'编辑机构: {org.name}')
        flash('机构信息已更新', 'success')
        return redirect(url_for('admin_org_list'))
    return render_template('admin/org_form.html', org=org, orgs=orgs, action='edit')

@app.route('/admin/org/delete/<int:org_id>', methods=['POST'])
@login_required
def admin_org_delete(org_id):
    err = _require_admin()
    if err: return err
    org = Organization.query.get_or_404(org_id)
    if org.departments.count() > 0:
        flash('该机构下存在部门，无法删除', 'danger')
        return redirect(url_for('admin_org_list'))
    if org.children.count() > 0:
        flash('该机构存在子机构，请先删除子机构', 'danger')
        return redirect(url_for('admin_org_list'))
    org.is_active = False
    db.session.commit()
    _log_op('org', 'delete', f'停用机构: {org.name}')
    flash('机构已停用', 'success')
    return redirect(url_for('admin_org_list'))

# ---------- 部门管理 ----------
@app.route('/admin/dept')
@login_required
def admin_dept_list():
    err = _require_admin()
    if err: return err
    org_id = request.args.get('org_id', type=int)
    query = Department.query.filter_by(is_active=True)
    if org_id:
        query = query.filter_by(org_id=org_id)
    depts = query.order_by(Department.sort_order).all()
    orgs = Organization.query.filter_by(is_active=True).all()
    return render_template('admin/dept_list.html', depts=depts, orgs=orgs, current_org_id=org_id)

@app.route('/admin/dept/add', methods=['GET', 'POST'])
@login_required
def admin_dept_add():
    err = _require_admin()
    if err: return err
    orgs = Organization.query.filter_by(is_active=True).order_by(Organization.sort_order).all()
    depts_all = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    users_all = User.query.filter_by(is_active=True).order_by(User.name).all()
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash('部门名称不能为空', 'danger')
        else:
            parent_id = request.form.get('parent_id') or None
            manager_id = request.form.get('manager_id') or None
            dept = Department(
                name=name,
                code=request.form.get('code', '').strip() or None,
                org_id=int(request.form.get('org_id')) if request.form.get('org_id') else None,
                parent_id=int(parent_id) if parent_id else None,
                dept_type=request.form.get('dept_type', 'functional'),
                manager_id=int(manager_id) if manager_id else None,
                sort_order=int(request.form.get('sort_order', 0)),
                description=request.form.get('description', '').strip(),
            )
            db.session.add(dept)
            db.session.commit()
            _log_op('dept', 'create', f'新建部门: {name}')
            flash('部门创建成功', 'success')
            return redirect(url_for('admin_dept_list'))
    return render_template('admin/dept_form.html', dept=None, orgs=orgs,
                           depts_all=depts_all, users_all=users_all, action='add')

@app.route('/admin/dept/edit/<int:dept_id>', methods=['GET', 'POST'])
@login_required
def admin_dept_edit(dept_id):
    err = _require_admin()
    if err: return err
    dept = Department.query.get_or_404(dept_id)
    orgs = Organization.query.filter_by(is_active=True).order_by(Organization.sort_order).all()
    depts_all = Department.query.filter(Department.id != dept_id, Department.is_active==True).order_by(Department.sort_order).all()
    users_all = User.query.filter_by(is_active=True).order_by(User.name).all()
    if request.method == 'POST':
        dept.name = request.form.get('name', '').strip()
        dept.code = request.form.get('code', '').strip() or None
        dept.org_id = int(request.form.get('org_id')) if request.form.get('org_id') else None
        parent_id = request.form.get('parent_id') or None
        dept.parent_id = int(parent_id) if parent_id else None
        dept.dept_type = request.form.get('dept_type', 'functional')
        manager_id = request.form.get('manager_id') or None
        dept.manager_id = int(manager_id) if manager_id else None
        dept.sort_order = int(request.form.get('sort_order', 0))
        dept.description = request.form.get('description', '').strip()
        dept.is_active = request.form.get('is_active') == '1'
        db.session.commit()
        _log_op('dept', 'update', f'编辑部门: {dept.name}')
        flash('部门信息已更新', 'success')
        return redirect(url_for('admin_dept_list'))
    return render_template('admin/dept_form.html', dept=dept, orgs=orgs,
                           depts_all=depts_all, users_all=users_all, action='edit')

@app.route('/admin/dept/delete/<int:dept_id>', methods=['POST'])
@login_required
def admin_dept_delete(dept_id):
    err = _require_admin()
    if err: return err
    dept = Department.query.get_or_404(dept_id)
    member_count = User.query.filter_by(dept_id=dept_id, is_active=True).count()
    if member_count > 0:
        flash(f'该部门下还有 {member_count} 名在职人员，请先调整人员归属', 'danger')
        return redirect(url_for('admin_dept_list'))
    dept.is_active = False
    db.session.commit()
    _log_op('dept', 'delete', f'停用部门: {dept.name}')
    flash('部门已停用', 'success')
    return redirect(url_for('admin_dept_list'))

@app.route('/admin/dept/<int:dept_id>/members')
@login_required
def admin_dept_members(dept_id):
    err = _require_admin()
    if err: return err
    dept = Department.query.get_or_404(dept_id)
    members = User.query.filter_by(dept_id=dept_id).order_by(User.name).all()
    positions = Position.query.filter_by(dept_id=dept_id, is_active=True).all()
    return render_template('admin/dept_members.html', dept=dept, members=members, positions=positions)

# ---------- 岗位管理 ----------
@app.route('/admin/position')
@login_required
def admin_position_list():
    err = _require_admin()
    if err: return err
    dept_id = request.args.get('dept_id', type=int)
    query = Position.query.filter_by(is_active=True)
    if dept_id:
        query = query.filter_by(dept_id=dept_id)
    positions = query.order_by(Position.sort_order).all()
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    roles = Role.query.order_by(Role.display_name).all()
    return render_template('admin/position_list.html', positions=positions, depts=depts,
                           roles=roles, current_dept_id=dept_id)

@app.route('/admin/position/add', methods=['GET', 'POST'])
@login_required
def admin_position_add():
    err = _require_admin()
    if err: return err
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    roles = Role.query.order_by(Role.display_name).all()
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash('岗位名称不能为空', 'danger')
        else:
            dept_id = request.form.get('dept_id') or None
            pos = Position(
                name=name,
                code=request.form.get('code', '').strip() or None,
                dept_id=int(dept_id) if dept_id else None,
                role_name=request.form.get('role_name', '').strip() or None,
                level=request.form.get('level', 'staff'),
                headcount=int(request.form.get('headcount', 1)),
                description=request.form.get('description', '').strip(),
                requirements=request.form.get('requirements', '').strip(),
                sort_order=int(request.form.get('sort_order', 0)),
            )
            db.session.add(pos)
            db.session.commit()
            _log_op('position', 'create', f'新建岗位: {name}')
            flash('岗位创建成功', 'success')
            return redirect(url_for('admin_position_list'))
    return render_template('admin/position_form.html', pos=None, depts=depts, roles=roles, action='add')

@app.route('/admin/position/edit/<int:pos_id>', methods=['GET', 'POST'])
@login_required
def admin_position_edit(pos_id):
    err = _require_admin()
    if err: return err
    pos = Position.query.get_or_404(pos_id)
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    roles = Role.query.order_by(Role.display_name).all()
    if request.method == 'POST':
        pos.name = request.form.get('name', '').strip()
        pos.code = request.form.get('code', '').strip() or None
        dept_id = request.form.get('dept_id') or None
        pos.dept_id = int(dept_id) if dept_id else None
        pos.role_name = request.form.get('role_name', '').strip() or None
        pos.level = request.form.get('level', 'staff')
        pos.headcount = int(request.form.get('headcount', 1))
        pos.description = request.form.get('description', '').strip()
        pos.requirements = request.form.get('requirements', '').strip()
        pos.sort_order = int(request.form.get('sort_order', 0))
        pos.is_active = request.form.get('is_active') == '1'
        db.session.commit()
        _log_op('position', 'update', f'编辑岗位: {pos.name}')
        flash('岗位信息已更新', 'success')
        return redirect(url_for('admin_position_list'))
    return render_template('admin/position_form.html', pos=pos, depts=depts, roles=roles, action='edit')

@app.route('/admin/position/delete/<int:pos_id>', methods=['POST'])
@login_required
def admin_position_delete(pos_id):
    err = _require_admin()
    if err: return err
    pos = Position.query.get_or_404(pos_id)
    count = User.query.filter_by(position_id=pos_id, is_active=True).count()
    if count > 0:
        flash(f'该岗位上还有 {count} 名在职人员，请先调岗', 'danger')
        return redirect(url_for('admin_position_list'))
    pos.is_active = False
    db.session.commit()
    _log_op('position', 'delete', f'停用岗位: {pos.name}')
    flash('岗位已停用', 'success')
    return redirect(url_for('admin_position_list'))

# ---------- 用户管理（升级版，支持组织架构） ----------
@app.route('/admin/users')
@login_required
def admin_users():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    keyword = request.args.get('kw', '').strip()
    dept_id = request.args.get('dept_id', type=int)
    role_filter = request.args.get('role', '').strip()
    status_filter = request.args.get('status', 'active')

    query = User.query
    if keyword:
        query = query.filter(
            db.or_(User.name.contains(keyword), User.username.contains(keyword),
                   User.employee_no.contains(keyword), User.phone.contains(keyword))
        )
    if dept_id:
        query = query.filter_by(dept_id=dept_id)
    if role_filter:
        query = query.filter_by(role=role_filter)
    if status_filter == 'active':
        query = query.filter(db.or_(User.is_active == True, User.is_active == None))
    elif status_filter == 'inactive':
        query = query.filter_by(is_active=False)

    users = query.order_by(User.dept_id, User.name).all()
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    roles = Role.query.order_by(Role.display_name).all()
    return render_template('admin/users.html', users=users, depts=depts, roles=roles,
                           keyword=keyword, dept_id=dept_id, role_filter=role_filter,
                           status_filter=status_filter)

@app.route('/admin/user/add', methods=['GET', 'POST'])
@login_required
def admin_user_add():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    positions = Position.query.filter_by(is_active=True).order_by(Position.sort_order).all()
    orgs = Organization.query.filter_by(is_active=True).order_by(Organization.sort_order).all()
    roles = Role.query.order_by(Role.display_name).all()
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        name = request.form.get('name', '').strip()
        password = request.form.get('password', '').strip()
        if not username or not name or not password:
            flash('用户名、姓名、密码不能为空', 'danger')
        elif User.query.filter_by(username=username).first():
            flash('用户名已存在', 'danger')
        else:
            dept_id = request.form.get('dept_id') or None
            position_id = request.form.get('position_id') or None
            org_id = request.form.get('org_id') or None
            # 联动：从岗位自动获取角色
            role = request.form.get('role', 'employee')
            if position_id:
                pos = Position.query.get(int(position_id))
                if pos and pos.role_name:
                    role = pos.role_name
            # 联动：从部门自动获取部门名称（兼容旧字段）
            dept_name = ''
            if dept_id:
                d = Department.query.get(int(dept_id))
                if d:
                    dept_name = d.name

            user = User(
                username=username,
                name=name,
                password=generate_password_hash(password),
                phone=request.form.get('phone', '').strip(),
                email=request.form.get('email', '').strip(),
                role=role,
                department=dept_name,
                dept_id=int(dept_id) if dept_id else None,
                position_id=int(position_id) if position_id else None,
                org_id=int(org_id) if org_id else None,
                employee_no=request.form.get('employee_no', '').strip() or None,
                gender=request.form.get('gender', '未知'),
                is_reporter=request.form.get('is_reporter') == '1',
                is_receiver=request.form.get('is_receiver') == '1',
                is_active=True,
                remark=request.form.get('remark', '').strip(),
            )
            db.session.add(user)
            db.session.commit()
            # 自动创建个人知识库
            personal_kb = KnowledgeBase(name=f"{user.name}的个人知识库", type='personal', owner_id=user.id)
            db.session.add(personal_kb)
            db.session.commit()
            _log_op('user', 'create', f'新建用户: {username}')
            flash('用户创建成功', 'success')
            return redirect(url_for('admin_users'))
    return render_template('admin/user_form.html', user=None, depts=depts,
                           positions=positions, orgs=orgs, roles=roles, action='add')

@app.route('/admin/user/edit/<int:user_id>', methods=['GET', 'POST'])
@login_required
def admin_user_edit(user_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    user = User.query.get_or_404(user_id)
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    positions = Position.query.filter_by(is_active=True).order_by(Position.sort_order).all()
    orgs = Organization.query.filter_by(is_active=True).order_by(Organization.sort_order).all()
    roles = Role.query.order_by(Role.display_name).all()
    if request.method == 'POST':
        user.name = request.form.get('name', '').strip()
        user.phone = request.form.get('phone', '').strip()
        user.email = request.form.get('email', '').strip()
        dept_id = request.form.get('dept_id') or None
        position_id = request.form.get('position_id') or None
        org_id = request.form.get('org_id') or None
        user.dept_id = int(dept_id) if dept_id else None
        user.position_id = int(position_id) if position_id else None
        user.org_id = int(org_id) if org_id else None
        user.employee_no = request.form.get('employee_no', '').strip() or None
        user.gender = request.form.get('gender', '未知')
        # 联动部门名称（兼容旧字段）
        if dept_id:
            d = Department.query.get(int(dept_id))
            if d:
                user.department = d.name
        # 联动岗位角色
        role = request.form.get('role', user.role)
        if position_id:
            pos = Position.query.get(int(position_id))
            if pos and pos.role_name:
                role = pos.role_name
        user.role = role
        user.is_reporter = request.form.get('is_reporter') == '1'
        user.is_receiver = request.form.get('is_receiver') == '1'
        user.is_active = request.form.get('is_active') == '1'
        user.remark = request.form.get('remark', '').strip()
        new_pwd = request.form.get('password', '').strip()
        if new_pwd:
            user.password = generate_password_hash(new_pwd)
        db.session.commit()
        _log_op('user', 'update', f'编辑用户: {user.username}')
        flash('用户信息已更新', 'success')
        return redirect(url_for('admin_users'))
    return render_template('admin/user_form.html', user=user, depts=depts,
                           positions=positions, orgs=orgs, roles=roles, action='edit')

@app.route('/admin/user/delete/<int:user_id>', methods=['POST'])
@login_required
def admin_user_delete(user_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    user = User.query.get_or_404(user_id)
    if user.username == 'admin':
        flash('不能删除超级管理员账号', 'danger')
        return redirect(url_for('admin_users'))
    user.is_active = False
    db.session.commit()
    _log_op('user', 'delete', f'停用用户: {user.username}')
    flash(f'用户 {user.name} 已停用', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/reset_pwd/<int:user_id>', methods=['POST'])
@login_required
def admin_user_reset_pwd(user_id):
    """重置用户密码"""
    if current_user.role != 'admin':
        return jsonify({'success': False, 'msg': '权限不足'}), 403
    user = User.query.get_or_404(user_id)
    new_pwd = request.form.get('new_pwd', '').strip()
    if not new_pwd or len(new_pwd) < 6:
        return jsonify({'success': False, 'msg': '密码不能少于6位'})
    user.password = generate_password_hash(new_pwd)
    db.session.commit()
    _log_op('user', 'reset_pwd', f'重置密码: {user.username}')
    return jsonify({'success': True, 'msg': f'用户 {user.name} 密码已重置'})

# ---------- 组织架构总览（树形结构） ----------
@app.route('/admin/org_tree')
@login_required
def admin_org_tree():
    err = _require_admin()
    if err: return err
    orgs = Organization.query.filter_by(is_active=True).order_by(Organization.level, Organization.sort_order).all()
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    positions = Position.query.filter_by(is_active=True).order_by(Position.sort_order).all()
    # 统计
    total_users = User.query.filter(db.or_(User.is_active==True, User.is_active==None)).count()
    total_depts = len(depts)
    total_positions = len(positions)
    return render_template('admin/org_tree.html', orgs=orgs, depts=depts, positions=positions,
                           total_users=total_users, total_depts=total_depts, total_positions=total_positions)

# ---------- API：获取部门下的岗位（动态联动） ----------
@app.route('/admin/api/positions_by_dept')
@login_required
@csrf.exempt
def api_positions_by_dept():
    dept_id = request.args.get('dept_id', type=int)
    if not dept_id:
        return jsonify([])
    positions = Position.query.filter_by(dept_id=dept_id, is_active=True).order_by(Position.sort_order).all()
    return jsonify([{'id': p.id, 'name': p.name, 'role_name': p.role_name or ''} for p in positions])

# ---------- API：获取部门下的用户（公文收发联动） ----------
@app.route('/admin/api/users_by_dept')
@login_required
@csrf.exempt
def api_users_by_dept():
    dept_id = request.args.get('dept_id', type=int)
    if not dept_id:
        users = User.query.filter(db.or_(User.is_active==True, User.is_active==None)).order_by(User.name).all()
    else:
        users = User.query.filter_by(dept_id=dept_id).filter(
            db.or_(User.is_active==True, User.is_active==None)
        ).order_by(User.name).all()
    return jsonify([{'id': u.id, 'name': u.name, 'dept': u.department or ''} for u in users])


# ==================== 错误处理 ====================
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


@app.errorhandler(500)
def internal_server_error(e):
    try:
        db.session.rollback()
    except Exception:
        pass
    return render_template('500.html'), 500

# ==================== AI 模型配置管理 ====================

@app.route('/ai/config')
@login_required
def ai_config_list():
    if current_user.role != 'admin':
        flash('只有管理员可以管理AI配置')
        return redirect(url_for('index'))
    configs = AIModelConfig.query.all()
    return render_template('ai_config.html', configs=configs)

@app.route('/ai/config/add', methods=['GET', 'POST'])
@login_required
def ai_config_add():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    if request.method == 'POST':
        name = request.form.get('name')
        provider = request.form.get('provider')
        api_key = request.form.get('api_key')
        api_base = request.form.get('api_base')
        model_name = request.form.get('model_name')
        temperature = float(request.form.get('temperature', 0.7))
        max_tokens = int(request.form.get('max_tokens', 2000))
        context_length = int(request.form.get('context_length', 4000))
        delay = int(request.form.get('delay', 0))
        is_active = request.form.get('is_active') == 'on'
        config = AIModelConfig(
            name=name, provider=provider, api_key=api_key, api_base=api_base,
            model_name=model_name, temperature=temperature, max_tokens=max_tokens,
            context_length=context_length, delay=delay, is_active=is_active
        )
        db.session.add(config)
        db.session.commit()
        flash('添加成功')
        return redirect(url_for('ai_config_list'))
    return render_template('ai_config_form.html')

@app.route('/ai/config/edit/<int:config_id>', methods=['GET', 'POST'])
@login_required
def ai_config_edit(config_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    config = AIModelConfig.query.get_or_404(config_id)
    if request.method == 'POST':
        config.name = request.form.get('name')
        config.provider = request.form.get('provider')
        config.api_key = request.form.get('api_key')
        config.api_base = request.form.get('api_base')
        config.model_name = request.form.get('model_name')
        config.temperature = float(request.form.get('temperature', 0.7))
        config.max_tokens = int(request.form.get('max_tokens', 2000))
        config.context_length = int(request.form.get('context_length', 4000))
        config.delay = int(request.form.get('delay', 0))
        config.is_active = request.form.get('is_active') == 'on'
        db.session.commit()
        flash('更新成功')
        return redirect(url_for('ai_config_list'))
    return render_template('ai_config_form.html', config=config)

@app.route('/ai/config/delete/<int:config_id>')
@login_required
def ai_config_delete(config_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))
    config = AIModelConfig.query.get_or_404(config_id)
    db.session.delete(config)
    db.session.commit()
    flash('删除成功')
    return redirect(url_for('ai_config_list'))

# ==================== 增强的AI对话 ====================
@app.route('/qa/ai_chat', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def ai_chat():
    # 获取所有已激活的模型配置
    models = AIModelConfig.query.filter_by(is_active=True).all()
    # 获取当前用户的对话会话列表
    sessions = ChatSession.query.filter_by(user_id=current_user.id).order_by(ChatSession.updated_at.desc()).all()

    # 确定当前会话：
    # ?new=1  → 强制空白（不加载任何会话）
    # ?session_id=X → 加载指定会话
    # 无参数  → 不自动加载（等待用户发送第一条消息后再创建）
    force_new = request.args.get('new') == '1'
    session_id = None if force_new else request.args.get('session_id')
    # POST 时也可通过 form 传 session_id（前端发送消息时继续已有会话）
    if request.method == 'POST':
        session_id = request.form.get('session_id') or session_id

    current_session = None
    messages = []
    if session_id and not force_new:
        current_session = ChatSession.query.get(session_id)
        if current_session and current_session.user_id == current_user.id:
            messages = ChatMessage.query.filter_by(session_id=session_id).order_by(ChatMessage.created_at).all()
        else:
            session_id = None
            current_session = None

    # 获取知识库列表（用于前端选择）
    personal_kb = KnowledgeBase.query.filter_by(owner_id=current_user.id, type='personal').first()
    shared_kbs = KnowledgeBase.query.filter_by(type='shared', is_public=True).all()
    policy_kbs = KnowledgeBase.query.filter_by(type='policy').all()

    if request.method == 'POST':
        # 处理文件上传
        uploaded_file = request.files.get('file')
        file_content = ''
        if uploaded_file:
            filepath = save_upload_file(uploaded_file, 'chat')
            if filepath:
                if filepath.endswith('.txt'):
                    with open(filepath, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                elif filepath.endswith('.pdf'):
                    file_content = pdf_to_text(filepath)
                elif filepath.endswith('.docx'):
                    import docx
                    doc = docx.Document(filepath)
                    file_content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    file_content = '不支持的文件类型'
        question = request.form.get('question')
        model_id = request.form.get('model_id')
        # 支持多值 FormData（前端用 formData.append('knowledge_bases', val) 逐个追加）
        selected_kbs = request.form.getlist('knowledge_bases')

        # 如果没有当前会话，创建新会话
        if not current_session:
            model_config = AIModelConfig.query.get(model_id) if model_id else models[0] if models else None
            if not model_config:
                return jsonify({'error': '没有可用的模型配置'}), 400
            current_session = ChatSession(
                user_id=current_user.id,
                title=question[:50] if question else '新对话',
                model_config_id=model_config.id
            )
            db.session.add(current_session)
            db.session.commit()
            session_id = current_session.id
        else:
            if model_id and current_session.model_config_id != int(model_id):
                current_session.model_config_id = int(model_id)
                db.session.commit()

        # ===== P0: RAG 智能检索（替代全文件读取）=====
        knowledge_context, citations = get_knowledge_context_for_ai(
            selected_kbs, question or '', max_snippets=5, max_chars=3000
        )
        # 追加上传的文件内容
        if file_content:
            max_length = config_manager.get("document.max_ai_sample_length", 5000)
            if max_length == -1:
                max_length = len(file_content)  # -1表示无限制
            if max_length > 0 and len(file_content) > max_length:
                file_content = file_content[:max_length] + "...（内容过长，已截断）"
            knowledge_context += f"\n\n【上传文件内容】\n{file_content}"

        # 构建消息上下文
        model_config = AIModelConfig.query.get(current_session.model_config_id)
        if not model_config:
            return jsonify({'error': '模型配置无效'}), 400

        # 获取历史消息
        history = ChatMessage.query.filter_by(session_id=current_session.id).order_by(ChatMessage.created_at).all()
        messages = [{'role': msg.role, 'content': msg.content} for msg in history]

        # 构建系统消息（包含知识库上下文）
        system_content = "你是一个智能助手。"
        if knowledge_context:
            system_content += f"\n以下是知识库中相关内容，请参考回答用户问题：\n{knowledge_context}"
        # 将系统消息放在开头
        messages.insert(0, {'role': 'system', 'content': system_content})

        # 添加用户消息（如果有文件内容，先附加）
        user_content = question
        if file_content:
            user_content = f"用户上传了文件内容：\n{file_content}\n\n问题：{question}"
        messages.append({'role': 'user', 'content': user_content})

        # 限制上下文长度
        total_len = sum(len(m['content']) for m in messages)
        if total_len > model_config.context_length:
            # 保留系统消息，删除最早的用户/助手消息
            while total_len > model_config.context_length and len(messages) > 1:
                removed = messages.pop(1)  # 保留系统消息
                total_len -= len(removed['content'])

        # 调用AI
        answer = call_ai_model(model_config, messages)

        # 保存用户消息和AI回复
        user_msg = ChatMessage(session_id=current_session.id, role='user', content=user_content)
        db.session.add(user_msg)
        ai_msg = ChatMessage(session_id=current_session.id, role='assistant', content=answer)
        db.session.add(ai_msg)
        db.session.commit()
        current_session.updated_at = datetime.now()
        db.session.commit()

        return jsonify({'answer': answer, 'session_id': current_session.id})

    return render_template('intelligent_qa/ai_chat.html',
                           models=models,
                           sessions=sessions,
                           current_session=current_session,
                           messages=messages,
                           personal_kb=personal_kb,
                           shared_kbs=shared_kbs,
                           policy_kbs=policy_kbs)

@app.route('/chat/session/delete/<int:session_id>', methods=['POST'])
@login_required
@csrf.exempt
def delete_session(session_id):
    session = ChatSession.query.get_or_404(session_id)
    if session.user_id == current_user.id:
        # 删除所有消息
        ChatMessage.query.filter_by(session_id=session.id).delete()
        db.session.delete(session)
        db.session.commit()
    return jsonify({'status': 'ok'})

# 静态文件缓存优化
@app.after_request
def add_cache_headers(response):
    """为静态文件添加缓存头"""
    if response.status_code == 200:
        # 检查是否为静态文件
        if request.path.startswith('/static/'):
            # direct_passthrough 模式下不能修改响应体，跳过
            if response.direct_passthrough:
                return response
            # 缓存1年（31536000秒）
            response.cache_control.public = True
            response.cache_control.max_age = 31536000
            response.cache_control.s_maxage = 31536000
    return response


# ==================== 简报生成系统 ====================

# 简报任务队列
import queue
import threading
import traceback
from urllib.parse import urljoin
from bs4 import BeautifulSoup

briefing_task_queue = queue.Queue()
briefing_task_status = {}

# 创建简报输出目录
BRIEFING_OUTPUT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'briefing_output')
os.makedirs(BRIEFING_OUTPUT_FOLDER, exist_ok=True)


# ==================== 简报页面路由 ====================

@app.route('/briefing')
@login_required
def briefing_index():
    """简报系统首页"""
    from utils import BriefingStatisticsCalculator
    
    today = datetime.now().strftime('%Y-%m-%d')
    stats = BriefingStatisticsCalculator.calculate_daily_stats(today)
    recent_briefings = Briefing.query.order_by(Briefing.start_time.desc()).limit(5).all()
    keywords = BriefingKeyword.query.filter_by(is_active=True).order_by(BriefingKeyword.use_count.desc()).limit(10).all()
    sources = BriefingSource.query.filter_by(is_active=True).all()
    
    return render_template('briefing/index.html',
                          stats=stats,
                          recent_briefings=recent_briefings,
                          keywords=keywords,
                          sources=sources,
                          now=datetime.now())


@app.route('/briefing/sources')
@login_required
def briefing_sources():
    """数据源管理页面"""
    sources = BriefingSource.query.order_by(BriefingSource.priority.desc(), BriefingSource.name).all()
    return render_template('briefing/sources.html', sources=sources)


@app.route('/briefing/keywords')
@login_required
def briefing_keywords():
    """关键词管理页面"""
    keywords = BriefingKeyword.query.order_by(BriefingKeyword.use_count.desc()).all()
    return render_template('briefing/keywords.html', keywords=keywords)


@app.route('/briefing/history')
@login_required
def briefing_history():
    """历史记录页面"""
    page = request.args.get('page', 1, type=int)
    briefings = Briefing.query.order_by(Briefing.start_time.desc())\
        .paginate(page=page, per_page=20)
    return render_template('briefing/history.html', briefings=briefings)


@app.route('/briefing/settings')
@login_required
def briefing_settings():
    """简报系统设置页面"""
    tasks = BriefingScheduledTask.query.all()
    return render_template('briefing/settings.html', tasks=tasks)


# ==================== 简报API路由 ====================

@app.route('/briefing/api/generate', methods=['POST'])
@login_required
@csrf.exempt
def briefing_generate():
    """生成简报"""
    try:
        data = request.json
        keywords = data.get('keywords', [])
        source_ids = data.get('sources', [])
        date = data.get('date', datetime.now().strftime('%Y%m%d'))
        title = data.get('title', f'学习资料汇编_{date}')
        
        if not keywords:
            return jsonify({'success': False, 'error': '请至少选择一个关键词'}), 400
        if not source_ids:
            return jsonify({'success': False, 'error': '请至少选择一个数据源'}), 400
        
        task_id = datetime.now().strftime('%Y%m%d%H%M%S')
        
        briefing = Briefing(
            task_id=task_id,
            title=title,
            keywords=json.dumps(keywords, ensure_ascii=False),
            sources=json.dumps(source_ids, ensure_ascii=False),
            target_date=date,
            status='pending',
            user_id=current_user.id
        )
        db.session.add(briefing)
        db.session.commit()
        
        briefing_task_queue.put({
            'task_id': task_id,
            'keywords': keywords,
            'source_ids': source_ids,
            'date': date,
            'title': title
        })
        
        return jsonify({'success': True, 'task_id': task_id, 'message': '任务已加入队列'})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/briefing/api/task_status/<task_id>')
@login_required
def get_briefing_task_status(task_id):
    """获取任务状态"""
    briefing = Briefing.query.filter_by(task_id=task_id).first()
    if not briefing:
        return jsonify({'error': 'Task not found'}), 404
    
    response = {
        'task_id': briefing.task_id,
        'status': briefing.status,
        'progress': 0,
        'message': '',
        'article_count': briefing.article_count,
        'start_time': briefing.start_time.strftime('%H:%M:%S') if briefing.start_time else None,
        'end_time': briefing.end_time.strftime('%H:%M:%S') if briefing.end_time else None,
        'error': briefing.error_message
    }
    
    if briefing.status == 'running' and task_id in briefing_task_status:
        response['progress'] = briefing_task_status[task_id].get('progress', 0)
        response['message'] = briefing_task_status[task_id].get('message', '')
    
    return jsonify(response)


@app.route('/briefing/api/briefing/<task_id>')
@login_required
def get_briefing_detail(task_id):
    """获取简报详情"""
    briefing = Briefing.query.filter_by(task_id=task_id).first()
    if not briefing:
        return jsonify({'error': 'Briefing not found'}), 404
    
    articles = BriefingArticle.query.filter_by(briefing_id=briefing.id).all()
    return jsonify({
        'briefing': {
            'id': briefing.id, 'task_id': briefing.task_id, 'title': briefing.title,
            'status': briefing.status, 'article_count': briefing.article_count,
            'keywords': briefing.get_keywords(), 'sources': briefing.get_sources(),
            'created_at': briefing.start_time.strftime('%Y-%m-%d %H:%M:%S') if briefing.start_time else None
        },
        'articles': [{'id': a.id, 'title': a.title, 'source': a.source_name, 'url': a.source_url, 'keyword': a.keyword} for a in articles]
    })


@app.route('/briefing/api/download/<task_id>/<format>')
@login_required
def briefing_download(task_id, format):
    """下载简报文档"""
    from flask import send_file
    briefing = Briefing.query.filter_by(task_id=task_id).first()
    if not briefing:
        return jsonify({'error': 'Briefing not found'}), 404
    
    if format == 'word' and briefing.docx_path and os.path.exists(briefing.docx_path):
        return send_file(briefing.docx_path, as_attachment=True, download_name=f"{briefing.title}.docx")
    
    return jsonify({'error': 'File not found'}), 404


# ==================== 数据源API ====================

@app.route('/briefing/api/sources', methods=['GET'])
@login_required
def get_briefing_sources():
    sources = BriefingSource.query.filter_by(is_active=True).all()
    return jsonify({'sources': [{'id': s.id, 'name': s.name, 'url': s.url, 'type': s.source_type} for s in sources]})


@app.route('/briefing/api/sources', methods=['POST'])
@login_required
@csrf.exempt
def add_briefing_source():
    try:
        data = request.json
        source = BriefingSource(
            name=data.get('name'),
            url=data.get('url'),
            source_type=data.get('type', 'website'),
            category=data.get('category')
        )
        db.session.add(source)
        db.session.commit()
        return jsonify({'success': True, 'id': source.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/briefing/api/sources/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
@csrf.exempt
def handle_briefing_source(id):
    source = BriefingSource.query.get_or_404(id)
    
    if request.method == 'GET':
        return jsonify({
            'id': source.id, 'name': source.name, 'url': source.url,
            'type': source.source_type, 'category': source.category
        })
    
    elif request.method == 'PUT':
        data = request.json
        source.name = data.get('name', source.name)
        source.url = data.get('url', source.url)
        source.source_type = data.get('type', source.source_type)
        source.category = data.get('category', source.category)
        db.session.commit()
        return jsonify({'success': True, 'message': '更新成功'})
    
    elif request.method == 'DELETE':
        db.session.delete(source)
        db.session.commit()
        return jsonify({'success': True, 'message': '删除成功'})


@app.route('/briefing/api/sources/toggle/<int:id>', methods=['POST'])
@login_required
@csrf.exempt
def toggle_briefing_source(id):
    source = BriefingSource.query.get_or_404(id)
    source.is_active = not source.is_active
    db.session.commit()
    return jsonify({'success': True, 'is_active': source.is_active})


# ==================== 关键词API ====================

@app.route('/briefing/api/keywords', methods=['GET'])
@login_required
def get_briefing_keywords():
    keywords = BriefingKeyword.query.filter_by(is_active=True).all()
    return jsonify({'keywords': [{'id': k.id, 'text': k.text, 'category': k.category, 'color': k.color} for k in keywords]})


@app.route('/briefing/api/keywords', methods=['POST'])
@login_required
@csrf.exempt
def add_briefing_keyword():
    try:
        data = request.json
        existing = BriefingKeyword.query.filter_by(text=data.get('text')).first()
        if existing:
            return jsonify({'success': False, 'error': '关键词已存在'}), 400
        
        keyword = BriefingKeyword(
            text=data.get('text'),
            category=data.get('category'),
            description=data.get('description'),
            color=data.get('color', '#3498db')
        )
        db.session.add(keyword)
        db.session.commit()
        return jsonify({'success': True, 'id': keyword.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/briefing/api/keywords/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
@csrf.exempt
def handle_briefing_keyword(id):
    keyword = BriefingKeyword.query.get_or_404(id)
    
    if request.method == 'GET':
        return jsonify({'id': keyword.id, 'text': keyword.text, 'category': keyword.category, 'description': keyword.description, 'color': keyword.color})
    
    elif request.method == 'PUT':
        data = request.json
        existing = BriefingKeyword.query.filter(BriefingKeyword.text == data.get('text'), BriefingKeyword.id != id).first()
        if existing:
            return jsonify({'success': False, 'error': '关键词文本已存在'}), 400
            
        keyword.text = data.get('text', keyword.text)
        keyword.category = data.get('category', keyword.category)
        keyword.description = data.get('description', keyword.description)
        keyword.color = data.get('color', keyword.color)
        db.session.commit()
        return jsonify({'success': True, 'message': '更新成功'})
    
    elif request.method == 'DELETE':
        db.session.delete(keyword)
        db.session.commit()
        return jsonify({'success': True, 'message': '删除成功'})


# ==================== 定时任务API ====================

@app.route('/briefing/api/schedule', methods=['POST'])
@login_required
@csrf.exempt
def create_briefing_schedule():
    """创建简报定时任务"""
    try:
        data = request.json
        task = BriefingScheduledTask(
            name=data.get('name'),
            cron_expression=data.get('cron'),
            keywords=json.dumps(data.get('keywords', []), ensure_ascii=False),
            sources=json.dumps(data.get('sources', []), ensure_ascii=False),
            email_recipients=json.dumps(data.get('emails', []), ensure_ascii=False)
        )
        db.session.add(task)
        db.session.commit()
        return jsonify({'success': True, 'id': task.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/briefing/api/task/<int:id>', methods=['DELETE'])
@login_required
@csrf.exempt
def delete_briefing_schedule(id):
    """删除简报定时任务"""
    task = BriefingScheduledTask.query.get_or_404(id)
    db.session.delete(task)
    db.session.commit()
    return jsonify({'success': True, 'message': '定时任务已删除'})


@app.route('/briefing/api/statistics/trend')
@login_required
def get_briefing_trend():
    """获取简报统计趋势"""
    from utils import BriefingStatisticsCalculator
    days = request.args.get('days', 30, type=int)
    trend = BriefingStatisticsCalculator.get_trend_data(days)
    return jsonify(trend)


@app.route('/briefing/api/logs')
@login_required
def get_briefing_logs():
    """获取简报系统日志"""
    page = request.args.get('page', 1, type=int)
    logs = BriefingSystemLog.query.order_by(BriefingSystemLog.created_at.desc())\
        .paginate(page=page, per_page=50)
    return jsonify({
        'logs': [{
            'id': l.id,
            'level': l.level,
            'module': l.module,
            'message': l.message,
            'created_at': l.created_at.strftime('%Y-%m-%d %H:%M:%S') if l.created_at else None
        } for l in logs.items],
        'total': logs.total,
        'pages': logs.pages,
        'current_page': logs.page
    })


@app.route('/briefing/api/all_briefings')
@login_required
def get_all_briefings():
    """获取所有简报列表"""
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    status = request.args.get('status', None)
    
    query = Briefing.query
    if status:
        query = query.filter_by(status=status)
    
    briefings = query.order_by(Briefing.start_time.desc()).paginate(page=page, per_page=per_page)
    
    return jsonify({
        'briefings': [{
            'id': b.id,
            'task_id': b.task_id,
            'title': b.title,
            'status': b.status,
            'article_count': b.article_count,
            'start_time': b.start_time.strftime('%Y-%m-%d %H:%M:%S') if b.start_time else None,
            'end_time': b.end_time.strftime('%Y-%m-%d %H:%M:%S') if b.end_time else None,
            'duration': b.duration
        } for b in briefings.items],
        'total': briefings.total,
        'pages': briefings.pages,
        'current_page': briefings.page
    })


@app.route('/briefing/api/briefing/<task_id>', methods=['DELETE'])
@login_required
@csrf.exempt
def delete_briefing(task_id):
    """删除简报"""
    briefing = Briefing.query.filter_by(task_id=task_id).first()
    if not briefing:
        return jsonify({'error': 'Briefing not found'}), 404
    
    db.session.delete(briefing)
    db.session.commit()
    return jsonify({'success': True, 'message': '简报已删除'})


# ==================== 后台任务处理器 ====================

def briefing_task_worker():
    """后台任务工作线程"""
    from scraper_engine import engine
    from utils import BriefingDocumentGenerator
    
    while True:
        try:
            task_data = briefing_task_queue.get()
            if task_data is None:
                break
            
            with app.app_context():
                try:
                    task_id = task_data['task_id']
                    
                    briefing = Briefing.query.filter_by(task_id=task_id).first()
                    if not briefing:
                        continue
                    
                    briefing.status = 'running'
                    db.session.commit()
                    briefing_task_status[task_id] = {'progress': 0, 'message': '正在初始化...'}
                    
                    source_ids = json.loads(briefing.sources)
                    keywords = [str(k) for k in json.loads(briefing.keywords)]
                    
                    # 获取关键词文本
                    kw_objects = BriefingKeyword.query.filter(BriefingKeyword.id.in_([int(k) for k in keywords if k.isdigit()])).all()
                    kw_texts = [kw.text for kw in kw_objects] if kw_objects else keywords
                    
                    sources = BriefingSource.query.filter(BriefingSource.id.in_([int(s) for s in source_ids if s.isdigit()])).all()
                    
                    all_articles = []
                    processed_urls = set()
                    link_pool = []
                    
                    # 阶段1: 扫描首页
                    briefing_task_status[task_id] = {'progress': 10, 'message': f'正在扫描 {len(sources)} 个数据源...'}
                    
                    for source in sources:
                        try:
                            result = engine.fetch_url(source.url, timeout=10)
                            if result['status'] == 'success':
                                soup = BeautifulSoup(result['html'], 'lxml')
                                links = soup.find_all('a', href=True)
                                
                                for link in links:
                                    href = link.get('href', '')
                                    title = link.get_text(strip=True)
                                    
                                    if not href or href.startswith('javascript:') or len(title) < 4:
                                        continue
                                    
                                    full_url = urljoin(source.url, href)
                                    
                                    # 关键词初筛
                                    is_match = any(kw in title for kw in kw_texts)
                                    
                                    if is_match and full_url not in processed_urls:
                                        processed_urls.add(full_url)
                                        link_pool.append({
                                            'url': full_url,
                                            'title': title,
                                            'source_name': source.name
                                        })
                        except Exception as e:
                            BriefingSystemLog.log('WARNING', 'crawler', f'扫描数据源失败 {source.name}: {str(e)}')

                    total_links = len(link_pool)
                    if total_links == 0:
                        briefing_task_status[task_id] = {'progress': 100, 'message': '未发现相关文章'}
                        briefing.status = 'completed'
                        briefing.article_count = 0
                        db.session.commit()
                        continue

                    # 阶段2: 并发抓取正文
                    briefing_task_status[task_id] = {'progress': 20, 'message': f'发现 {total_links} 篇疑似文章，开始抓取...'}
                    
                    count = 0
                    for item in link_pool[:50]:  # 限制最多抓取50篇
                        try:
                            result = engine.fetch_url(item['url'], timeout=10)
                            count += 1
                            
                            progress = 20 + int((count / min(total_links, 50)) * 70)
                            briefing_task_status[task_id] = {'progress': progress, 'message': f'已抓取 {count}/{min(total_links, 50)}...'}
                            
                            if result['status'] == 'success':
                                parsed = engine.parse_article(result['html'])
                                if parsed and parsed['content']:
                                    full_text = parsed['content'] + parsed['title']
                                    matched_kw = next((kw for kw in kw_texts if kw in full_text), None)
                                    
                                    if matched_kw:
                                        all_articles.append({
                                            'title': item['title'] or parsed['title'],
                                            'content': parsed['content'],
                                            'url': item['url'],
                                            'source_name': item['source_name'],
                                            'keyword': matched_kw
                                        })
                        except Exception:
                            pass

                    # 阶段3: 保存结果
                    briefing_task_status[task_id] = {'progress': 95, 'message': '正在生成文档...'}
                    
                    if all_articles:
                        unique_articles = []
                        seen_titles = set()
                        for art in all_articles:
                            if art['title'] not in seen_titles:
                                seen_titles.add(art['title'])
                                unique_articles.append(art)
                        
                        output_filename = f"{briefing.title}.docx"
                        output_path = os.path.join(BRIEFING_OUTPUT_FOLDER, output_filename)
                        BriefingDocumentGenerator.create_word_document(unique_articles, output_path)
                        briefing.docx_path = output_path
                        
                        for article in unique_articles:
                            art = BriefingArticle(
                                briefing_id=briefing.id,
                                title=article['title'],
                                content=article['content'],
                                source_name=article['source_name'],
                                source_url=article['url'],
                                keyword=article['keyword'],
                                word_count=len(article['content'])
                            )
                            db.session.add(art)
                        
                        briefing.article_count = len(unique_articles)
                    
                    briefing.status = 'completed'
                    briefing.end_time = datetime.now()
                    briefing.duration = (briefing.end_time - briefing.start_time).seconds if briefing.start_time else 0
                    db.session.commit()
                    
                    briefing_task_status[task_id] = {'progress': 100, 'message': f'任务完成！共抓取 {len(all_articles)} 篇文章。'}
                
                except Exception as e:
                    BriefingSystemLog.log('ERROR', 'task', f'任务执行失败: {str(e)}')
                    traceback.print_exc()
                    if 'task_id' in dir():
                        briefing = Briefing.query.filter_by(task_id=task_id).first()
                        if briefing:
                            briefing.status = 'failed'
                            briefing.error_message = str(e)
                            db.session.commit()

        except Exception as outer_e:
            print(f"严重错误: {outer_e}")
            traceback.print_exc()


# 启动后台工作线程
briefing_worker_thread = threading.Thread(target=briefing_task_worker, daemon=True)
briefing_worker_thread.start()


def _parse_datetime_local(value):
    """解析 datetime-local / 普通日期时间字符串"""
    if not value:
        return None
    for fmt in ('%Y-%m-%dT%H:%M', '%Y-%m-%d %H:%M', '%Y-%m-%d %H:%M:%S'):
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            continue
    return None


def _parse_date_local(value):
    """解析日期字符串"""
    if not value:
        return None
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except ValueError:
        return None


def _parse_float_local(value, default=None):
    """解析浮点数字符串，非法时返回 None"""
    if value is None:
        return default
    value = str(value).strip()
    if not value:
        return default
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _user_dept_name(user):

    if not user:
        return ''
    return user.department or (user.dept.name if getattr(user, 'dept', None) else '') or ''


def _sync_meeting_attendance(meeting):
    attendee_ids = [int(uid) for uid in meeting.get_attendee_ids() if uid]
    existing_records = {
        record.user_id: record for record in MeetingAttendance.query.filter_by(meeting_id=meeting.id).all()
    }

    for user_id in attendee_ids:
        if user_id not in existing_records:
            db.session.add(MeetingAttendance(meeting_id=meeting.id, user_id=user_id, attendance_status='pending'))

    for user_id, record in existing_records.items():
        if user_id not in attendee_ids:
            db.session.delete(record)


def _meeting_attendance_summary(meeting):
    attendee_ids = [int(uid) for uid in meeting.get_attendee_ids() if uid]
    records = {
        record.user_id: record for record in MeetingAttendance.query.filter_by(meeting_id=meeting.id).all()
    }
    summary = {
        'total': len(attendee_ids),
        'signed': 0,
        'leave': 0,
        'absent': 0,
        'pending': 0,
        'feedback_count': 0,
    }
    for user_id in attendee_ids:
        record = records.get(user_id)
        status = record.attendance_status if record else 'pending'
        if status not in ('signed', 'leave', 'absent'):
            status = 'pending'
        summary[status] += 1
        if record and record.feedback:
            summary['feedback_count'] += 1
    return summary


def _find_meeting_room_conflict(room_id, start_time, end_time, exclude_meeting_id=None):
    if not room_id or not start_time or not end_time:
        return None
    query = Meeting.query.filter(
        Meeting.room_id == room_id,
        Meeting.status.in_(['pending', 'confirmed', 'in_progress']),
        Meeting.start_time.isnot(None),
        Meeting.end_time.isnot(None),
        Meeting.start_time < end_time,
        Meeting.end_time > start_time,
    )
    if exclude_meeting_id:
        query = query.filter(Meeting.id != exclude_meeting_id)
    return query.order_by(Meeting.start_time.asc()).first()



def _format_deadline_delta(target_time, now=None):

    if not target_time:
        return '未设置时限'
    now = now or datetime.now()
    delta_seconds = int((target_time - now).total_seconds())
    overdue = delta_seconds < 0
    delta_seconds = abs(delta_seconds)
    days, remainder = divmod(delta_seconds, 86400)
    hours, _ = divmod(remainder, 3600)
    if days:
        text = f'{days}天{hours}小时'
    else:
        text = f'{max(hours, 0)}小时'
    return f'已逾期{text}' if overdue else f'剩余{text}'


def _get_supervision_alert_meta(task, now=None):
    now = now or datetime.now()
    is_active = task.status not in ('completed', 'cancelled')
    due_date = task.due_date
    is_overdue = bool(due_date and due_date < now and is_active)
    hours_left = None
    if due_date:
        hours_left = (due_date - now).total_seconds() / 3600
    is_due_soon = bool(due_date and is_active and not is_overdue and hours_left is not None and hours_left <= 24)
    latest_reminder = task.progress_logs.filter(
        SupervisionProgress.action.in_(['manual_remind', 'auto_remind', 'overdue_warn'])
    ).order_by(SupervisionProgress.created_at.desc()).first()
    return {
        'is_overdue': is_overdue,
        'is_due_soon': is_due_soon,
        'deadline_text': _format_deadline_delta(due_date, now) if due_date else '未设置时限',
        'badge_class': 'danger' if is_overdue else 'warning' if is_due_soon else 'secondary',
        'row_class': 'table-danger' if is_overdue else 'table-warning' if is_due_soon else '',
        'latest_reminder_text': latest_reminder.created_at.strftime('%m-%d %H:%M') if latest_reminder else '',
    }


def _refresh_supervision_alerts(tasks=None):
    now = datetime.now()
    if tasks is None:
        tasks = SupervisionTask.query.filter(SupervisionTask.status.in_(['issued', 'processing', 'review'])).all()
    else:
        tasks = list(tasks)

    changed = False
    for task in tasks:
        if not task.due_date or task.status in ('completed', 'cancelled'):
            continue

        alert_action = None
        alert_note = ''
        hours_left = (task.due_date - now).total_seconds() / 3600
        if task.due_date < now:
            alert_action = 'overdue_warn'
            alert_note = f'系统自动预警：任务已逾期，请负责人和协办人员立即处理。'
        elif hours_left <= 24:
            alert_action = 'auto_remind'
            alert_note = f'系统自动催办：任务将在24小时内到期，请及时反馈办理进展。'

        if not alert_action:
            continue

        recent_log = task.progress_logs.filter_by(action=alert_action).order_by(SupervisionProgress.created_at.desc()).first()
        if recent_log and (now - recent_log.created_at).total_seconds() < 24 * 3600:
            continue

        db.session.add(SupervisionProgress(
            task_id=task.id,
            operator_id=task.creator_id,
            action=alert_action,
            progress_percent=task.progress_percent or 0,
            note=alert_note,
        ))
        changed = True

    if changed:
        db.session.commit()


def _build_supervision_leader_board(tasks, now=None):
    now = now or datetime.now()
    summary = {}
    risk_tasks = []
    for task in tasks:
        dept_name = _user_dept_name(task.owner) or '未分配部门'
        bucket = summary.setdefault(dept_name, {
            'department': dept_name,
            'total': 0,
            'completed': 0,
            'processing': 0,
            'overdue': 0,
            'due_soon': 0,
            'progress_sum': 0,
        })
        bucket['total'] += 1
        bucket['progress_sum'] += task.progress_percent or 0
        if task.status == 'completed':
            bucket['completed'] += 1
        elif task.status in ('issued', 'processing', 'review'):
            bucket['processing'] += 1

        due_date = task.due_date
        if due_date and task.status not in ('completed', 'cancelled'):
            if due_date < now:
                bucket['overdue'] += 1
                risk_tasks.append(task)
            elif (due_date - now).total_seconds() / 3600 <= 24:
                bucket['due_soon'] += 1
                risk_tasks.append(task)

    rows = []
    for item in summary.values():
        item['avg_progress'] = round(item['progress_sum'] / item['total'], 1) if item['total'] else 0
        rows.append(item)
    rows.sort(key=lambda item: (-item['overdue'], -item['due_soon'], -item['total'], item['department']))
    risk_tasks = sorted({task.id: task for task in risk_tasks}.values(), key=lambda item: item.due_date or datetime.max)
    return rows[:8], risk_tasks[:6]


def _build_performance_dashboard(records):
    dept_summary = {}
    period_summary = {}
    user_summary = {}
    user_period_summary = {}
    period_meta = {}
    alert_rows = []

    for record in records:
        full_score = record.full_score or 100
        score = record.score or 0
        score_rate = round((score / full_score) * 100, 1) if full_score else 0
        dept_name = _user_dept_name(record.employee) or '未分配部门'
        period_name = record.period.name if record.period else '未设置周期'
        user_name = record.employee.name if record.employee else f'用户#{record.user_id or "-"}'
        period_meta[period_name] = (
            record.period.start_date if record.period and record.period.start_date else datetime.max.date(),
            record.period.created_at if record.period and record.period.created_at else datetime.max,
        )

        dept_bucket = dept_summary.setdefault(dept_name, {
            'department': dept_name,
            'count': 0,
            'employees': set(),
            'score_sum': 0,
            'rate_sum': 0,
            'published': 0,
        })
        dept_bucket['count'] += 1
        dept_bucket['employees'].add(record.user_id)
        dept_bucket['score_sum'] += score
        dept_bucket['rate_sum'] += score_rate
        if record.status == 'published':
            dept_bucket['published'] += 1

        period_bucket = period_summary.setdefault(period_name, {
            'period': period_name,
            'count': 0,
            'score_sum': 0,
            'rate_sum': 0,
            'published': 0,
        })
        period_bucket['count'] += 1
        period_bucket['score_sum'] += score
        period_bucket['rate_sum'] += score_rate
        if record.status == 'published':
            period_bucket['published'] += 1

        user_bucket = user_summary.setdefault(user_name, {
            'employee': user_name,
            'count': 0,
            'score_sum': 0,
            'rate_sum': 0,
            'published': 0,
        })
        user_bucket['count'] += 1
        user_bucket['score_sum'] += score
        user_bucket['rate_sum'] += score_rate
        if record.status == 'published':
            user_bucket['published'] += 1

        user_period_bucket = user_period_summary.setdefault(user_name, {}).setdefault(period_name, {
            'count': 0,
            'rate_sum': 0,
        })
        user_period_bucket['count'] += 1
        user_period_bucket['rate_sum'] += score_rate

        if record.status == 'published':
            if score_rate < 75:
                alert_rows.append({
                    'level': 'danger',
                    'label': '低分预警',
                    'employee': user_name,
                    'project_name': record.project_name,
                    'period': period_name,
                    'score_rate': score_rate,
                    'score_text': f'{score} / {full_score}',
                    'evaluation': record.evaluation or '建议尽快复盘扣分项并明确改进措施。',
                })
            elif score_rate >= 95:
                alert_rows.append({
                    'level': 'success',
                    'label': '高分亮点',
                    'employee': user_name,
                    'project_name': record.project_name,
                    'period': period_name,
                    'score_rate': score_rate,
                    'score_text': f'{score} / {full_score}',
                    'evaluation': record.evaluation or '当前记录表现突出，可作为经验亮点沉淀。',
                })

    dept_rows = []
    for item in dept_summary.values():
        dept_rows.append({
            'department': item['department'],
            'count': item['count'],
            'employees': len(item['employees']),
            'avg_score': round(item['score_sum'] / item['count'], 1) if item['count'] else 0,
            'avg_rate': round(item['rate_sum'] / item['count'], 1) if item['count'] else 0,
            'published': item['published'],
        })
    dept_rows.sort(key=lambda item: (-item['avg_rate'], -item['count'], item['department']))

    period_rows = []
    for item in period_summary.values():
        period_rows.append({
            'period': item['period'],
            'count': item['count'],
            'avg_score': round(item['score_sum'] / item['count'], 1) if item['count'] else 0,
            'avg_rate': round(item['rate_sum'] / item['count'], 1) if item['count'] else 0,
            'published': item['published'],
        })
    period_rows.sort(key=lambda item: (period_meta.get(item['period'], (datetime.max.date(), datetime.max))[0], period_meta.get(item['period'], (datetime.max.date(), datetime.max))[1], item['period']))

    user_rows = []
    for item in user_summary.values():
        user_rows.append({
            'employee': item['employee'],
            'count': item['count'],
            'avg_score': round(item['score_sum'] / item['count'], 1) if item['count'] else 0,
            'avg_rate': round(item['rate_sum'] / item['count'], 1) if item['count'] else 0,
            'published': item['published'],
        })
    user_rows.sort(key=lambda item: (-item['avg_rate'], -item['count'], item['employee']))

    trend_labels = [item['period'] for item in period_rows[:8]]
    trend_datasets = []
    for index, user in enumerate(user_rows[:5]):
        user_periods = user_period_summary.get(user['employee'], {})
        trend_datasets.append({
            'label': user['employee'],
            'data': [
                round(user_periods.get(period_name, {}).get('rate_sum', 0) / user_periods.get(period_name, {}).get('count', 1), 1)
                if user_periods.get(period_name, {}).get('count') else None
                for period_name in trend_labels
            ],
            'borderColor': [
                'rgba(13, 110, 253, 0.9)',
                'rgba(25, 135, 84, 0.9)',
                'rgba(255, 193, 7, 0.9)',
                'rgba(220, 53, 69, 0.9)',
                'rgba(111, 66, 193, 0.9)',
            ][index % 5],
            'backgroundColor': [
                'rgba(13, 110, 253, 0.15)',
                'rgba(25, 135, 84, 0.15)',
                'rgba(255, 193, 7, 0.18)',
                'rgba(220, 53, 69, 0.15)',
                'rgba(111, 66, 193, 0.15)',
            ][index % 5],
        })

    alert_rows.sort(key=lambda item: (0 if item['level'] == 'danger' else 1, item['score_rate'], item['employee'], item['project_name']))

    return {
        'dept_summary': dept_rows[:8],
        'period_summary': period_rows[:8],
        'user_summary': user_rows[:8],
        'user_trend': {
            'labels': trend_labels,
            'datasets': trend_datasets,
        },
        'score_alerts': alert_rows[:8],
    }



def _add_worklog_review_log(log_id, operator_id, action, from_status, to_status, comment=''):
    db.session.add(WorkLogReview(
        log_id=log_id,
        operator_id=operator_id,
        action=action,
        from_status=from_status,
        to_status=to_status,
        comment=(comment or '').strip(),
    ))


def _gen_supervision_no():
    today = datetime.now().strftime('%Y%m%d')
    count = SupervisionTask.query.filter(SupervisionTask.task_no.like(f'DB-{today}-%')).count()
    return f'DB-{today}-{count + 1:03d}'



# ============================================================
#  公文工作台 + 会议管理 + 督查督办 + 绩效考核 + 工作日志
# ============================================================
@app.route('/official_doc/dashboard')
@login_required
def doc_dashboard():
    """公文工作台，联动会议/督办模块"""
    _track_usage('official_doc', 'view_dashboard')
    _refresh_supervision_alerts()

    my_doc_query = OfficialDoc.query.filter_by(sender_id=current_user.id)
    stats = {
        'draft': my_doc_query.filter_by(status='draft').count(),
        'pending': my_doc_query.filter_by(status='pending_approve').count(),
        'approved': my_doc_query.filter_by(status='approved').count(),
        'sent': my_doc_query.filter_by(status='sent').count(),
        'archived': my_doc_query.filter_by(status='archived').count(),
        'unread': DocReadRecord.query.filter_by(user_id=current_user.id, handle_status='unread').count(),
        'pending_approval': OfficialDoc.query.filter_by(status='pending_approve').count() if current_user.role in ('admin', 'manager') else 0,
    }

    recent_outbox = my_doc_query.order_by(OfficialDoc.updated_at.desc()).limit(6).all()
    recent_inbox_records = DocReadRecord.query.filter_by(user_id=current_user.id).order_by(DocReadRecord.created_at.desc()).limit(6).all()
    inbox_docs = {d.id: d for d in OfficialDoc.query.filter(OfficialDoc.id.in_([r.doc_id for r in recent_inbox_records])).all()} if recent_inbox_records else {}

    meeting_query = Meeting.query.filter(Meeting.status.in_(['pending', 'confirmed', 'in_progress']))
    if current_user.role not in ('admin', 'manager'):
        meeting_query = meeting_query.filter(db.or_(Meeting.creator_id == current_user.id, Meeting.host_id == current_user.id, Meeting.status != 'draft'))
    upcoming_meetings = meeting_query.order_by(Meeting.start_time.asc()).limit(5).all()
    meeting_attendance_map = {meeting.id: _meeting_attendance_summary(meeting) for meeting in upcoming_meetings}

    supervision_candidates = SupervisionTask.query.filter(SupervisionTask.status.in_(['issued', 'processing', 'review'])).order_by(SupervisionTask.due_date.asc()).all()
    if current_user.role in ('admin', 'manager'):
        upcoming_supervisions = supervision_candidates[:5]
    else:
        visible = []
        for task in supervision_candidates:
            if task.creator_id == current_user.id or task.owner_id == current_user.id or current_user.id in task.get_helper_ids():
                visible.append(task)
        upcoming_supervisions = visible[:5]
    supervision_alerts = {task.id: _get_supervision_alert_meta(task) for task in upcoming_supervisions}

    returned_logs = WorkLog.query.filter_by(user_id=current_user.id, status='returned').order_by(WorkLog.updated_at.desc()).limit(5).all()
    pending_review_logs = []
    if current_user.role in ('admin', 'manager'):
        pending_review_logs = WorkLog.query.filter_by(status='submitted').order_by(WorkLog.log_date.desc(), WorkLog.updated_at.desc()).limit(5).all()

    reminder_items = []
    if stats['unread']:
        reminder_items.append({
            'category': '公文',
            'level': 'danger',
            'title': f'待阅公文 {stats["unread"]} 份',
            'meta': '收件箱存在未读公文，建议优先处理。',
            'link': url_for('doc_inbox'),
            'action_text': '去收文',
            'sort_time': datetime.now(),
        })
    if stats['pending']:
        reminder_items.append({
            'category': '发文',
            'level': 'warning',
            'title': f'待审批发文 {stats["pending"]} 份',
            'meta': '你发起的公文还在审批流程中，可关注审批进度。',
            'link': url_for('doc_outbox', status='pending_approve'),
            'action_text': '看发文',
            'sort_time': datetime.now(),
        })
    if stats['pending_approval']:
        reminder_items.append({
            'category': '审批',
            'level': 'warning',
            'title': f'待审批公文 {stats["pending_approval"]} 份',
            'meta': '当前有待审批公文，建议尽快处理审批队列。',
            'link': url_for('doc_pending'),

            'action_text': '去审批',
            'sort_time': datetime.now(),
        })

    now = datetime.now()
    for meeting in upcoming_meetings:
        summary = meeting_attendance_map.get(meeting.id, {})
        starts_soon = bool(meeting.start_time and (meeting.start_time - now).total_seconds() <= 24 * 3600)
        need_follow_up = summary.get('pending', 0) > 0 and meeting.require_signin
        if not starts_soon and not need_follow_up:
            continue
        reminder_items.append({
            'category': '会议',
            'level': 'warning' if need_follow_up else 'info',
            'title': meeting.subject,
            'meta': f'{meeting.start_time.strftime("%m-%d %H:%M") if meeting.start_time else "待定"} · 待反馈 {summary.get("pending", 0)} 人',
            'link': url_for('meeting_detail', meeting_id=meeting.id),
            'action_text': '看会议',
            'sort_time': meeting.start_time or now,
        })

    for task in upcoming_supervisions:
        alert = supervision_alerts.get(task.id) or _get_supervision_alert_meta(task)
        if not alert['is_overdue'] and not alert['is_due_soon']:
            continue
        reminder_items.append({
            'category': '督办',
            'level': 'danger' if alert['is_overdue'] else 'warning',
            'title': task.title,
            'meta': f'{alert["deadline_text"]} · 负责人：{task.owner.name if task.owner else "-"}',
            'link': url_for('supervision_detail', task_id=task.id),
            'action_text': '去督办',
            'sort_time': task.due_date or now,
        })

    for log in returned_logs:
        reminder_items.append({
            'category': '日志',
            'level': 'danger',
            'title': f'日志"{log.title}"被退回',
            'meta': f'{log.log_date.strftime("%Y-%m-%d") if log.log_date else "-"} · 请尽快修改后重新提交。',
            'link': url_for('worklog_new', id=log.id),
            'action_text': '去修改',
            'sort_time': log.updated_at or now,
        })

    for log in pending_review_logs:
        reminder_items.append({
            'category': '日志审批',
            'level': 'info',
            'title': f'待阅日志：{log.title}',
            'meta': f'{log.user.name if log.user else "-"} · {log.log_date.strftime("%Y-%m-%d") if log.log_date else "-"}',
            'link': url_for('worklog_detail', log_id=log.id),
            'action_text': '去查看',
            'sort_time': log.updated_at or now,
        })

    level_order = {'danger': 0, 'warning': 1, 'info': 2, 'success': 3, 'secondary': 4}
    reminder_items.sort(key=lambda item: (level_order.get(item['level'], 9), item['sort_time']))
    reminder_stats = {
        'total': len(reminder_items),
        'urgent': len([item for item in reminder_items if item['level'] == 'danger']),
        'meeting': len([item for item in reminder_items if item['category'] == '会议']),
        'supervision': len([item for item in reminder_items if item['category'] == '督办']),
        'worklog': len([item for item in reminder_items if item['category'] in ('日志', '日志审批')]),
    }

    return render_template(
        'official_doc/dashboard.html',
        stats=stats,
        recent_outbox=recent_outbox,
        recent_inbox_records=recent_inbox_records,
        inbox_docs=inbox_docs,
        upcoming_meetings=upcoming_meetings,
        meeting_attendance_map=meeting_attendance_map,
        upcoming_supervisions=upcoming_supervisions,
        supervision_alerts=supervision_alerts,
        reminder_items=reminder_items[:10],
        reminder_stats=reminder_stats,
        returned_logs=returned_logs,
        pending_review_logs=pending_review_logs,
    )




@app.route('/meeting')
@login_required
def meeting_index():
    """会议管理首页"""
    _track_usage('meeting', 'index')
    keyword = request.args.get('keyword', '').strip()
    status_filter = request.args.get('status', '').strip()
    scope = request.args.get('scope', 'all').strip()
    quick_filter = request.args.get('quick', '').strip()

    q = Meeting.query
    if current_user.role not in ('admin', 'manager'):
        q = q.filter(db.or_(Meeting.creator_id == current_user.id, Meeting.host_id == current_user.id, Meeting.status != 'draft'))
    if status_filter:
        q = q.filter(Meeting.status == status_filter)
    if keyword:
        q = q.filter(db.or_(Meeting.subject.like(f'%{keyword}%'), Meeting.agenda.like(f'%{keyword}%')))

    meetings = q.order_by(Meeting.start_time.asc(), Meeting.created_at.desc()).all()
    if scope == 'mine':
        meetings = [m for m in meetings if m.creator_id == current_user.id or m.host_id == current_user.id]
    elif scope == 'participated':
        meetings = [m for m in meetings if current_user.id in m.get_attendee_ids()]

    now = datetime.now()
    if quick_filter == 'today':
        meetings = [m for m in meetings if m.start_time and m.start_time.date() == now.date()]
    elif quick_filter == 'upcoming':
        meetings = [m for m in meetings if m.start_time and m.start_time >= now and m.status in ('pending', 'confirmed')]
    elif quick_filter == 'need_signin':
        meetings = [m for m in meetings if m.require_signin]
    elif quick_filter == 'pending_feedback':
        meetings = [m for m in meetings if _meeting_attendance_summary(m).get('pending', 0) > 0]

    visible_meetings = Meeting.query.all()
    if current_user.role not in ('admin', 'manager'):
        visible_meetings = [m for m in visible_meetings if m.creator_id == current_user.id or m.host_id == current_user.id or m.status != 'draft']
    stats = {
        'today': len([m for m in visible_meetings if m.start_time and m.start_time.date() == now.date() and m.status not in ('cancelled',)]),
        'upcoming': len([m for m in visible_meetings if m.start_time and m.start_time >= now and m.status in ('pending', 'confirmed')]),
        'in_progress': len([m for m in visible_meetings if m.status == 'in_progress']),
        'completed': len([m for m in visible_meetings if m.status == 'completed']),
    }
    attendance_overview = {'signed': 0, 'leave': 0, 'absent': 0, 'pending': 0}
    meeting_attendance_map = {}
    for meeting in meetings:
        summary = _meeting_attendance_summary(meeting)
        meeting_attendance_map[meeting.id] = summary
        for key in attendance_overview:
            attendance_overview[key] += summary.get(key, 0)

    meeting_manage_ids = {
        meeting.id for meeting in meetings
        if current_user.role in ('admin', 'manager') or current_user.id in [meeting.creator_id, meeting.host_id]
    }
    rooms = MeetingRoom.query.order_by(MeetingRoom.name.asc()).all()
    return render_template(
        'meeting/index.html',
        meetings=meetings,
        stats=stats,
        attendance_overview=attendance_overview,
        meeting_attendance_map=meeting_attendance_map,
        keyword=keyword,
        status_filter=status_filter,
        scope=scope,
        quick_filter=quick_filter,
        meeting_manage_ids=meeting_manage_ids,
        rooms=rooms,
    )




@app.route('/meeting/new', methods=['GET', 'POST'])
@login_required
def meeting_new():
    """发起会议"""
    related_doc = OfficialDoc.query.get(request.args.get('doc_id', type=int)) if request.args.get('doc_id') else None
    if request.method == 'POST':
        subject = request.form.get('subject', '').strip()
        start_time = _parse_datetime_local(request.form.get('start_time'))
        end_time = _parse_datetime_local(request.form.get('end_time'))
        attendee_ids = [int(i) for i in request.form.getlist('attendee_ids') if i]
        action = request.form.get('action') or 'save'
        room_id = request.form.get('room_id', type=int)
        host_id = request.form.get('host_id', type=int) or current_user.id
        room = MeetingRoom.query.get(room_id) if room_id else None
        if not subject or not start_time or not end_time:
            flash('请完整填写会议主题和起止时间', 'danger')
            return redirect(request.url)
        if end_time <= start_time:
            flash('会议结束时间必须晚于开始时间', 'danger')
            return redirect(request.url)
        if action == 'submit' and room and room.status != 'available':
            flash(f'会议室"{room.name}"当前状态不可排期，请更换会议室或先保存草稿', 'danger')
            return redirect(request.url)
        if action == 'submit' and room and room.capacity and attendee_ids and len(attendee_ids) > room.capacity:
            flash(f'当前会议室容量为 {room.capacity} 人，少于已选参会人员 {len(attendee_ids)} 人，请调整会场或参会范围', 'danger')
            return redirect(request.url)
        conflict_meeting = _find_meeting_room_conflict(room_id, start_time, end_time) if action == 'submit' else None
        if conflict_meeting:
            flash(
                f'会议室"{room.name if room else "已选会议室"}"在 {conflict_meeting.start_time.strftime("%m-%d %H:%M")} '
                f'至 {conflict_meeting.end_time.strftime("%H:%M")} 已安排"{conflict_meeting.subject}"，请调整时间或会议室',
                'danger'
            )
            return redirect(request.url)

        meeting = Meeting(
            subject=subject,
            meeting_type=request.form.get('meeting_type', '办公会'),
            level=request.form.get('level', '部门级'),
            priority=request.form.get('priority', '普通'),
            room_id=room_id,
            host_id=host_id,
            creator_id=current_user.id,
            attendee_depts=request.form.get('attendee_depts', '').strip(),
            agenda=request.form.get('agenda', '').strip(),
            minutes=request.form.get('minutes', '').strip(),
            related_doc_id=request.form.get('related_doc_id', type=int),
            require_signin=bool(request.form.get('require_signin')),
            start_time=start_time,
            end_time=end_time,
            status='pending' if action == 'submit' else 'draft',
        )

        meeting.set_attendee_ids(attendee_ids)
        db.session.add(meeting)
        db.session.flush()
        _sync_meeting_attendance(meeting)
        db.session.commit()
        _log_op('meeting', 'create', f'会议 #{meeting.subject}')
        flash('会议已发起' if meeting.status == 'pending' else '会议草稿已保存', 'success')
        return redirect(url_for('meeting_index'))

    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.dept_id, User.name).all()
    rooms = MeetingRoom.query.filter(MeetingRoom.status != 'disabled').order_by(MeetingRoom.name.asc()).all()
    return render_template('meeting/new.html', users=users, rooms=rooms, related_doc=related_doc, meeting=None)


@app.route('/meeting/edit/<int:meeting_id>', methods=['GET', 'POST'])
@login_required
def meeting_edit(meeting_id):
    """编辑会议安排"""
    meeting = Meeting.query.get_or_404(meeting_id)
    if current_user.role not in ('admin', 'manager') and current_user.id not in [meeting.creator_id, meeting.host_id]:
        flash('无权编辑该会议', 'danger')
        return redirect(url_for('meeting_detail', meeting_id=meeting.id))
    if meeting.status not in ('draft', 'pending', 'confirmed'):
        flash('当前会议状态下不支持编辑，请先完成或取消后再处理', 'warning')
        return redirect(url_for('meeting_detail', meeting_id=meeting.id))

    if request.method == 'POST':
        subject = request.form.get('subject', '').strip()
        start_time = _parse_datetime_local(request.form.get('start_time'))
        end_time = _parse_datetime_local(request.form.get('end_time'))
        attendee_ids = [int(i) for i in request.form.getlist('attendee_ids') if i]
        action = request.form.get('action') or 'save'
        room_id = request.form.get('room_id', type=int)
        host_id = request.form.get('host_id', type=int) or current_user.id
        room = MeetingRoom.query.get(room_id) if room_id else None
        if not subject or not start_time or not end_time:
            flash('请完整填写会议主题和起止时间', 'danger')
            return redirect(request.url)
        if end_time <= start_time:
            flash('会议结束时间必须晚于开始时间', 'danger')
            return redirect(request.url)

        should_validate_room = meeting.status in ('pending', 'confirmed') or action == 'submit'
        if should_validate_room and room and room.status != 'available':
            flash(f'会议室"{room.name}"当前状态不可排期，请更换会议室后再保存', 'danger')
            return redirect(request.url)
        if should_validate_room and room and room.capacity and attendee_ids and len(attendee_ids) > room.capacity:
            flash(f'当前会议室容量为 {room.capacity} 人，少于已选参会人员 {len(attendee_ids)} 人，请调整会场或参会范围', 'danger')
            return redirect(request.url)
        conflict_meeting = _find_meeting_room_conflict(room_id, start_time, end_time, exclude_meeting_id=meeting.id) if should_validate_room else None
        if conflict_meeting:
            flash(
                f'会议室"{room.name if room else "已选会议室"}"在 {conflict_meeting.start_time.strftime("%m-%d %H:%M")} '
                f'至 {conflict_meeting.end_time.strftime("%H:%M")} 已安排"{conflict_meeting.subject}"，请调整时间或会议室',
                'danger'
            )
            return redirect(request.url)

        previous_status = meeting.status
        meeting.subject = subject
        meeting.meeting_type = request.form.get('meeting_type', '办公会')
        meeting.level = request.form.get('level', '部门级')
        meeting.priority = request.form.get('priority', '普通')
        meeting.room_id = room_id
        meeting.host_id = host_id
        meeting.attendee_depts = request.form.get('attendee_depts', '').strip()
        meeting.agenda = request.form.get('agenda', '').strip()
        meeting.minutes = request.form.get('minutes', '').strip()
        meeting.related_doc_id = request.form.get('related_doc_id', type=int)
        meeting.require_signin = bool(request.form.get('require_signin'))
        meeting.start_time = start_time
        meeting.end_time = end_time
        if previous_status == 'draft':
            meeting.status = 'pending' if action == 'submit' else 'draft'
        meeting.set_attendee_ids(attendee_ids)

        _sync_meeting_attendance(meeting)
        db.session.commit()
        _log_op('meeting', 'edit', f'会议 #{meeting.subject}')
        flash('会议安排已更新', 'success')
        return redirect(url_for('meeting_detail', meeting_id=meeting.id))

    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.dept_id, User.name).all()
    rooms = MeetingRoom.query.filter(MeetingRoom.status != 'disabled').order_by(MeetingRoom.name.asc()).all()
    return render_template('meeting/new.html', users=users, rooms=rooms, related_doc=meeting.related_doc, meeting=meeting)


@app.route('/meeting/detail/<int:meeting_id>')
@login_required
def meeting_detail(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    if current_user.role not in ('admin', 'manager') and meeting.status == 'draft' and current_user.id not in [meeting.creator_id, meeting.host_id]:
        flash('无权查看该会议', 'danger')
        return redirect(url_for('meeting_index'))

    _sync_meeting_attendance(meeting)
    db.session.commit()

    attendee_ids = [int(uid) for uid in meeting.get_attendee_ids() if uid]
    attendee_users = User.query.filter(User.id.in_(attendee_ids)).all() if attendee_ids else []
    attendee_user_map = {user.id: user for user in attendee_users}
    attendance_map = {
        record.user_id: record for record in MeetingAttendance.query.filter_by(meeting_id=meeting.id).all()
    }
    attendee_rows = []
    for user_id in attendee_ids:
        attendee_rows.append({
            'user': attendee_user_map.get(user_id),
            'record': attendance_map.get(user_id),
        })

    can_manage_attendance = current_user.role in ('admin', 'manager') or current_user.id in [meeting.creator_id, meeting.host_id]
    my_attendance = attendance_map.get(current_user.id) if current_user.id in attendee_ids else None
    attendance_summary = _meeting_attendance_summary(meeting)
    return render_template(
        'meeting/detail.html',
        meeting=meeting,
        attendee_rows=attendee_rows,
        attendance_map=attendance_map,
        attendance_summary=attendance_summary,
        can_manage_attendance=can_manage_attendance,
        my_attendance=my_attendance,
    )


@app.route('/meeting/attendance/<int:meeting_id>', methods=['POST'])
@login_required
def meeting_update_attendance(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    if current_user.role not in ('admin', 'manager') and meeting.status == 'draft' and current_user.id not in [meeting.creator_id, meeting.host_id]:
        flash('当前会议尚未开放签到', 'warning')
        return redirect(url_for('meeting_detail', meeting_id=meeting_id))

    _sync_meeting_attendance(meeting)
    attendee_ids = [int(uid) for uid in meeting.get_attendee_ids() if uid]
    target_user_id = request.form.get('user_id', type=int) or current_user.id
    can_manage_attendance = current_user.role in ('admin', 'manager') or current_user.id in [meeting.creator_id, meeting.host_id]

    if target_user_id not in attendee_ids:
        flash('该人员不在参会范围内', 'danger')
        return redirect(url_for('meeting_detail', meeting_id=meeting_id))
    if target_user_id != current_user.id and not can_manage_attendance:
        flash('无权代他人维护签到状态', 'danger')
        return redirect(url_for('meeting_detail', meeting_id=meeting_id))

    record = MeetingAttendance.query.filter_by(meeting_id=meeting.id, user_id=target_user_id).first()
    if not record:
        record = MeetingAttendance(meeting_id=meeting.id, user_id=target_user_id, attendance_status='pending')
        db.session.add(record)

    action = request.form.get('action', 'feedback')
    feedback = request.form.get('feedback', '').strip()
    remark = request.form.get('remark', '').strip()
    message = '参会反馈已保存'

    if action == 'sign_in':
        record.attendance_status = 'signed'
        record.signin_time = datetime.now()
        message = '签到已完成'
    elif action == 'leave':
        record.attendance_status = 'leave'
        record.signin_time = None
        message = '请假反馈已提交'
    elif action == 'update_status':
        if not can_manage_attendance:
            flash('无权维护参会状态', 'danger')
            return redirect(url_for('meeting_detail', meeting_id=meeting_id))
        new_status = request.form.get('attendance_status', 'pending')
        if new_status not in ('pending', 'signed', 'leave', 'absent'):
            flash('无效的参会状态', 'warning')
            return redirect(url_for('meeting_detail', meeting_id=meeting_id))
        record.attendance_status = new_status
        record.signin_time = datetime.now() if new_status == 'signed' else None
        message = '参会状态已更新'

    if feedback:
        record.feedback = feedback
    if remark:
        record.remark = remark

    db.session.commit()
    _log_op('meeting', 'attendance', f'会议 #{meeting.subject}')
    flash(message, 'success')
    return redirect(url_for('meeting_detail', meeting_id=meeting_id))


@app.route('/meeting/status/<int:meeting_id>', methods=['POST'])
@login_required
def meeting_update_status(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    if current_user.role not in ('admin', 'manager') and current_user.id not in [meeting.creator_id, meeting.host_id]:
        flash('无权操作该会议', 'danger')
        return redirect(url_for('meeting_detail', meeting_id=meeting_id))

    action = request.form.get('action')
    status_map = {
        'confirm': 'confirmed',
        'start': 'in_progress',
        'complete': 'completed',
        'cancel': 'cancelled',
        'reopen': 'confirmed',
    }
    new_status = status_map.get(action)
    if not new_status:
        flash('无效操作', 'warning')
        return redirect(url_for('meeting_detail', meeting_id=meeting_id))

    _sync_meeting_attendance(meeting)
    meeting.status = new_status
    auto_absent_count = 0
    if action == 'complete':
        if request.form.get('minutes', '').strip():
            meeting.minutes = request.form.get('minutes', '').strip()
        if meeting.require_signin:
            pending_records = MeetingAttendance.query.filter_by(meeting_id=meeting.id, attendance_status='pending').all()
            for record in pending_records:
                record.attendance_status = 'absent'
                record.remark = record.remark or '会议结束后系统自动标记为缺席'
                auto_absent_count += 1
    db.session.commit()
    _log_op('meeting', action, f'会议 #{meeting.subject}')
    flash('会议状态已更新' + (f'，并自动标记 {auto_absent_count} 人缺席' if auto_absent_count else ''), 'success')
    return redirect(url_for('meeting_detail', meeting_id=meeting_id))



@app.route('/meeting/calendar')
@login_required
def meeting_calendar():
    """会议日历视图"""
    _track_usage('meeting', 'calendar')
    start = _parse_date_local(request.args.get('start_date')) or datetime.now().date()
    end = _parse_date_local(request.args.get('end_date'))
    if not end:
        end = datetime(start.year, start.month, start.day).date()
        from datetime import timedelta
        end = end + timedelta(days=14)

    q = Meeting.query.filter(Meeting.start_time != None)
    if current_user.role not in ('admin', 'manager'):
        q = q.filter(db.or_(Meeting.creator_id == current_user.id, Meeting.host_id == current_user.id, Meeting.status != 'draft'))
    meetings = q.order_by(Meeting.start_time.asc()).all()
    meetings = [m for m in meetings if start <= m.start_time.date() <= end]

    grouped = {}
    for meeting in meetings:
        day_key = meeting.start_time.strftime('%Y-%m-%d')
        grouped.setdefault(day_key, []).append(meeting)

    return render_template('meeting/calendar.html', grouped=grouped, start_date=start, end_date=end)


@app.route('/meeting/rooms', methods=['GET', 'POST'])
@login_required
def meeting_rooms():
    """会议室管理"""
    if request.method == 'POST':
        if current_user.role not in ('admin', 'manager'):
            flash('仅管理员或部门负责人可维护会议室', 'danger')
            return redirect(url_for('meeting_rooms'))
        room = MeetingRoom(
            name=request.form.get('name', '').strip(),
            location=request.form.get('location', '').strip(),
            capacity=request.form.get('capacity', type=int) or 0,
            equipment=request.form.get('equipment', '').strip(),
            manager_id=request.form.get('manager_id', type=int),
            status=request.form.get('status', 'available'),
            remark=request.form.get('remark', '').strip(),
        )
        if not room.name:
            flash('会议室名称不能为空', 'danger')
            return redirect(url_for('meeting_rooms'))
        db.session.add(room)
        db.session.commit()
        _log_op('meeting_room', 'create', room.name)
        flash('会议室已保存', 'success')
        return redirect(url_for('meeting_rooms'))

    rooms = MeetingRoom.query.order_by(MeetingRoom.status.asc(), MeetingRoom.name.asc()).all()
    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.name).all()
    return render_template('meeting/rooms.html', rooms=rooms, users=users)


@app.route('/supervision')
@login_required
def supervision_index():
    """督查督办首页"""
    _track_usage('supervision', 'index')
    _refresh_supervision_alerts()

    keyword = request.args.get('keyword', '').strip()
    status_filter = request.args.get('status', '').strip()
    scope = request.args.get('scope', 'all').strip()
    quick_filter = request.args.get('quick', '').strip()
    now = datetime.now()

    visible_tasks = SupervisionTask.query.order_by(SupervisionTask.due_date.asc(), SupervisionTask.created_at.desc()).all()
    if current_user.role not in ('admin', 'manager'):
        visible_tasks = [t for t in visible_tasks if t.creator_id == current_user.id or t.owner_id == current_user.id or current_user.id in t.get_helper_ids()]

    tasks = list(visible_tasks)
    if scope == 'mine':
        tasks = [t for t in tasks if t.owner_id == current_user.id]
    elif scope == 'created':
        tasks = [t for t in tasks if t.creator_id == current_user.id]

    if status_filter == 'overdue':
        tasks = [t for t in tasks if _get_supervision_alert_meta(t, now)['is_overdue']]
    elif status_filter == 'due_soon':
        tasks = [t for t in tasks if _get_supervision_alert_meta(t, now)['is_due_soon']]
    elif status_filter:
        tasks = [t for t in tasks if t.status == status_filter]
    if keyword:
        tasks = [t for t in tasks if keyword.lower() in (t.title or '').lower() or keyword.lower() in (t.content or '').lower()]

    if quick_filter == 'my_due_soon':
        tasks = [t for t in tasks if t.owner_id == current_user.id and _get_supervision_alert_meta(t, now)['is_due_soon']]
    elif quick_filter == 'my_overdue':
        tasks = [t for t in tasks if t.owner_id == current_user.id and _get_supervision_alert_meta(t, now)['is_overdue']]
    elif quick_filter == 'high_priority':
        tasks = [t for t in tasks if t.priority == '高']
    elif quick_filter == 'drafts':
        tasks = [t for t in tasks if t.status == 'draft']
    elif quick_filter == 'waiting_review':
        tasks = [t for t in tasks if t.status == 'review']

    stats = {
        'issued': len([t for t in visible_tasks if t.status == 'issued']),
        'processing': len([t for t in visible_tasks if t.status == 'processing']),
        'review': len([t for t in visible_tasks if t.status == 'review']),
        'overdue': len([t for t in visible_tasks if _get_supervision_alert_meta(t, now)['is_overdue']]),
        'due_soon': len([t for t in visible_tasks if _get_supervision_alert_meta(t, now)['is_due_soon']]),
    }
    task_alerts = {task.id: _get_supervision_alert_meta(task, now) for task in tasks}
    task_manage_ids = {
        task.id for task in tasks
        if current_user.role in ('admin', 'manager') or current_user.id == task.creator_id
    }
    leader_board, risk_tasks = _build_supervision_leader_board(visible_tasks, now) if current_user.role in ('admin', 'manager') else ([], [])
    return render_template(
        'supervision/index.html',
        tasks=tasks,
        stats=stats,
        task_alerts=task_alerts,
        leader_board=leader_board,
        risk_tasks=risk_tasks,
        keyword=keyword,
        status_filter=status_filter,
        scope=scope,
        quick_filter=quick_filter,
        task_manage_ids=task_manage_ids,
    )




@app.route('/supervision/new', methods=['GET', 'POST'])
@login_required
def supervision_new():
    """新建督办任务"""
    source_doc = OfficialDoc.query.get(request.args.get('doc_id', type=int)) if request.args.get('doc_id') else None
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        owner_id = request.form.get('owner_id', type=int)
        due_date = _parse_datetime_local(request.form.get('due_date'))
        helper_ids = [int(i) for i in request.form.getlist('helper_ids') if i]
        if not title or not owner_id or not due_date:
            flash('请至少填写标题、负责人和截止时间', 'danger')
            return redirect(request.url)

        task = SupervisionTask(
            task_no=_gen_supervision_no(),
            title=title,
            category=request.form.get('category', '重点工作'),
            source=request.form.get('source', '领导交办'),
            content=request.form.get('content', '').strip(),
            creator_id=current_user.id,
            owner_id=owner_id,
            priority=request.form.get('priority', '中'),
            status='issued' if request.form.get('action') == 'issue' else 'draft',
            progress_percent=0,
            due_date=due_date,
            result_summary=request.form.get('result_summary', '').strip(),
            source_doc_id=request.form.get('source_doc_id', type=int),
        )
        task.set_helper_ids(helper_ids)
        db.session.add(task)
        db.session.flush()
        db.session.add(SupervisionProgress(
            task_id=task.id,
            operator_id=current_user.id,
            action='issue' if task.status == 'issued' else 'create',
            progress_percent=task.progress_percent,
            note='新建督办任务' if task.status == 'draft' else '任务已下发，请按期办理',
        ))
        db.session.commit()
        _log_op('supervision', 'create', task.task_no)
        flash('督办任务已下发' if task.status == 'issued' else '督办任务草稿已保存', 'success')
        return redirect(url_for('supervision_index'))

    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.dept_id, User.name).all()
    return render_template('supervision/new.html', users=users, source_doc=source_doc, task_form=None)


@app.route('/supervision/edit/<int:task_id>', methods=['GET', 'POST'])
@login_required
def supervision_edit(task_id):
    """编辑督办任务"""
    task = SupervisionTask.query.get_or_404(task_id)
    if current_user.role not in ('admin', 'manager') and current_user.id != task.creator_id:
        flash('仅下达人或管理员可编辑该任务', 'danger')
        return redirect(url_for('supervision_detail', task_id=task.id))
    if task.status in ('completed', 'cancelled'):
        flash('已关闭任务不支持继续编辑', 'warning')
        return redirect(url_for('supervision_detail', task_id=task.id))

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        owner_id = request.form.get('owner_id', type=int)
        due_date = _parse_datetime_local(request.form.get('due_date'))
        helper_ids = [int(i) for i in request.form.getlist('helper_ids') if i]
        if not title or not owner_id or not due_date:
            flash('请至少填写标题、负责人和截止时间', 'danger')
            return redirect(request.url)

        previous_status = task.status
        action = request.form.get('action') or 'save'
        task.title = title
        task.category = request.form.get('category', '重点工作')
        task.source = request.form.get('source', '领导交办')
        task.content = request.form.get('content', '').strip()
        task.owner_id = owner_id
        task.priority = request.form.get('priority', '中')
        task.due_date = due_date
        task.result_summary = request.form.get('result_summary', '').strip()
        task.source_doc_id = request.form.get('source_doc_id', type=int)
        if previous_status == 'draft' and action == 'issue':
            task.status = 'issued'
        task.set_helper_ids(helper_ids)

        db.session.add(SupervisionProgress(
            task_id=task.id,
            operator_id=current_user.id,
            action='edit',
            progress_percent=task.progress_percent or 0,
            note='已更新督办任务信息',
        ))
        db.session.commit()
        _log_op('supervision', 'edit', task.task_no)
        flash('督办任务已更新', 'success')
        return redirect(url_for('supervision_detail', task_id=task.id))

    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.dept_id, User.name).all()
    return render_template('supervision/new.html', users=users, source_doc=task.source_doc, task_form=task)


@app.route('/supervision/batch_remind', methods=['POST'])
@login_required
def supervision_batch_remind():
    raw_ids = request.form.getlist('task_ids')
    task_ids = [int(item) for item in raw_ids if str(item).isdigit()]
    if not task_ids:
        flash('请先选择需要催办的任务', 'warning')
        return redirect(url_for('supervision_index'))

    tasks = SupervisionTask.query.filter(SupervisionTask.id.in_(task_ids)).order_by(SupervisionTask.due_date.asc()).all()
    note = request.form.get('note', '').strip()
    reminded_count = 0
    skipped_count = 0
    for task in tasks:
        if current_user.role not in ('admin', 'manager') and current_user.id != task.creator_id:
            skipped_count += 1
            continue
        if task.status in ('completed', 'cancelled'):
            skipped_count += 1
            continue
        db.session.add(SupervisionProgress(
            task_id=task.id,
            operator_id=current_user.id,
            action='manual_remind',
            progress_percent=task.progress_percent or 0,
            note=note or f'批量催办：请负责人尽快反馈"{task.title}"的办理进展，确保按期办结。',
        ))
        reminded_count += 1

    if reminded_count:
        db.session.commit()
        _log_op('supervision', 'batch_remind', f'{reminded_count} tasks')
        flash(f'已发送 {reminded_count} 条催办提醒' + (f'，另有 {skipped_count} 条因权限或状态被跳过' if skipped_count else ''), 'success')
    else:
        flash('未发送催办提醒，请检查所选任务权限或状态', 'warning')
    return redirect(request.referrer or url_for('supervision_index'))


@app.route('/supervision/close/<int:task_id>', methods=['POST'])
@login_required
def supervision_close(task_id):
    task = SupervisionTask.query.get_or_404(task_id)
    if current_user.role not in ('admin', 'manager') and current_user.id != task.creator_id:
        flash('仅下达人或管理员可关闭该任务', 'danger')
        return redirect(url_for('supervision_detail', task_id=task.id))
    if task.status in ('completed', 'cancelled'):
        flash('该任务已关闭，无需重复操作', 'warning')
        return redirect(url_for('supervision_detail', task_id=task.id))

    close_type = request.form.get('close_type', 'completed')
    if close_type not in ('completed', 'cancelled'):
        flash('无效的关闭类型', 'warning')
        return redirect(url_for('supervision_detail', task_id=task.id))

    note = request.form.get('note', '').strip() or ('任务已确认办结' if close_type == 'completed' else '任务已关闭归档')
    result_summary = request.form.get('result_summary', '').strip()
    task.status = close_type
    if close_type == 'completed':
        task.progress_percent = 100
        task.completed_at = datetime.now()
    else:
        task.completed_at = None
    if result_summary:
        task.result_summary = result_summary

    db.session.add(SupervisionProgress(
        task_id=task.id,
        operator_id=current_user.id,
        action='close',
        progress_percent=task.progress_percent or 0,
        note=note,
    ))
    db.session.commit()
    _log_op('supervision', 'close', task.task_no)
    flash('督办任务已办结' if close_type == 'completed' else '督办任务已关闭', 'success')
    return redirect(url_for('supervision_detail', task_id=task.id))


@app.route('/supervision/detail/<int:task_id>')

@login_required
def supervision_detail(task_id):
    task = SupervisionTask.query.get_or_404(task_id)
    if current_user.role not in ('admin', 'manager') and current_user.id not in [task.creator_id, task.owner_id] and current_user.id not in task.get_helper_ids():
        flash('无权查看该督办任务', 'danger')
        return redirect(url_for('supervision_index'))

    _refresh_supervision_alerts([task])
    helpers = User.query.filter(User.id.in_(task.get_helper_ids())).all() if task.get_helper_ids() else []
    logs = task.progress_logs.order_by(SupervisionProgress.created_at.desc()).all()
    alert_meta = _get_supervision_alert_meta(task)
    can_remind = current_user.role in ('admin', 'manager') or current_user.id == task.creator_id
    can_manage = can_remind
    return render_template('supervision/detail.html', task=task, helpers=helpers, logs=logs, alert_meta=alert_meta, can_remind=can_remind, can_manage=can_manage)



@app.route('/supervision/remind/<int:task_id>', methods=['POST'])
@login_required
def supervision_remind(task_id):
    task = SupervisionTask.query.get_or_404(task_id)
    if current_user.role not in ('admin', 'manager') and current_user.id != task.creator_id:
        flash('仅下达人或管理员可催办该任务', 'danger')
        return redirect(url_for('supervision_detail', task_id=task_id))
    if task.status in ('completed', 'cancelled'):
        flash('当前任务状态无需催办', 'warning')
        return redirect(url_for('supervision_detail', task_id=task_id))

    note = request.form.get('note', '').strip() or '督办催办：请负责人尽快反馈办理进度，确保按期办结。'
    db.session.add(SupervisionProgress(
        task_id=task.id,
        operator_id=current_user.id,
        action='manual_remind',
        progress_percent=task.progress_percent or 0,
        note=note,
    ))
    db.session.commit()
    _log_op('supervision', 'remind', task.task_no)
    flash('催办提醒已发送', 'success')
    return redirect(url_for('supervision_detail', task_id=task_id))


@app.route('/supervision/progress/<int:task_id>', methods=['POST'])
@login_required
def supervision_progress(task_id):
    task = SupervisionTask.query.get_or_404(task_id)
    if current_user.role not in ('admin', 'manager') and current_user.id not in [task.creator_id, task.owner_id] and current_user.id not in task.get_helper_ids():
        flash('无权更新该督办任务', 'danger')
        return redirect(url_for('supervision_detail', task_id=task_id))

    status = request.form.get('status', task.status)
    progress_percent = request.form.get('progress_percent', type=int)
    note = request.form.get('note', '').strip()
    result_summary = request.form.get('result_summary', '').strip()

    if progress_percent is None:
        progress_percent = task.progress_percent or 0
    progress_percent = max(0, min(progress_percent, 100))

    if status == 'issued' and progress_percent > 0:
        status = 'processing'

    task.status = status
    task.progress_percent = progress_percent
    if result_summary:
        task.result_summary = result_summary
    if status == 'completed' or progress_percent == 100:
        task.status = 'completed'
        task.progress_percent = 100
        task.completed_at = datetime.now()
    else:
        task.completed_at = None

    db.session.add(SupervisionProgress(
        task_id=task.id,
        operator_id=current_user.id,
        action='update',
        progress_percent=task.progress_percent,
        note=note or f'任务状态更新为 {task.status}',
    ))
    db.session.commit()
    _log_op('supervision', 'progress', task.task_no)
    flash('督办进展已更新', 'success')
    return redirect(url_for('supervision_detail', task_id=task_id))



@app.route('/performance')
@login_required
def performance_index():
    """绩效考核台账"""
    _track_usage('performance', 'index')
    period_id = request.args.get('period_id', type=int)
    status_filter = request.args.get('status', '').strip()

    q = PerformanceAssessment.query
    if current_user.role not in ('admin', 'manager'):
        q = q.filter(PerformanceAssessment.user_id == current_user.id)
    if period_id:
        q = q.filter(PerformanceAssessment.period_id == period_id)
    if status_filter:
        q = q.filter(PerformanceAssessment.status == status_filter)

    records = q.order_by(PerformanceAssessment.created_at.desc()).all()
    periods = PerformancePeriod.query.order_by(PerformancePeriod.start_date.desc(), PerformancePeriod.created_at.desc()).all()
    scores = [r.score for r in records if r.score is not None]
    stats = {
        'total': len(records),
        'published': len([r for r in records if r.status == 'published']),
        'avg_score': round(sum(scores) / len(scores), 1) if scores else 0,
        'best_score': max(scores) if scores else 0,
    }
    dashboard = _build_performance_dashboard(records)
    return render_template(
        'performance/index.html',
        records=records,
        periods=periods,
        period_id=period_id,
        status_filter=status_filter,
        stats=stats,
        dashboard=dashboard,
        can_manage_records=current_user.role in ('admin', 'manager'),
    )




@app.route('/performance/new', methods=['GET', 'POST'])
@login_required
def performance_new():
    """新增绩效记录"""
    if current_user.role not in ('admin', 'manager'):
        flash('仅管理员或部门负责人可录入绩效记录', 'danger')
        return redirect(url_for('performance_index'))

    if request.method == 'POST':
        project_name = request.form.get('project_name', '').strip()
        user_id = request.form.get('user_id', type=int)
        period_id = request.form.get('period_id', type=int)
        action = request.form.get('action') or 'save'
        score = _parse_float_local(request.form.get('score'), 0)
        full_score = _parse_float_local(request.form.get('full_score'), 100)
        weight = _parse_float_local(request.form.get('weight'), 1)
        period = PerformancePeriod.query.get(period_id) if period_id else None
        if not project_name or not user_id:
            flash('请填写考核项目和被考核人', 'danger')
            return redirect(request.url)
        if score is None or full_score is None or weight is None:
            flash('得分、满分和权重必须为合法数字', 'danger')
            return redirect(request.url)
        if full_score <= 0:
            flash('满分必须大于 0', 'danger')
            return redirect(request.url)
        if weight <= 0:
            flash('权重必须大于 0', 'danger')
            return redirect(request.url)
        if score < 0 or score > full_score:
            flash('得分不能小于 0，且不能高于满分', 'danger')
            return redirect(request.url)
        if period_id and not period:
            flash('所选考核周期不存在，请刷新页面后重试', 'danger')
            return redirect(request.url)
        if action == 'publish' and period and period.status != 'active':
            flash('所选考核周期当前未启用，不能直接发布绩效记录', 'danger')
            return redirect(request.url)

        record = PerformanceAssessment(
            period_id=period_id,
            user_id=user_id,
            assessor_id=current_user.id,
            project_name=project_name,
            category=request.form.get('category', '重点工作'),
            score=score,
            full_score=full_score,
            weight=weight,
            evaluation=request.form.get('evaluation', '').strip(),
            highlights=request.form.get('highlights', '').strip(),
            status='published' if action == 'publish' else 'draft',
        )

        db.session.add(record)
        db.session.commit()
        _log_op('performance', 'create', project_name)
        flash('绩效记录已保存', 'success')
        return redirect(url_for('performance_index'))

    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.dept_id, User.name).all()
    periods = PerformancePeriod.query.order_by(PerformancePeriod.start_date.desc(), PerformancePeriod.created_at.desc()).all()
    return render_template('performance/new.html', users=users, periods=periods, record=None)


@app.route('/performance/edit/<int:record_id>', methods=['GET', 'POST'])
@login_required
def performance_edit(record_id):
    """编辑绩效记录"""
    if current_user.role not in ('admin', 'manager'):
        flash('仅管理员或部门负责人可编辑绩效记录', 'danger')
        return redirect(url_for('performance_index'))

    record = PerformanceAssessment.query.get_or_404(record_id)
    if request.method == 'POST':
        project_name = request.form.get('project_name', '').strip()
        user_id = request.form.get('user_id', type=int)
        period_id = request.form.get('period_id', type=int)
        action = request.form.get('action') or 'save'
        score = _parse_float_local(request.form.get('score'), 0)
        full_score = _parse_float_local(request.form.get('full_score'), 100)
        weight = _parse_float_local(request.form.get('weight'), 1)
        period = PerformancePeriod.query.get(period_id) if period_id else None
        if not project_name or not user_id:
            flash('请填写考核项目和被考核人', 'danger')
            return redirect(request.url)
        if score is None or full_score is None or weight is None:
            flash('得分、满分和权重必须为合法数字', 'danger')
            return redirect(request.url)
        if full_score <= 0:
            flash('满分必须大于 0', 'danger')
            return redirect(request.url)
        if weight <= 0:
            flash('权重必须大于 0', 'danger')
            return redirect(request.url)
        if score < 0 or score > full_score:
            flash('得分不能小于 0，且不能高于满分', 'danger')
            return redirect(request.url)
        if period_id and not period:
            flash('所选考核周期不存在，请刷新页面后重试', 'danger')
            return redirect(request.url)
        if action == 'publish' and period and period.status != 'active':
            flash('所选考核周期当前未启用，不能直接发布绩效记录', 'danger')
            return redirect(request.url)

        record.period_id = period_id
        record.user_id = user_id
        record.project_name = project_name
        record.category = request.form.get('category', '重点工作')
        record.score = score
        record.full_score = full_score
        record.weight = weight
        record.evaluation = request.form.get('evaluation', '').strip()
        record.highlights = request.form.get('highlights', '').strip()
        if record.status == 'draft':
            record.status = 'published' if action == 'publish' else 'draft'

        db.session.commit()
        _log_op('performance', 'edit', project_name)
        flash('绩效记录已更新', 'success')
        return redirect(url_for('performance_index', period_id=record.period_id or None))

    users = User.query.filter(db.or_(User.is_active == True, User.is_active == None)).order_by(User.dept_id, User.name).all()
    periods = PerformancePeriod.query.order_by(PerformancePeriod.start_date.desc(), PerformancePeriod.created_at.desc()).all()
    return render_template('performance/new.html', users=users, periods=periods, record=record)


@app.route('/performance/retract/<int:record_id>', methods=['POST'])
@login_required
def performance_retract(record_id):
    if current_user.role not in ('admin', 'manager'):
        flash('仅管理员或部门负责人可撤回绩效记录', 'danger')
        return redirect(url_for('performance_index'))

    record = PerformanceAssessment.query.get_or_404(record_id)
    if record.status != 'published':
        flash('当前绩效记录尚未发布，无需撤回', 'warning')
        return redirect(url_for('performance_index', period_id=record.period_id or None))

    record.status = 'draft'
    db.session.commit()
    _log_op('performance', 'retract', record.project_name)
    flash('绩效记录已撤回为草稿', 'success')
    return redirect(url_for('performance_index', period_id=record.period_id or None))


@app.route('/performance/period/toggle/<int:period_id>', methods=['POST'])
@login_required
def performance_period_toggle(period_id):
    if current_user.role not in ('admin', 'manager'):
        flash('仅管理员或部门负责人可切换考核周期状态', 'danger')
        return redirect(url_for('performance_periods'))

    period = PerformancePeriod.query.get_or_404(period_id)
    target_status = request.form.get('status', 'active' if period.status != 'active' else 'closed')
    if target_status not in ('active', 'closed'):
        flash('无效的周期状态', 'warning')
        return redirect(url_for('performance_periods'))

    if target_status == 'active':
        overlapping_period = PerformancePeriod.query.filter(
            PerformancePeriod.id != period.id,
            PerformancePeriod.period_type == period.period_type,
            PerformancePeriod.status == 'active',
            PerformancePeriod.start_date.isnot(None),
            PerformancePeriod.end_date.isnot(None),
            PerformancePeriod.start_date <= period.end_date,
            PerformancePeriod.end_date >= period.start_date,
        ).order_by(PerformancePeriod.start_date.asc()).first()
        if overlapping_period:
            flash(f'与已启用周期"{overlapping_period.name}"时间重叠，请先关闭原周期或调整时间范围', 'danger')
            return redirect(url_for('performance_periods'))

    period.status = target_status
    db.session.commit()
    _log_op('performance', 'toggle_period', f'{period.name}:{target_status}')
    flash('考核周期已启用' if target_status == 'active' else '考核周期已关闭', 'success')
    return redirect(url_for('performance_periods'))


@app.route('/performance/periods', methods=['GET', 'POST'])
@login_required
def performance_periods():
    """考核周期管理"""
    if request.method == 'POST':
        if current_user.role not in ('admin', 'manager'):
            flash('仅管理员或部门负责人可维护考核周期', 'danger')
            return redirect(url_for('performance_periods'))

        name = request.form.get('name', '').strip()
        period_type = request.form.get('period_type', 'monthly')
        start_date = _parse_date_local(request.form.get('start_date'))
        end_date = _parse_date_local(request.form.get('end_date'))
        status = request.form.get('status', 'active')
        remark = request.form.get('remark', '').strip()
        if not name:
            flash('周期名称不能为空', 'danger')
            return redirect(url_for('performance_periods'))
        if not start_date or not end_date:
            flash('请完整填写开始日期和结束日期', 'danger')
            return redirect(url_for('performance_periods'))
        if end_date < start_date:
            flash('结束日期不能早于开始日期', 'danger')
            return redirect(url_for('performance_periods'))
        if PerformancePeriod.query.filter(PerformancePeriod.name == name).first():
            flash('已存在同名考核周期，请直接使用或更换名称', 'danger')
            return redirect(url_for('performance_periods'))
        if status == 'active':
            overlapping_period = PerformancePeriod.query.filter(
                PerformancePeriod.period_type == period_type,
                PerformancePeriod.status == 'active',
                PerformancePeriod.start_date.isnot(None),
                PerformancePeriod.end_date.isnot(None),
                PerformancePeriod.start_date <= end_date,
                PerformancePeriod.end_date >= start_date,
            ).order_by(PerformancePeriod.start_date.asc()).first()
            if overlapping_period:
                flash(f'与已启用周期"{overlapping_period.name}"时间重叠，请先关闭原周期或调整时间范围', 'danger')
                return redirect(url_for('performance_periods'))

        period = PerformancePeriod(
            name=name,
            period_type=period_type,
            start_date=start_date,
            end_date=end_date,
            status=status,
            remark=remark,
            created_by=current_user.id,
        )
        db.session.add(period)
        db.session.commit()
        _log_op('performance', 'create_period', period.name)
        flash('考核周期已保存', 'success')
        return redirect(url_for('performance_periods'))


    periods = PerformancePeriod.query.order_by(PerformancePeriod.start_date.desc(), PerformancePeriod.created_at.desc()).all()
    return render_template('performance/periods.html', periods=periods)


@app.route('/worklog')
@login_required
def worklog_index():
    """工作日志列表"""
    _track_usage('worklog', 'index')
    keyword = request.args.get('keyword', '').strip()
    status_filter = request.args.get('status', '').strip()
    log_date = _parse_date_local(request.args.get('log_date'))
    quick_filter = request.args.get('quick', '').strip()

    q = WorkLog.query
    if current_user.role not in ('admin', 'manager'):
        q = q.filter(WorkLog.user_id == current_user.id)
    if status_filter:
        q = q.filter(WorkLog.status == status_filter)
    if log_date:
        q = q.filter(WorkLog.log_date == log_date)
    if keyword:
        q = q.filter(db.or_(WorkLog.title.like(f'%{keyword}%'), WorkLog.content.like(f'%{keyword}%'), WorkLog.achievements.like(f'%{keyword}%')))

    today = datetime.now().date()
    if quick_filter == 'today':
        q = q.filter(WorkLog.log_date == today)
    elif quick_filter == 'returned':
        q = q.filter(WorkLog.status == 'returned')
    elif quick_filter == 'task_related':
        q = q.filter(WorkLog.related_task_id.isnot(None))
    elif quick_filter == 'need_review' and current_user.role in ('admin', 'manager'):
        q = q.filter(WorkLog.status == 'submitted')

    logs = q.order_by(WorkLog.log_date.desc(), WorkLog.created_at.desc()).all()
    visible_logs = q.all()

    stats = {
        'today': len([i for i in visible_logs if i.log_date == today]),
        'submitted': len([i for i in visible_logs if i.status == 'submitted']),
        'returned': len([i for i in visible_logs if i.status == 'returned']),
        'reviewed': len([i for i in visible_logs if i.status == 'reviewed']),
        'task_related': len([i for i in visible_logs if i.related_task_id]),
        'hours': round(sum([i.hours or 0 for i in visible_logs]), 1),
    }
    latest_reviews = {
        log.id: log.review_logs.order_by(WorkLogReview.created_at.desc()).first() for log in logs
    }
    tasks = SupervisionTask.query.order_by(SupervisionTask.due_date.asc(), SupervisionTask.created_at.desc()).all()
    if current_user.role not in ('admin', 'manager'):
        tasks = [t for t in tasks if t.owner_id == current_user.id or current_user.id in t.get_helper_ids() or t.creator_id == current_user.id]
    return render_template(
        'work_log/index.html',
        logs=logs,
        stats=stats,
        latest_reviews=latest_reviews,
        keyword=keyword,
        status_filter=status_filter,
        log_date=log_date,
        quick_filter=quick_filter,
        tasks=tasks,
    )




@app.route('/worklog/new', methods=['GET', 'POST'])
@login_required
def worklog_new():
    """新建工作日志"""
    log_id = request.args.get('id', type=int)
    log = WorkLog.query.get(log_id) if log_id else None
    if log and log.user_id != current_user.id and current_user.role not in ('admin', 'manager'):
        flash('无权编辑该日志', 'danger')
        return redirect(url_for('worklog_index'))

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        if not title:
            flash('日志标题不能为空', 'danger')
            return redirect(request.url)

        action = request.form.get('action', 'draft').strip().lower()
        if action not in ('draft', 'submit'):
            action = 'draft'
        previous_status = log.status if log else None
        is_new = False


        hours = _parse_float_local(request.form.get('hours'), 8)

        if hours is None:
            flash('投入工时必须为合法数字', 'danger')
            return redirect(request.url)
        if hours < 0 or hours > 24:
            flash('投入工时必须在 0 到 24 小时之间', 'danger')
            return redirect(request.url)

        if not log:
            log = WorkLog(user_id=current_user.id)
            db.session.add(log)
            is_new = True

        log.log_date = _parse_date_local(request.form.get('log_date')) or datetime.now().date()

        log.title = title
        log.category = request.form.get('category', '日常工作')
        log.content = request.form.get('content', '').strip()
        log.achievements = request.form.get('achievements', '').strip()
        log.issues = request.form.get('issues', '').strip()
        log.tomorrow_plan = request.form.get('tomorrow_plan', '').strip()
        log.hours = hours
        log.related_task_id = request.form.get('related_task_id', type=int)

        review_action = 'save_draft'

        flash_message = '工作日志草稿已保存'
        if action == 'submit':
            log.status = 'submitted'
            review_action = 'resubmit' if previous_status == 'returned' else 'submit'
            flash_message = '工作日志已重新提交' if previous_status == 'returned' else '工作日志已提交'
        else:
            log.status = 'draft'

        db.session.flush()
        should_log = is_new or previous_status != log.status or review_action in ('submit', 'resubmit')
        if should_log:
            _add_worklog_review_log(log.id, current_user.id, review_action, previous_status, log.status)
        db.session.commit()
        _log_op('worklog', review_action, title)
        flash(flash_message, 'success')
        return redirect(url_for('worklog_detail', log_id=log.id))

    tasks = SupervisionTask.query.order_by(SupervisionTask.due_date.asc(), SupervisionTask.created_at.desc()).all()
    if current_user.role not in ('admin', 'manager'):
        tasks = [t for t in tasks if t.owner_id == current_user.id or current_user.id in t.get_helper_ids() or t.creator_id == current_user.id]
    latest_return = log.review_logs.filter_by(action='return').order_by(WorkLogReview.created_at.desc()).first() if log else None
    review_logs = log.review_logs.order_by(WorkLogReview.created_at.desc()).all() if log else []
    return render_template('work_log/new.html', log=log, tasks=tasks, latest_return=latest_return, review_logs=review_logs)



@app.route('/worklog/detail/<int:log_id>')
@login_required
def worklog_detail(log_id):
    log = WorkLog.query.get_or_404(log_id)
    if current_user.role not in ('admin', 'manager') and log.user_id != current_user.id:
        flash('无权查看该日志', 'danger')
        return redirect(url_for('worklog_index'))
    review_logs = log.review_logs.order_by(WorkLogReview.created_at.desc()).all()
    return render_template('work_log/detail.html', log=log, review_logs=review_logs)


@app.route('/worklog/review/<int:log_id>', methods=['POST'])
@login_required
def worklog_review(log_id):
    log = WorkLog.query.get_or_404(log_id)
    if current_user.role not in ('admin', 'manager'):
        flash('无审核权限', 'danger')
        return redirect(url_for('worklog_detail', log_id=log_id))

    action = request.form.get('action', 'review').strip()
    comment = request.form.get('comment', '').strip()
    previous_status = log.status

    if action == 'return':
        log.status = 'returned'
        _add_worklog_review_log(log.id, current_user.id, 'return', previous_status, log.status, comment or '审核退回，请补充完善后重新提交。')
        message = '工作日志已退回修改'
        op_action = 'return'
    else:
        log.status = 'reviewed'
        _add_worklog_review_log(log.id, current_user.id, 'review', previous_status, log.status, comment or '日志已阅。')
        message = '工作日志已标记为已阅'
        op_action = 'review'

    db.session.commit()
    _log_op('worklog', op_action, log.title)
    flash(message, 'success')
    return redirect(url_for('worklog_detail', log_id=log_id))



# ============================================================
#  电子公文收发模块
# ============================================================


def _gen_doc_no():
    """生成公文编号，格式：OA-年份-6位序号"""
    year = datetime.now().strftime('%Y')
    count = OfficialDoc.query.filter(OfficialDoc.doc_no.like(f'OA-{year}-%')).count()
    return f'OA-{year}-{count + 1:06d}'


OFFICIAL_DOC_ALLOWED_EXTENSIONS = {
    'pdf', 'doc', 'docx', 'xls', 'xlsx', 'txt',
    'zip', 'rar', '7z', 'jpg', 'jpeg', 'png'
}



def _allowed_official_doc_file(filename):
    return allowed_file(filename, OFFICIAL_DOC_ALLOWED_EXTENSIONS)


# ---------- 公文工作台（发件箱 + 草稿） ----------

@app.route('/official_doc/outbox')
@login_required
def doc_outbox():
    """我的发件箱（我起草的公文）"""
    _track_usage('official_doc', 'view_outbox')
    status_filter = request.args.get('status', '').strip()
    keyword = request.args.get('keyword', '').strip()

    q = OfficialDoc.query.filter_by(sender_id=current_user.id)
    if status_filter:
        q = q.filter(OfficialDoc.status == status_filter)
    if keyword:
        q = q.filter(db.or_(OfficialDoc.title.like(f'%{keyword}%'), OfficialDoc.doc_no.like(f'%{keyword}%')))

    docs = q.order_by(OfficialDoc.updated_at.desc()).all()
    all_docs = OfficialDoc.query.filter_by(sender_id=current_user.id).all()
    stats = {
        'draft': len([d for d in all_docs if d.status == 'draft']),
        'pending': len([d for d in all_docs if d.status == 'pending_approve']),
        'approved': len([d for d in all_docs if d.status == 'approved']),
        'sent': len([d for d in all_docs if d.status == 'sent']),
        'archived': len([d for d in all_docs if d.status == 'archived']),
    }
    return render_template('official_doc/outbox.html', docs=docs, status_filter=status_filter, keyword=keyword, stats=stats)



# ---------- 收件箱 ----------
@app.route('/official_doc/inbox')
@login_required
def doc_inbox():
    """收件箱（发给我的公文）"""
    _track_usage('official_doc', 'view_inbox')
    handle_filter = request.args.get('handle', '').strip()
    keyword = request.args.get('keyword', '').strip()

    q = DocReadRecord.query.join(OfficialDoc, OfficialDoc.id == DocReadRecord.doc_id).filter(DocReadRecord.user_id == current_user.id)
    if handle_filter:
        q = q.filter(DocReadRecord.handle_status == handle_filter)
    if keyword:
        q = q.filter(db.or_(OfficialDoc.title.like(f'%{keyword}%'), OfficialDoc.doc_no.like(f'%{keyword}%')))

    records = q.order_by(DocReadRecord.created_at.desc()).all()
    doc_ids = [r.doc_id for r in records]
    docs_map = {d.id: d for d in OfficialDoc.query.filter(OfficialDoc.id.in_(doc_ids)).all()} if doc_ids else {}
    all_records = DocReadRecord.query.filter_by(user_id=current_user.id).all()
    stats = {
        'unread': len([r for r in all_records if r.handle_status == 'unread']),
        'read': len([r for r in all_records if r.handle_status == 'read']),
        'handled': len([r for r in all_records if r.handle_status == 'handled']),
        'returned': len([r for r in all_records if r.handle_status == 'returned']),
    }
    unread_count = stats['unread']
    return render_template('official_doc/inbox.html', records=records, docs_map=docs_map,
                           handle_filter=handle_filter, unread_count=unread_count, keyword=keyword, stats=stats)



# ---------- 起草公文 ----------
@app.route('/official_doc/compose', methods=['GET', 'POST'])
@login_required
def doc_compose():
    """起草/编辑公文"""
    doc_id = request.args.get('id', type=int)
    doc = OfficialDoc.query.get(doc_id) if doc_id else None
    if doc and doc.sender_id != current_user.id:
        flash('无权编辑该公文', 'danger')
        return redirect(url_for('doc_outbox'))
    if doc and doc.status not in ('draft', 'recalled'):
        flash('当前状态下不可编辑，请先撤回或等待流程结束', 'warning')
        return redirect(url_for('doc_detail', doc_id=doc.id))

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        receiver_ids = request.form.getlist('receiver_ids')
        receiver_depts = request.form.get('receiver_depts', '').strip()
        action = request.form.get('action', 'save')  # save / submit

        if not title:
            flash('公文标题不能为空', 'danger')
            users = User.query.filter(
                User.id != current_user.id,
                db.or_(User.is_active == True, User.is_active == None)
            ).order_by(User.dept_id, User.name).all()
            depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
            return render_template('official_doc/compose.html', doc=doc, users=users, depts=depts)

        if action == 'submit' and not receiver_ids and not receiver_depts:
            flash('提交审批前请至少选择一位收件人或填写收件部门', 'danger')
            users = User.query.filter(
                User.id != current_user.id,
                db.or_(User.is_active == True, User.is_active == None)
            ).order_by(User.dept_id, User.name).all()
            depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
            return render_template('official_doc/compose.html', doc=doc, users=users, depts=depts)

        attachments = doc.get_attachments() if doc else []
        files = request.files.getlist('attachments')
        for f in files:
            if f and f.filename:
                safe_name = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{f.filename}"
                save_dir = os.path.join(app.root_path, 'uploads', 'official_doc')
                os.makedirs(save_dir, exist_ok=True)
                f.save(os.path.join(save_dir, safe_name))
                attachments.append({'name': f.filename, 'path': safe_name})

        if doc:
            doc.title = title
            doc.doc_type = request.form.get('doc_type', '通知')
            doc.urgency = request.form.get('urgency', '普通')
            doc.secrecy = request.form.get('secrecy', '普通')
            doc.content = request.form.get('content', '')
            doc.sender_dept = request.form.get('sender_dept', current_user.department or '')
            doc.sign_dept = request.form.get('sign_dept', '').strip()
            doc.receiver_ids = json.dumps([int(i) for i in receiver_ids if i], ensure_ascii=False)
            doc.receiver_depts = receiver_depts
            doc.set_attachments(attachments)
        else:
            doc = OfficialDoc(
                doc_no=_gen_doc_no(),
                title=title,
                doc_type=request.form.get('doc_type', '通知'),
                urgency=request.form.get('urgency', '普通'),
                secrecy=request.form.get('secrecy', '普通'),
                content=request.form.get('content', ''),
                sender_id=current_user.id,
                sender_dept=request.form.get('sender_dept', current_user.department or ''),
                sign_dept=request.form.get('sign_dept', '').strip(),
                receiver_depts=receiver_depts,
            )
            doc.set_receiver_ids([int(i) for i in receiver_ids if i])
            doc.set_attachments(attachments)
            db.session.add(doc)

        if action == 'submit':
            doc.status = 'pending_approve'
            db.session.flush()
            flow = DocFlow(doc_id=doc.id, operator_id=current_user.id,
                           action='submit', opinion='提交审批')
            db.session.add(flow)
            _log_op('official_doc', 'submit', f'提交公文 #{doc.doc_no}')
            flash('公文已提交审批', 'success')
        else:
            doc.status = 'draft'
            _log_op('official_doc', 'save_draft', f'保存草稿 #{doc.doc_no if doc.doc_no else "新建"}')
            flash('草稿已保存', 'success')

        db.session.commit()
        return redirect(url_for('doc_outbox'))

    users = User.query.filter(
        User.id != current_user.id,
        db.or_(User.is_active == True, User.is_active == None)
    ).order_by(User.dept_id, User.name).all()
    depts = Department.query.filter_by(is_active=True).order_by(Department.sort_order).all()
    return render_template('official_doc/compose.html', doc=doc, users=users, depts=depts)



# ---------- 公文详情 ----------
@app.route('/official_doc/detail/<int:doc_id>')
@login_required
def doc_detail(doc_id):
    """公文详情页（含流转记录）"""
    doc = OfficialDoc.query.get_or_404(doc_id)
    receiver_ids = doc.get_receiver_ids()

    if (doc.sender_id != current_user.id and
            current_user.id not in receiver_ids and
            current_user.role not in ('admin', 'manager')):
        flash('无权查看该公文', 'danger')
        return redirect(url_for('doc_inbox'))

    if current_user.id in receiver_ids:
        record = DocReadRecord.query.filter_by(doc_id=doc_id, user_id=current_user.id).first()
        if record and record.handle_status == 'unread':
            record.handle_status = 'read'
            record.read_at = datetime.now()
            db.session.commit()

    flows = doc.flows.order_by(DocFlow.created_at.asc()).all()
    receivers = User.query.filter(User.id.in_(receiver_ids)).all() if receiver_ids else []
    read_records = doc.read_records.all()
    read_map = {r.user_id: r for r in read_records}
    linked_supervisions = SupervisionTask.query.filter_by(source_doc_id=doc.id).order_by(SupervisionTask.created_at.desc()).all()
    linked_meetings = Meeting.query.filter_by(related_doc_id=doc.id).order_by(Meeting.start_time.desc(), Meeting.created_at.desc()).all()
    return render_template('official_doc/detail.html', doc=doc, flows=flows,
                           receivers=receivers, read_map=read_map,
                           linked_supervisions=linked_supervisions, linked_meetings=linked_meetings)



# ---------- 审批公文（管理员/经理） ----------
@app.route('/official_doc/approve/<int:doc_id>', methods=['POST'])
@login_required
def doc_approve(doc_id):
    """审批公文"""
    if current_user.role not in ('admin', 'manager'):
        return jsonify({'success': False, 'msg': '无审批权限'}), 403
    doc = OfficialDoc.query.get_or_404(doc_id)
    if doc.status != 'pending_approve':
        return jsonify({'success': False, 'msg': '公文状态不允许审批'}), 400

    action = request.form.get('action')  # approve / reject
    opinion = request.form.get('opinion', '')

    if action == 'approve':
        doc.status = 'approved'
        flow_action = 'approve'
        msg = '审批通过'
    else:
        doc.status = 'draft'
        flow_action = 'reject'
        msg = '已退回起草人'

    flow = DocFlow(doc_id=doc.id, operator_id=current_user.id,
                   action=flow_action, opinion=opinion or msg)
    db.session.add(flow)
    _log_op('official_doc', flow_action, f'审批公文 #{doc.doc_no}：{msg}')
    db.session.commit()
    flash(msg, 'success')
    return redirect(url_for('doc_detail', doc_id=doc_id))


# ---------- 发送公文 ----------
@app.route('/official_doc/send/<int:doc_id>', methods=['POST'])
@login_required
def doc_send(doc_id):
    """发送公文（审批通过后）"""
    doc = OfficialDoc.query.get_or_404(doc_id)
    if doc.sender_id != current_user.id:
        flash('无权发送该公文', 'danger')
        return redirect(url_for('doc_outbox'))
    if doc.status != 'approved':
        flash('公文需审批通过后方可正式发送', 'warning')
        return redirect(url_for('doc_detail', doc_id=doc_id))


    receiver_ids = doc.get_receiver_ids()
    
    # 如果有按部门发送，自动查找部门下的所有活跃用户补充进收件人
    if doc.receiver_depts:
        dept_names = [d.strip() for d in doc.receiver_depts.split(',') if d.strip()]
        if dept_names:
            dept_users = User.query.filter(
                User.department.in_(dept_names),
                db.or_(User.is_active == True, User.is_active == None),
                User.id != current_user.id
            ).all()
            dept_user_ids = [u.id for u in dept_users]
            # 合并去重
            all_ids = list(set(receiver_ids + dept_user_ids))
            if all_ids != receiver_ids:
                doc.set_receiver_ids(all_ids)
                receiver_ids = all_ids
    
    if not receiver_ids:
        flash('请先设置收件人或接收部门', 'warning')
        return redirect(url_for('doc_compose', id=doc_id))

    doc.status = 'sent'
    doc.sent_at = datetime.now()

    # 为每个收件人创建阅读记录
    for uid in receiver_ids:
        exists = DocReadRecord.query.filter_by(doc_id=doc_id, user_id=uid).first()
        if not exists:
            db.session.add(DocReadRecord(doc_id=doc_id, user_id=uid, handle_status='unread'))

    flow = DocFlow(doc_id=doc.id, operator_id=current_user.id,
                   action='send', opinion=f'发送给 {len(receiver_ids)} 位接收人')
    db.session.add(flow)
    _log_op('official_doc', 'send', f'发送公文 #{doc.doc_no}')
    db.session.commit()
    flash('公文已发送', 'success')
    return redirect(url_for('doc_outbox'))


# ---------- 办理公文（收件人操作） ----------
@app.route('/official_doc/handle/<int:doc_id>', methods=['POST'])
@login_required
def doc_handle(doc_id):
    """收件人办理/退回公文"""
    doc = OfficialDoc.query.get_or_404(doc_id)
    record = DocReadRecord.query.filter_by(doc_id=doc_id, user_id=current_user.id).first()
    if not record:
        flash('您不是该公文的接收人', 'danger')
        return redirect(url_for('doc_inbox'))

    action = request.form.get('action')  # handle / return
    opinion = request.form.get('opinion', '')

    if action == 'handle':
        record.handle_status = 'handled'
        record.handle_opinion = opinion
        record.handled_at = datetime.now()
        flow_action = 'handle'
        msg = '已办理'
    else:
        record.handle_status = 'returned'
        record.handle_opinion = opinion
        record.handled_at = datetime.now()
        flow_action = 'return'
        msg = '已退回'

    flow = DocFlow(doc_id=doc.id, operator_id=current_user.id,
                   action=flow_action, opinion=opinion or msg)
    db.session.add(flow)
    db.session.commit()
    flash(msg, 'success')
    return redirect(url_for('doc_detail', doc_id=doc_id))


# ---------- 撤回公文 ----------
@app.route('/official_doc/recall/<int:doc_id>', methods=['POST'])
@login_required
def doc_recall(doc_id):
    """撤回已发送的公文"""
    doc = OfficialDoc.query.get_or_404(doc_id)
    if doc.sender_id != current_user.id:
        flash('只有起草人可撤回', 'danger')
        return redirect(url_for('doc_outbox'))
    if doc.status not in ('sent', 'pending_approve'):
        flash('当前状态不允许撤回', 'warning')
        return redirect(url_for('doc_outbox'))

    doc.status = 'recalled'
    flow = DocFlow(doc_id=doc.id, operator_id=current_user.id,
                   action='recall', opinion='起草人撤回')
    db.session.add(flow)
    _log_op('official_doc', 'recall', f'撤回公文 #{doc.doc_no}')
    db.session.commit()
    flash('公文已撤回', 'success')
    return redirect(url_for('doc_outbox'))


# ---------- 归档公文 ----------
@app.route('/official_doc/archive/<int:doc_id>', methods=['POST'])
@login_required
def doc_archive(doc_id):
    """归档公文（管理员/发件人）"""
    doc = OfficialDoc.query.get_or_404(doc_id)
    if doc.sender_id != current_user.id and current_user.role not in ('admin', 'manager'):
        flash('无归档权限', 'danger')
        return redirect(url_for('doc_outbox'))

    doc.status = 'archived'
    doc.archived_at = datetime.now()
    flow = DocFlow(doc_id=doc.id, operator_id=current_user.id,
                   action='archive', opinion='已归档')
    db.session.add(flow)
    _log_op('official_doc', 'archive', f'归档公文 #{doc.doc_no}')
    db.session.commit()
    flash('公文已归档', 'success')
    return redirect(url_for('doc_outbox'))


# ---------- 删除草稿 ----------
@app.route('/official_doc/delete/<int:doc_id>', methods=['POST'])
@login_required
def doc_delete(doc_id):
    """删除草稿"""
    doc = OfficialDoc.query.get_or_404(doc_id)
    if doc.sender_id != current_user.id:
        flash('无权删除', 'danger')
        return redirect(url_for('doc_outbox'))
    if doc.status not in ('draft', 'recalled'):
        flash('只能删除草稿或已撤回的公文', 'warning')
        return redirect(url_for('doc_outbox'))
    db.session.delete(doc)
    db.session.commit()
    flash('已删除', 'success')
    return redirect(url_for('doc_outbox'))


# ---------- 公文归档列表 ----------
@app.route('/official_doc/archive_list')
@login_required
def doc_archive_list():
    """已归档公文列表"""
    if current_user.role not in ('admin', 'manager'):
        flash('无权查看归档列表', 'danger')
        return redirect(url_for('doc_inbox'))
    _track_usage('official_doc', 'view_archive')
    keyword = request.args.get('keyword', '')
    doc_type = request.args.get('doc_type', '')
    q = OfficialDoc.query.filter_by(status='archived')
    if keyword:
        q = q.filter(OfficialDoc.title.like(f'%{keyword}%'))
    if doc_type:
        q = q.filter(OfficialDoc.doc_type == doc_type)
    docs = q.order_by(OfficialDoc.archived_at.desc()).all()
    return render_template('official_doc/archive_list.html', docs=docs, keyword=keyword, doc_type=doc_type)


# ---------- 待审批列表（管理员/经理专用） ----------
@app.route('/official_doc/pending')
@login_required
def doc_pending():
    """待审批公文列表"""
    if current_user.role not in ('admin', 'manager'):
        flash('无审批权限', 'danger')
        return redirect(url_for('doc_inbox'))
    _track_usage('official_doc', 'view_pending')
    docs = OfficialDoc.query.filter_by(status='pending_approve').order_by(OfficialDoc.created_at.asc()).all()
    return render_template('official_doc/pending.html', docs=docs)


# ---------- 公文下载附件 ----------
@app.route('/official_doc/attachment/<path:filename>')
@login_required
def doc_attachment(filename):
    """下载公文附件"""
    att_dir = os.path.join(app.root_path, 'uploads', 'official_doc')
    return send_from_directory(att_dir, filename, as_attachment=True)


# ============================================================
#  电子公文模块 END
# ============================================================

# ============================================================
#  网站爬虫模块
# ============================================================
@app.route('/crawler')
@login_required
@csrf.exempt
def crawler_index():
    """爬虫管理首页"""
    from models import CrawlerTask, CrawlerPage
    tasks = CrawlerTask.query.order_by(CrawlerTask.updated_at.desc()).all()
    # 统计每个任务的页面数
    task_stats = {}
    for t in tasks:
        cnt = CrawlerPage.query.filter_by(task_id=t.id, status='completed').count()
        task_stats[t.id] = cnt
    return render_template('crawler/index.html', tasks=tasks, task_stats=task_stats)


@app.route('/crawler/task/new', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def crawler_task_new():
    """新建爬虫任务"""
    from models import CrawlerTask
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        base_url = request.form.get('base_url', '').strip()
        max_threads = int(request.form.get('max_threads', 5))
        attachment_handling = request.form.get('attachment_handling', 'skip')
        schedule_enabled = request.form.get('schedule_enabled') == '1'
        schedule_interval = int(request.form.get('schedule_interval', 24))
        schedule_time = request.form.get('schedule_time', '00:00')
        try:
            max_depth = int(request.form.get('max_depth', 3))
        except (TypeError, ValueError):
            max_depth = 3
        max_depth = max(0, min(max_depth, 10))

        if not name or not base_url:
            flash('任务名称和网址不能为空', 'warning')
            return redirect(url_for('crawler_index'))

        task = CrawlerTask(
            name=name,
            base_url=base_url,
            max_threads=max_threads,
            max_depth=max_depth,
            attachment_handling=attachment_handling,
            schedule_enabled=schedule_enabled,
            schedule_interval=schedule_interval,
            schedule_time=schedule_time,
            status='idle',
        )

        db.session.add(task)
        db.session.commit()
        flash(f'任务「{name}」创建成功', 'success')
        return redirect(url_for('crawler_index'))

    return render_template('crawler/task_form.html', task=None)


@app.route('/crawler/task/edit/<int:tid>', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def crawler_task_edit(tid):
    """编辑爬虫任务"""
    from models import CrawlerTask
    task = CrawlerTask.query.get_or_404(tid)

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        base_url = request.form.get('base_url', '').strip()
        max_threads = int(request.form.get('max_threads', 5))
        attachment_handling = request.form.get('attachment_handling', 'skip')
        schedule_enabled = request.form.get('schedule_enabled') == '1'
        schedule_interval = int(request.form.get('schedule_interval', 24))
        schedule_time = request.form.get('schedule_time', '00:00')
        try:
            max_depth = int(request.form.get('max_depth', 3))
        except (TypeError, ValueError):
            max_depth = 3
        max_depth = max(0, min(max_depth, 10))

        if not name or not base_url:
            flash('任务名称和网址不能为空', 'warning')
            return redirect(url_for('crawler_index'))

        task.name = name
        task.base_url = base_url
        task.max_threads = max_threads
        task.max_depth = max_depth
        task.attachment_handling = attachment_handling
        task.schedule_enabled = schedule_enabled
        task.schedule_interval = schedule_interval
        task.schedule_time = schedule_time
        db.session.commit()

        flash(f'任务「{name}」修改成功', 'success')
        return redirect(url_for('crawler_index'))

    return render_template('crawler/task_form.html', task=task)


@app.route('/crawler/task/<int:tid>/action', methods=['POST'])
@login_required
@csrf.exempt
def crawler_task_action(tid):
    """控制爬虫任务：start/pause/resume/stop"""
    from models import CrawlerTask
    from crawler_core import get_crawler, stop_crawler
    import json

    task = CrawlerTask.query.get_or_404(tid)
    action = request.form.get('action', '')

    if action == 'start':
        task.status = 'running'
        db.session.commit()
        crawler = get_crawler(task.id)
        if crawler:
            crawler.start()
        flash('爬虫已启动', 'success')

    elif action == 'pause':
        task.status = 'paused'
        db.session.commit()
        crawler = get_crawler(tid)
        if crawler:
            crawler.pause()
        flash('爬虫已暂停', 'info')

    elif action == 'resume':
        task.status = 'running'
        db.session.commit()
        crawler = get_crawler(tid)
        if crawler:
            crawler.resume()
        flash('爬虫已恢复', 'success')

    elif action == 'stop':
        task.status = 'stopped'
        db.session.commit()
        stop_crawler(tid)
        flash('爬虫已停止', 'info')

    elif action == 'delete':
        stop_crawler(tid)
        from models import CrawlerPage
        CrawlerPage.query.filter_by(task_id=tid).delete()
        CrawlerTask.query.filter_by(id=tid).delete()
        db.session.commit()
        flash('任务已删除', 'info')
        return redirect(url_for('crawler_index'))

    elif action == 'status':
        crawler = get_crawler(tid)
        if crawler:
            return jsonify(crawler.get_status())
        return jsonify({'error': 'no crawler'})

    return redirect(url_for('crawler_index'))


@app.route('/crawler/pages/<int:tid>')
@login_required
@csrf.exempt
def crawler_pages(tid):
    """查看任务爬取的页面"""
    from models import CrawlerTask, CrawlerPage
    task = CrawlerTask.query.get_or_404(tid)
    pages = CrawlerPage.query.filter_by(task_id=tid).order_by(CrawlerPage.updated_at.desc()).limit(200).all()
    return render_template('crawler/pages.html', task=task, pages=pages)


@app.route('/crawler/page/<int:pid>')
@login_required
def crawler_page_detail(pid):
    """获取页面详情（JSON，供弹窗使用）"""
    from models import CrawlerPage
    page = CrawlerPage.query.get_or_404(pid)
    return jsonify({
        'title': page.title or '无标题',
        'url': page.url,
        'status': page.status,
        'text_content': page.text_content or '（无正文内容）',
        'content': page.content or '',
        'updated_at': page.updated_at.strftime('%Y-%m-%d %H:%M') if page.updated_at else '-',
        'error_msg': page.error_msg or '',
    })


@app.route('/crawler/search', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def crawler_search():
    """关键词搜索"""
    from models import CrawlerTask, CrawlerPage

    results = []
    keyword = ''
    task_id = None

    if request.method == 'POST':
        keyword = request.form.get('keyword', '').strip()
        task_id = request.form.get('task_id')

        if keyword:
            q = CrawlerPage.query.filter(
                CrawlerPage.text_content.isnot(None),
                CrawlerPage.text_content != '',
            )
            if task_id and task_id != 'all':
                q = q.filter_by(task_id=int(task_id))
            # 分词搜索
            kw = keyword.replace(' ', '%')
            q = q.filter(CrawlerPage.text_content.like(f'%{kw}%'))
            results = q.order_by(CrawlerPage.updated_at.desc()).limit(100).all()

    tasks = CrawlerTask.query.order_by(CrawlerTask.name).all()
    return render_template('crawler/search.html', results=results, keyword=keyword,
                           tasks=tasks, selected_task=int(task_id) if task_id else None)


@app.route('/crawler/advanced_search', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def crawler_advanced_search():
    """高级检索"""
    from models import CrawlerTask, CrawlerPage

    results = []
    keyword = ''
    task_id = None
    date_from = ''
    date_to = ''
    status = ''

    if request.method == 'POST':
        keyword = request.form.get('keyword', '').strip()
        task_id = request.form.get('task_id')
        date_from = request.form.get('date_from', '')
        date_to = request.form.get('date_to', '')
        status = request.form.get('status', '')

        q = CrawlerPage.query.filter(
            CrawlerPage.text_content.isnot(None),
            CrawlerPage.text_content != '',
        )
        if keyword:
            kw = keyword.replace(' ', '%')
            q = q.filter(CrawlerPage.text_content.like(f'%{kw}%'))
        if task_id and task_id != 'all':
            q = q.filter_by(task_id=int(task_id))
        if date_from:
            q = q.filter(CrawlerPage.updated_at >= date_from)
        if date_to:
            q = q.filter(CrawlerPage.updated_at <= date_to + ' 23:59:59')
        if status:
            q = q.filter_by(status=status)

        results = q.order_by(CrawlerPage.updated_at.desc()).limit(200).all()

    tasks = CrawlerTask.query.order_by(CrawlerTask.name).all()
    return render_template('crawler/advanced_search.html',
                           results=results, keyword=keyword, tasks=tasks,
                           date_from=date_from, date_to=date_to, status=status)


@app.route('/crawler/api/progress/<int:tid>')
@login_required
@csrf.exempt
def crawler_progress(tid):
    """实时进度 SSE"""
    from models import CrawlerTask, CrawlerPage
    from flask import Response
    import json

    task = CrawlerTask.query.get(tid)
    if not task:
        return '', 404

    def generate():
        import time
        from crawler_core import get_crawler

        while True:
            crawler = get_crawler(tid)
            if crawler:
                s = crawler.get_status()
                CrawlerPage.query.filter_by(task_id=tid, status='completed').update({'status': 'completed'})
                db.session.commit()
            else:
                s = {'queued': 0, 'completed': CrawlerPage.query.filter_by(task_id=tid, status='completed').count(), 'errors': 0}

            data = json.dumps(s)
            yield f"data: {data}\n\n"
            time.sleep(2)

            if task.status not in ('running', 'paused'):
                break

    return Response(generate(), mimetype='text/event-stream')


# ============================================================
#  栏目监测模块
# ============================================================
@app.route('/monitor')
@login_required
@csrf.exempt
def monitor_index():
    """栏目监测首页"""
    from models import UrlLibrary, MonitorResult
    libs = UrlLibrary.query.order_by(UrlLibrary.updated_at.desc()).all()

    # 汇总统计
    lib_stats = {}
    for lib in libs:
        latest = MonitorResult.query.filter_by(library_id=lib.id).order_by(MonitorResult.monitor_time.desc()).first()
        overdue = MonitorResult.query.filter_by(library_id=lib.id, is_overdue=True).count() if latest else 0
        expiring = MonitorResult.query.filter_by(library_id=lib.id, is_expiring=True).count() if latest else 0
        total = MonitorResult.query.filter_by(library_id=lib.id).count() if latest else 0
        lib_stats[lib.id] = {'overdue': overdue, 'expiring': expiring, 'total': total,
                              'last_monitor': latest.monitor_time if latest else None}
    return render_template('monitor/index.html', libs=libs, lib_stats=lib_stats)


@app.route('/monitor/library/new', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def monitor_library_new():
    """新建网址库"""
    from models import UrlLibrary, UrlItem
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        category = request.form.get('category', '').strip()

        if not name:
            flash('库名称不能为空', 'warning')
            return redirect(url_for('monitor_index'))

        lib = UrlLibrary(name=name, description=description, category=category)
        db.session.add(lib)
        db.session.commit()
        flash(f'网址库「{name}」创建成功', 'success')
        return redirect(url_for('monitor_library_edit', lib_id=lib.id))

    return render_template('monitor/library_form.html', lib=None)


@app.route('/monitor/library/edit/<int:lib_id>', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def monitor_library_edit(lib_id):
    """编辑网址库"""
    from models import UrlLibrary, UrlItem
    lib = UrlLibrary.query.get_or_404(lib_id)

    if request.method == 'POST':
        action = request.form.get('action', '')

        if action == 'update_lib':
            lib.name = request.form.get('name', '').strip()
            lib.description = request.form.get('description', '').strip()
            lib.category = request.form.get('category', '').strip()
            expiring_days_str = request.form.get('expiring_days', '').strip()
            if expiring_days_str:
                try:
                    lib.expiring_days = max(1, min(365, int(expiring_days_str)))
                except ValueError:
                    lib.expiring_days = 8
            else:
                lib.expiring_days = 8
            db.session.commit()
            flash('库信息已更新', 'success')

        elif action == 'import_excel':
            file = request.files.get('excel_file')
            if file:
                import pandas as pd
                import io
                try:
                    df = pd.read_excel(file)
                    # 识别列名
                    col_map = {}
                    for col in df.columns:
                        cl = col.lower()
                        if '序号' in col: col_map['serial_no'] = col
                        elif '栏目名称' in col: col_map['column_name'] = col
                        elif 'url' in cl or '网址' in col: col_map['url'] = col
                        elif '栏目分类' in col: col_map['column_category'] = col
                        elif '更新期限' in col: col_map['update_deadline'] = col
                        elif '期限' in col and '天' not in col: col_map['deadline_days'] = col
                        elif '网站名称' in col: col_map['website_name'] = col
                        elif '标识码' in col: col_map['website_code'] = col

                    imported = 0
                    for _, row in df.iterrows():
                        url = str(row.get(col_map.get('url', ''), '')).strip()
                        if not url or url == 'nan':
                            continue
                        col_name = str(row.get(col_map.get('column_name', ''), '')).strip()
                        deadline_str = str(row.get(col_map.get('update_deadline', ''), '')).strip()
                        # 解析期限天数
                        days = None
                        if deadline_str and deadline_str != 'nan':
                            if '天' in deadline_str:
                                days = int(re.search(r'(\d+)', deadline_str).group(1)) if re.search(r'(\d+)', deadline_str) else None
                            elif '周' in deadline_str:
                                days = int(re.search(r'(\d+)', deadline_str).group(1)) * 7 if re.search(r'(\d+)', deadline_str) else None
                            elif '月' in deadline_str:
                                days = int(re.search(r'(\d+)', deadline_str).group(1)) * 30 if re.search(r'(\d+)', deadline_str) else None

                        item = UrlItem(
                            library_id=lib.id,
                            serial_no=str(row.get(col_map.get('serial_no', ''), '')),
                            column_name=col_name,
                            url=url,
                            column_category=str(row.get(col_map.get('column_category', ''), '')).strip(),
                            update_deadline=deadline_str,
                            deadline_days=days,
                            website_name=str(row.get(col_map.get('website_name', ''), '')).strip(),
                            website_code=str(row.get(col_map.get('website_code', ''), '')).strip(),
                        )
                        db.session.add(item)
                        imported += 1

                    lib.item_count = UrlItem.query.filter_by(library_id=lib.id).count()
                    db.session.commit()
                    flash(f'成功导入 {imported} 条网址', 'success')
                except Exception as e:
                    flash(f'导入失败: {e}', 'danger')

        elif action == 'add_item':
            url = request.form.get('url', '').strip()
            if not url:
                flash('网址不能为空', 'warning')
            else:
                deadline_str = request.form.get('update_deadline', '').strip()
                days = None
                if deadline_str:
                    if '天' in deadline_str:
                        m = re.search(r'(\d+)', deadline_str)
                        days = int(m.group(1)) if m else None
                    elif '周' in deadline_str:
                        m = re.search(r'(\d+)', deadline_str)
                        days = int(m.group(1)) * 7 if m else None
                    elif '月' in deadline_str:
                        m = re.search(r'(\d+)', deadline_str)
                        days = int(m.group(1)) * 30 if m else None

                item = UrlItem(
                    library_id=lib.id,
                    serial_no=request.form.get('serial_no', '').strip(),
                    column_name=request.form.get('column_name', '').strip(),
                    url=url,
                    column_category=request.form.get('column_category', '').strip(),
                    update_deadline=deadline_str,
                    deadline_days=days,
                )
                db.session.add(item)
                lib.item_count = UrlItem.query.filter_by(library_id=lib.id).count()
                db.session.commit()
                flash('网址已添加', 'success')

        elif action == 'delete_item':
            item_id = request.form.get('item_id', type=int)
            if item_id:
                UrlItem.query.filter_by(id=item_id).delete()
                lib.item_count = UrlItem.query.filter_by(library_id=lib.id).count()
                db.session.commit()

        elif action == 'delete_all':
            UrlItem.query.filter_by(library_id=lib.id).delete()
            lib.item_count = 0
            db.session.commit()
            flash('已清空所有网址', 'info')

        return redirect(url_for('monitor_library_edit', lib_id=lib.id))

    items = UrlItem.query.filter_by(library_id=lib_id).order_by(UrlItem.sort_order, UrlItem.id).all()
    return render_template('monitor/library_edit.html', lib=lib, items=items)


@app.route('/monitor/library/delete/<int:lib_id>', methods=['POST'])
@login_required
@csrf.exempt
def monitor_library_delete(lib_id):
    """删除网址库"""
    from models import UrlLibrary, UrlItem, MonitorResult, MonitorLog
    UrlItem.query.filter_by(library_id=lib_id).delete()
    MonitorResult.query.filter_by(library_id=lib_id).delete()
    MonitorLog.query.filter_by(library_id=lib_id).delete()
    UrlLibrary.query.filter_by(id=lib_id).delete()
    db.session.commit()
    flash('网址库已删除', 'info')
    return redirect(url_for('monitor_index'))


@app.route('/monitor/run/<int:lib_id>', methods=['POST'])
@login_required
@csrf.exempt
def monitor_run(lib_id):
    """执行监测"""
    from models import UrlLibrary
    from monitor_core import get_monitor
    import threading

    try:
        lib = UrlLibrary.query.get_or_404(lib_id)
        expiring_days = lib.expiring_days or 8
        monitor = get_monitor(lib_id, expiring_days=expiring_days)
        count = monitor.load_items()

        if count == 0:
            flash('网址库为空，请先添加网址', 'warning')
            return redirect(url_for('monitor_library_edit', lib_id=lib_id))

        # 后台执行
        def run_bg():
            with app.app_context():
                monitor.run()

        t = threading.Thread(target=run_bg, daemon=True)
        t.start()

        flash(f'监测已启动，正在处理 {count} 个网址...', 'success')
        return redirect(url_for('monitor_results', lib_id=lib_id))
    except Exception as e:
        import traceback
        traceback.print_exc()
        flash(f'监测启动失败: {str(e)}', 'error')
        return redirect(url_for('monitor_index'))


@app.route('/monitor/results/<int:lib_id>')
@login_required
@csrf.exempt
def monitor_results(lib_id):
    """查看监测结果"""
    from models import UrlLibrary, MonitorResult
    lib = UrlLibrary.query.get_or_404(lib_id)
    results = MonitorResult.query.filter_by(library_id=lib_id).order_by(MonitorResult.is_overdue.desc(), MonitorResult.days_since_update.desc()).all()

    overdue = [r for r in results if r.is_overdue]
    expiring = [r for r in results if r.is_expiring]
    normal = [r for r in results if not r.is_overdue and not r.is_expiring]

    return render_template('monitor/results.html', lib=lib, results=results,
                           overdue=overdue, expiring=expiring, normal=normal)


@app.route('/monitor/logs/<int:lib_id>')
@login_required
@csrf.exempt
def monitor_logs(lib_id):
    """查看监测日志"""
    from models import UrlLibrary, MonitorLog
    lib = UrlLibrary.query.get_or_404(lib_id)
    logs = MonitorLog.query.filter_by(library_id=lib_id).order_by(MonitorLog.created_at.desc()).limit(200).all()
    return render_template('monitor/logs.html', lib=lib, logs=logs)


@app.route('/monitor/export/<int:lib_id>')
@login_required
@csrf.exempt
def monitor_export(lib_id):
    """导出监测结果"""
    from models import MonitorResult
    import pandas as pd
    from flask import make_response

    results = MonitorResult.query.filter_by(library_id=lib_id).order_by(MonitorResult.is_overdue.desc()).all()

    data = [{
        '序号': r.url_item_id,
        '栏目名称': r.column_name,
        '网址': r.url,
        '栏目分类': r.column_category,
        '更新期限': r.update_deadline,
        '期限天数': r.deadline_days,
        '最大日期': r.last_max_date,
        '距今天数': r.days_since_update,
        '是否逾期': '是' if r.is_overdue else '否',
        '即将逾期': '是' if r.is_expiring else '否',
        '状态': '成功' if r.status == 'completed' else '失败',
        '监测时间': r.monitor_time,
    } for r in results]

    df = pd.DataFrame(data)
    output = io.StringIO()
    df.to_csv(output, index=False, encoding='utf-8-sig')
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = f'attachment; filename=monitor_result_{lib_id}.csv'
    return response


# ============================================================
# 栏目监测定时任务管理
# ============================================================
@app.route('/monitor/schedule')
@login_required
def monitor_schedule_index():
    """定时任务管理首页"""
    from models import MonitorScheduledTask, UrlLibrary
    tasks = MonitorScheduledTask.query.order_by(MonitorScheduledTask.created_at.desc()).all()
    libraries = UrlLibrary.query.order_by(UrlLibrary.name).all()
    return render_template('monitor/schedule.html', tasks=tasks, libraries=libraries)


@app.route('/monitor/schedule/new', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def monitor_schedule_new():
    """新建定时任务"""
    from models import MonitorScheduledTask, UrlLibrary
    import json
    
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        library_id = request.form.get('library_id', type=int)
        cron_expression = request.form.get('cron_expression', '').strip()
        email_recipients = request.form.get('email_recipients', '').strip()
        is_active = request.form.get('is_active') == 'on'
        
        if not name:
            flash('任务名称不能为空', 'warning')
            return redirect(url_for('monitor_schedule_index'))
        
        if not library_id:
            flash('请选择监测网址库', 'warning')
            return redirect(url_for('monitor_schedule_index'))
        
        # 验证网址库存在
        library = UrlLibrary.query.get(library_id)
        if not library:
            flash('指定的网址库不存在', 'error')
            return redirect(url_for('monitor_schedule_index'))
        
        # 解析邮件收件人
        recipients_list = [e.strip() for e in email_recipients.split(',') if e.strip()]
        
        task = MonitorScheduledTask(
            name=name,
            library_id=library_id,
            cron_expression=cron_expression,
            email_recipients=json.dumps(recipients_list),
            is_active=is_active
        )
        db.session.add(task)
        db.session.commit()
        
        # 如果启用，添加到调度器
        if is_active:
            try:
                from monitor_scheduler import scheduler, execute_scheduled_monitor_task, parse_cron_expression
                job_id = f'monitor_task_{task.id}'
                if scheduler.get_job(job_id):
                    scheduler.remove_job(job_id)
                scheduler.add_job(
                    id=job_id,
                    func=execute_scheduled_monitor_task,
                    args=[task.id],
                    trigger='cron',
                    **parse_cron_expression(cron_expression)
                )
            except Exception as e:
                flash(f'任务已创建但调度器启动失败: {str(e)}', 'warning')
        
        flash(f'定时任务「{name}」创建成功', 'success')
        return redirect(url_for('monitor_schedule_index'))
    
    libraries = UrlLibrary.query.order_by(UrlLibrary.name).all()
    return render_template('monitor/schedule_form.html', task=None, libraries=libraries)


@app.route('/monitor/schedule/edit/<int:task_id>', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def monitor_schedule_edit(task_id):
    """编辑定时任务"""
    from models import MonitorScheduledTask, UrlLibrary
    import json
    
    task = MonitorScheduledTask.query.get_or_404(task_id)
    
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        library_id = request.form.get('library_id', type=int)
        cron_expression = request.form.get('cron_expression', '').strip()
        email_recipients = request.form.get('email_recipients', '').strip()
        is_active = request.form.get('is_active') == 'on'
        
        if not name:
            flash('任务名称不能为空', 'warning')
            return redirect(url_for('monitor_schedule_index'))
        
        task.name = name
        task.library_id = library_id
        task.cron_expression = cron_expression
        task.email_recipients = json.dumps([e.strip() for e in email_recipients.split(',') if e.strip()])
        task.is_active = is_active
        db.session.commit()
        
        # 更新调度器
        try:
            from monitor_scheduler import scheduler, execute_scheduled_monitor_task, parse_cron_expression
            job_id = f'monitor_task_{task.id}'
            if scheduler.get_job(job_id):
                scheduler.remove_job(job_id)
            if is_active:
                scheduler.add_job(
                    id=job_id,
                    func=execute_scheduled_monitor_task,
                    args=[task.id],
                    trigger='cron',
                    **parse_cron_expression(cron_expression)
                )
        except Exception as e:
            flash(f'任务已更新但调度器更新失败: {str(e)}', 'warning')
        
        flash(f'定时任务「{name}」更新成功', 'success')
        return redirect(url_for('monitor_schedule_index'))
    
    libraries = UrlLibrary.query.order_by(UrlLibrary.name).all()
    return render_template('monitor/schedule_form.html', task=task, libraries=libraries)


@app.route('/monitor/schedule/delete/<int:task_id>', methods=['POST'])
@login_required
@csrf.exempt
def monitor_schedule_delete(task_id):
    """删除定时任务"""
    from models import MonitorScheduledTask
    
    task = MonitorScheduledTask.query.get_or_404(task_id)
    
    # 从调度器移除
    try:
        from monitor_scheduler import scheduler
        job_id = f'monitor_task_{task.id}'
        if scheduler.get_job(job_id):
            scheduler.remove_job(job_id)
    except:
        pass
    
    db.session.delete(task)
    db.session.commit()
    
    flash('定时任务已删除', 'info')
    return redirect(url_for('monitor_schedule_index'))


@app.route('/monitor/schedule/toggle/<int:task_id>', methods=['POST'])
@login_required
@csrf.exempt
def monitor_schedule_toggle(task_id):
    """启用/禁用定时任务"""
    from models import MonitorScheduledTask
    
    task = MonitorScheduledTask.query.get_or_404(task_id)
    task.is_active = not task.is_active
    db.session.commit()
    
    # 更新调度器
    try:
        from monitor_scheduler import scheduler, execute_scheduled_monitor_task, parse_cron_expression
        job_id = f'monitor_task_{task.id}'
        if scheduler.get_job(job_id):
            scheduler.remove_job(job_id)
        if task.is_active:
            scheduler.add_job(
                id=job_id,
                func=execute_scheduled_monitor_task,
                args=[task.id],
                trigger='cron',
                **parse_cron_expression(task.cron_expression)
            )
        status = '启用' if task.is_active else '禁用'
        flash(f'定时任务已{status}', 'success')
    except Exception as e:
        flash(f'状态已更新但调度器更新失败: {str(e)}', 'warning')
    
    return redirect(url_for('monitor_schedule_index'))


@app.route('/monitor/schedule/logs')
@login_required
@csrf.exempt
def monitor_schedule_logs():
    """查看调度日志"""
    from models import MonitorSystemLog
    logs = MonitorSystemLog.query.order_by(MonitorSystemLog.created_at.desc()).limit(100).all()
    return render_template('monitor/schedule_logs.html', logs=logs)


# ==================== 系统配置管理 ====================

@app.route('/admin/system_config')
@login_required
def system_config():
    """系统配置页面"""
    if current_user.role != 'admin':
        flash('只有管理员可以管理系统配置')
        return redirect(url_for('index'))
    
    config = config_manager.get_all()
    
    # 将配置数据转换为模板期望的分类格式
    categories = {}
    
    # 配置项的定义
    config_definitions = {
        'knowledge_base': [
            {'key': 'embedding_model_path', 'type': 'string', 'description': '嵌入模型路径', 'module': 'AI'},
            {'key': 'use_local_model', 'type': 'boolean', 'description': '使用本地模型', 'module': 'AI'},
            {'key': 'max_file_size_mb', 'type': 'integer', 'description': '最大文件大小(MB)', 'module': '上传'},
            {'key': 'batch_size', 'type': 'integer', 'description': '批量处理大小', 'module': '处理'},
            {'key': 'auto_extract_keywords', 'type': 'boolean', 'description': '自动抽取关键词', 'module': 'AI'},
            {'key': 'auto_generate_summary', 'type': 'boolean', 'description': '自动生成摘要', 'module': 'AI'},
            {'key': 'auto_tag', 'type': 'boolean', 'description': '自动标签', 'module': 'AI'}
        ],
        'ocr': [
            {'key': 'enabled', 'type': 'boolean', 'description': 'OCR功能启用', 'module': 'OCR'},
            {'key': 'tesseract_cmd', 'type': 'string', 'description': 'Tesseract路径', 'module': 'OCR'},
            {'key': 'tessdata_dir', 'type': 'string', 'description': '字库目录', 'module': 'OCR'},
            {'key': 'languages', 'type': 'json', 'description': '支持语言', 'module': 'OCR'},
            {'key': 'dpi', 'type': 'integer', 'description': '图片DPI', 'module': 'OCR'},
            {'key': 'psm_mode', 'type': 'integer', 'description': 'PSM模式', 'module': 'OCR'}
        ],
        'ai': [
            {'key': 'default_model', 'type': 'string', 'description': '默认AI模型', 'module': 'AI'},
            {'key': 'temperature', 'type': 'float', 'description': '温度参数', 'module': 'AI'},
            {'key': 'max_tokens', 'type': 'integer', 'description': '最大token数', 'module': 'AI'}
        ],
        'monitoring': [
            {'key': 'auto_run', 'type': 'boolean', 'description': '自动运行监控', 'module': '监控'},
            {'key': 'check_interval_hours', 'type': 'integer', 'description': '检查间隔(小时)', 'module': '监控'},
            {'key': 'notify_on_overdue', 'type': 'boolean', 'description': '超时通知', 'module': '监控'}
        ],
        'document': [
            {'key': 'max_extracted_length', 'type': 'integer', 'description': '最大提取长度(-1无限制)', 'module': '文档'},
            {'key': 'max_preview_length', 'type': 'integer', 'description': '预览页面最大长度', 'module': '文档'},
            {'key': 'max_ai_sample_length', 'type': 'integer', 'description': 'AI分析最大长度', 'module': '文档'},
            {'key': 'max_file_preview_length', 'type': 'integer', 'description': '文件预览最大长度', 'module': '文档'},
            {'key': 'max_upload_size_mb', 'type': 'integer', 'description': '最大上传大小(MB)', 'module': '文档'}
        ]
    }
    
    # 构建categories数据
    for category_name, config_items in config_definitions.items():
        if category_name in config:
            categories[category_name] = []
            for item_def in config_items:
                value = config[category_name].get(item_def['key'])
                if value is not None:
                    item = item_def.copy()
                    item['key'] = f"{category_name}.{item_def['key']}"  # 使用点分隔的完整路径
                    item['value'] = value
                    item['updated_at'] = datetime.now().isoformat()
                    categories[category_name].append(item)
    
    # 检查各项配置状态
    status = {
        'embedding_model': {
            'exists': os.path.exists(config_manager.get_embedding_model_path()),
            'path': config_manager.get_embedding_model_path()
        },
        'tesseract': {
            'exists': os.path.exists(config_manager.get_tesseract_cmd()),
            'path': config_manager.get_tesseract_cmd()
        },
        'tessdata': {
            'exists': os.path.exists(config_manager.get_tessdata_dir()) if config_manager.get_tessdata_dir() else False,
            'path': config_manager.get_tessdata_dir()
        }
    }
    
    # 检查中文字库
    if status['tessdata']['exists']:
        chi_sim_path = os.path.join(config_manager.get_tessdata_dir(), 'chi_sim.traineddata')
        status['chi_sim'] = {
            'exists': os.path.exists(chi_sim_path),
            'path': chi_sim_path
        }
    
    return render_template('admin/system_config.html', categories=categories, status=status)


@app.route('/admin/system_config/save', methods=['POST'])
@login_required
@csrf.exempt
def system_config_save():
    """保存系统配置"""
    if current_user.role != 'admin':
        return jsonify({'error': '权限不足'}), 403
    
    try:
        data = request.get_json()
        section = data.get('section')
        values = data.get('values', {})
        
        if section and values:
            config_manager.update_section(section, values)
            return jsonify({'success': True, 'message': '配置已保存'})
        
        return jsonify({'error': '参数错误'}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/admin/system_config/test', methods=['POST'])
@login_required
@csrf.exempt
def system_config_test():
    """测试配置"""
    if current_user.role != 'admin':
        return jsonify({'error': '权限不足'}), 403
    
    test_type = request.json.get('type')
    
    if test_type == 'embedding':
        try:
            from sentence_transformers import SentenceTransformer
            model_path = config_manager.get_embedding_model_path()
            if os.path.exists(model_path):
                model = SentenceTransformer(model_path, local_files_only=True)
                # 测试编码
                test_vector = model.encode("测试文本")
                return jsonify({
                    'success': True,
                    'message': f'模型加载成功，向量维度: {len(test_vector)}'
                })
            else:
                return jsonify({'error': '模型路径不存在'}), 400
        except Exception as e:
            return jsonify({'error': f'模型测试失败: {str(e)}'}), 500
    
    elif test_type == 'ocr':
        try:
            import pytesseract
            tesseract_cmd = config_manager.get_tesseract_cmd()
            if os.path.exists(tesseract_cmd):
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                version = pytesseract.get_tesseract_version()
                return jsonify({
                    'success': True,
                    'message': f'Tesseract 版本: {version}'
                })
            else:
                return jsonify({'error': 'Tesseract 路径不存在'}), 400
        except Exception as e:
            return jsonify({'error': f'OCR测试失败: {str(e)}'}), 500
    
    return jsonify({'error': '未知的测试类型'}), 400




# ============================================================
#  知识库首页 - 重定向到个人知识库
# ============================================================

@app.route('/knowledge')
@login_required
def knowledge_index():
    """知识库首页 - 重定向到个人知识库"""
    return redirect(url_for('personal_knowledge_base'))



# ============================================================
#  档案管理模块 - 兼容性路由（重定向到蓝图路由，保留 endpoint 名称供 url_for 使用）
# ============================================================

@app.route('/archive/fonds')
@login_required
def archive_fonds_list():
    """全宗管理页面 - 转发到蓝图"""
    return redirect(url_for('archive.fonds_list'))


@app.route('/archive/files')
@login_required
def archive_file_list():
    """档案列表页面 - 转发到蓝图"""
    qs = request.query_string.decode('utf-8')
    target = url_for('archive.file_list')
    return redirect(f"{target}?{qs}" if qs else target)


@app.route('/archive/search')
@login_required
def archive_search():
    """档案检索页面 - 转发到蓝图"""
    qs = request.query_string.decode('utf-8')
    target = url_for('archive.search')
    return redirect(f"{target}?{qs}" if qs else target)


@app.route('/archive/file/<int:file_id>')
@login_required
def archive_file_detail(file_id):
    """档案详情页面 - 转发到蓝图"""
    return redirect(url_for('archive.file_detail', file_id=file_id))


@app.route('/archive/tasks')
@login_required
def archive_task_list():
    """任务管理页面 - 转发到蓝图"""
    return redirect(url_for('archive.task_list'))


@app.route('/archive/my_borrows')
@login_required
def archive_my_borrows():
    """我的借阅页面 - 转发到蓝图"""
    qs = request.query_string.decode('utf-8')
    target = url_for('archive.my_borrows')
    return redirect(f"{target}?{qs}" if qs else target)


@app.route('/archive/statistics')
@login_required
def archive_statistics():
    """统计报表页面 - 转发到蓝图"""
    return redirect(url_for('archive.statistics'))





# ============================================================
#  档案管理模块 END
# ============================================================


if __name__ == '__main__':
    # DEBUG 模式：从环境变量读取，默认关闭（公网访问时不暴露调试信息）
    debug_mode = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    # HOST：0.0.0.0 监听所有网络接口，允许公网访问
    host = os.environ.get('FLASK_HOST', '0.0.0.0')
    # PORT：默认 5000，可通过环境变量覆盖
    port = int(os.environ.get('FLASK_PORT', 5000))
    app.run(host=host, port=port, debug=debug_mode)