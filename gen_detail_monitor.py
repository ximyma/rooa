import sys
sys.stdout.reconfigure(encoding="utf-8")
import os
os.chdir(r"C:\Users\Administrator\Desktop\ooa")
from datetime import datetime
from app import app, db
from models import MonitorResult

print("Generating detailed monitoring report...")
date_str = datetime.now().strftime("%Y%m%d")
output_dir = r"C:\Users\Administrator\Desktop\ooa\briefing_output"

with app.app_context():
    # 获取所有逾期栏目
    overdue = MonitorResult.query.filter_by(is_overdue=True).order_by(MonitorResult.days_since_update.desc()).all()
    expiring = MonitorResult.query.filter_by(is_expiring=True, is_overdue=False).order_by(MonitorResult.days_since_update.desc()).all()
    normal = MonitorResult.query.filter_by(is_overdue=False, is_expiring=False).count()
    
    print(f"Overdue: {len(overdue)}, Expiring: {len(expiring)}, Normal: {normal}")
    
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    doc.add_heading(f"石城县政府网站栏目监测报告", 0)
    doc.add_paragraph(f"监测时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"总监测栏目: {MonitorResult.query.count()}个")
    
    # 汇总表
    doc.add_heading("一、监测汇总", level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"🔴 逾期栏目: {len(overdue)}个\n").bold = True
    summary.add_run(f"🟡 预警栏目: {len(expiring)}个\n").bold = True
    summary.add_run(f"🟢 正常栏目: {normal}个\n").bold = True
    
    # 逾期栏目详情
    if overdue:
        doc.add_heading("二、逾期栏目详情（需要立即更新）", level=1)
        for i, r in enumerate(overdue, 1):
            p = doc.add_paragraph()
            p.add_run(f"[逾期] {r.column_name}").bold = True
            doc.add_paragraph(f"    栏目路径: {r.column_category or 'N/A'}")
            doc.add_paragraph(f"    URL: {r.url}")
            doc.add_paragraph(f"    最后更新: {r.last_max_date or '未知'}")
            doc.add_paragraph(f"    逾期天数: {r.days_since_update}天")
            doc.add_paragraph(f"    更新期限: {r.update_deadline}")
            doc.add_paragraph("")
    
    # 预警栏目
    if expiring:
        doc.add_heading("三、预警栏目（即将到期）", level=1)
        for i, r in enumerate(expiring[:30], 1):
            p = doc.add_paragraph()
            p.add_run(f"[预警] {r.column_name}").bold = True
            doc.add_paragraph(f"    栏目路径: {r.column_category or 'N/A'}")
            doc.add_paragraph(f"    URL: {r.url}")
            doc.add_paragraph(f"    最后更新: {r.last_max_date or '未知'}")
            doc.add_paragraph(f"    剩余天数: {r.days_since_update}天")
            doc.add_paragraph("")
    
    out = f"{output_dir}\\石城县政府网站监测详情_{date_str}.docx"
    doc.save(out)
    print(f"Saved: {out}")
