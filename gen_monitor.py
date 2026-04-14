import sys
sys.stdout.reconfigure(encoding="utf-8")
import os
os.chdir(r"C:\Users\Administrator\Desktop\ooa")
from datetime import datetime
from app import app, db
from models import MonitorResult

print("Generating website monitoring report...")
date_str = datetime.now().strftime("%Y%m%d")
output_dir = r"C:\Users\Administrator\Desktop\ooa\briefing_output"

with app.app_context():
    results = MonitorResult.query.filter(
        MonitorResult.monitor_time >= datetime.now().replace(hour=0, minute=0, second=0)
    ).order_by(MonitorResult.monitor_time.desc()).limit(200).all()
    
    print(f"Today results: {len(results)}")
    
    # 统计
    overdue = [r for r in results if r.is_overdue]
    expiring = [r for r in results if r.is_expiring and not r.is_overdue]
    normal = [r for r in results if not r.is_overdue and not r.is_expiring]
    
    articles = []
    
    # 汇总
    articles.append({
        "title": f"监测汇总 - 共{len(results)}个栏目",
        "content": f"逾期: {len(overdue)}个\n预警: {len(expiring)}个\n正常: {len(normal)}个",
        "url": "",
        "source_name": "系统",
        "keyword": "汇总"
    })
    
    # 逾期列表
    for r in overdue[:20]:
        articles.append({
            "title": f"[逾期] {r.column_name}",
            "content": f"URL: {r.url}\n分类: {r.column_category or 'N/A'}\n逾期天数: {r.days_since_update}天\n截止日期: {r.update_deadline}",
            "url": r.url,
            "source_name": "逾期预警",
            "keyword": "逾期"
        })
    
    # 预警列表
    for r in expiring[:10]:
        articles.append({
            "title": f"[预警] {r.column_name}",
            "content": f"URL: {r.url}\n分类: {r.column_category or 'N/A'}\n距截止: {r.days_since_update}天",
            "url": r.url,
            "source_name": "预警",
            "keyword": "预警"
        })
    
    # 生成docx
    from docx import Document
    doc = Document()
    doc.add_heading(f"政府网站监测报告 - {date_str}", 0)
    
    for a in articles:
        doc.add_heading(a["title"], level=1)
        doc.add_paragraph(a["content"])
        doc.add_paragraph("")
    
    out = f"{output_dir}\\政府网站监测报告_{date_str}.docx"
    doc.save(out)
    print(f"Saved: {out}")
    print(f"Total: {len(articles)}")
