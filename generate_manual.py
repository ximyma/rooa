import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level, number):
    """添加带编号的标题"""
    heading = doc.add_heading(f"{'.'.join(map(str, number))} {text}", level)

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 标题页
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('OOA智能办公系统', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
subtitle = doc.add_paragraph('用户手册')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)
subtitle.runs[0].font.bold = True
doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'版本: 2.0\n')
info.add_run(f'编制日期: {datetime.now().strftime("%Y年%m月%d日")}\n')
info.add_run(f'系统状态: 运行中\n')
doc.add_page_break()

# 目录
doc.add_heading('目录', 1)
toc_items = [
    "第一章 系统概述",
    "第二章 系统功能模块",
    "  2.1 公文管理系统",
    "  2.2 智能简报系统",
    "  2.3 网站监测系统",
    "  2.4 档案管理系统",
    "  2.5 知识管理系统",
    "  2.6 智能问答系统",
    "  2.7 会议管理系统",
    "  2.8 工作日志系统",
    "  2.9 绩效考核系统",
    "  2.10 专报管理系统",
    "  2.11 督办管理系统",
    "  2.12 智能办公工具",
    "  2.13 系统管理",
    "第三章 优化升级建议",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
doc.add_page_break()

# 第一章 系统概述
add_heading_with_number(doc, "系统概述", 1, [1])
doc.add_paragraph("OOA智能办公系统是一套面向政府机关、企事业单位的综合办公管理平台，集成公文流转、智能简报、网站监测、档案管理、知识库等核心功能，支持AI智能辅助决策。")

# 系统架构表
doc.add_heading("1.1 系统架构", 2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
headers = ["组件", "描述"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, "4472C4")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    cell.paragraphs[0].runs[0].font.bold = True

data = [
    ["后端框架", "Flask + SQLAlchemy + SQLite"],
    ["前端框架", "Bootstrap + jQuery + Layui"],
    ["AI引擎", "支持多种大模型对接"],
    ["数据规模", "48个数据模型，5651+监测记录"],
]
for i, row_data in enumerate(data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

# 功能统计
doc.add_heading("1.2 功能统计", 2)
stats = doc.add_paragraph()
stats.add_run("• 数据模型: 48个\n")
stats.add_run("• 前端页面: 130+个\n")
stats.add_run("• 简报数据源: 61个\n")
stats.add_run("• 网站监测栏目: 5651个\n")
stats.add_run("• 历史简报: 84+条\n")

# 第二章 系统功能模块
add_heading_with_number(doc, "系统功能模块", 1, [2])

# 2.1 公文管理
add_heading_with_number(doc, "公文管理系统", 2, [2, 1])
doc.add_paragraph("公文管理是系统的核心模块，支持发文、收文、签批、归档全流程电子化管理。")
features1 = doc.add_paragraph()
features1.add_run("【主要功能】\n").bold = True
features1.add_run("• 发文管理：起草、审核、签发、打印\n")
features1.add_run("• 收文管理：登记、拟办、批示、传阅\n")
features1.add_run("• 公文模板：支持自定义格式模板\n")
features1.add_run("• 公文检索：多条件快速查询\n")
features1.add_run("• 归档管理：自动归档与档案对接\n")
doc.add_paragraph()

# 2.2 智能简报
add_heading_with_number(doc, "智能简报系统", 2, [2, 2])
doc.add_paragraph("智能简报系统自动采集、整合、生成各类专题简报，支持定时推送。")
features2 = doc.add_paragraph()
features2.add_run("【主要功能】\n").bold = True
features2.add_run("• 多源采集：支持61个新闻/政策数据源\n")
features2.add_run("• 智能分类：自动归类到新闻/金融/AI/政策等类别\n")
features2.add_run("• 简报生成：一键生成Word格式简报\n")
features2.add_run("• 定时任务：支持早间/午间/晚间自动推送\n")
features2.add_run("• 关键词管理：自定义简报关键词\n")
doc.add_paragraph()

# 2.3 网站监测
add_heading_with_number(doc, "网站监测系统", 2, [2, 3])
doc.add_paragraph("网站栏目更新监测是政府网站绩效考核的核心工具。")
features3 = doc.add_paragraph()
features3.add_run("【主要功能】\n").bold = True
features3.add_run("• 全站扫描：自动扫描网站所有栏目\n")
features3.add_run("• 逾期预警：自动识别并标记逾期栏目\n")
features3.add_run("• 进度统计：实时统计更新完成率\n")
features3.add_run("• 报表导出：支持Excel/Word格式导出\n")
features3.add_run("• 定时监测：支持每日定时自动监测\n")
doc.add_paragraph()

# 2.4 档案管理
add_heading_with_number(doc, "档案管理系统", 2, [2, 4])
doc.add_paragraph("档案数字化管理，支持全宗、案卷、文件三级管理。")
features4 = doc.add_paragraph()
features4.add_run("【主要功能】\n").bold = True
features4.add_run("• 全宗管理：档案全宗信息维护\n")
features4.add_run("• 案卷管理：案卷创建、编目、整理\n")
features4.add_run("• 文件管理：档案文件上传、检索、借阅\n")
features4.add_run("• 批量导入：支持Excel批量导入\n")
features4.add_run("• 借阅管理：在线申请、审批、归还\n")
doc.add_paragraph()

# 2.5 知识管理
add_heading_with_number(doc, "知识管理系统", 2, [2, 5])
doc.add_paragraph("构建单位知识库，支持文档上传、智能检索、共享协作。")
features5 = doc.add_paragraph()
features5.add_run("【主要功能】\n").bold = True
features5.add_run("• 个人知识库：个人文档分类管理\n")
features5.add_run("• 共享知识库：团队协作共享\n")
features5.add_run("• 政策文库：政策文件集中管理\n")
features5.add_run("• 智能检索：全文搜索+语义搜索\n")
features5.add_run("• 收藏浏览：收藏、点赞、浏览记录\n")
doc.add_paragraph()

# 2.6 智能问答
add_heading_with_number(doc, "智能问答系统", 2, [2, 6])
doc.add_paragraph("基于AI的智能问答助手，支持知识库检索和政策解读。")
features6 = doc.add_paragraph()
features6.add_run("【主要功能】\n").bold = True
features6.add_run("• AI对话：智能问答交互\n")
features6.add_run("• 知识检索：RAG增强检索\n")
features6.add_run("• 政策解读：政策文件智能分析\n")
features6.add_run("• 多轮对话：上下文连续对话\n")
doc.add_paragraph()

# 2.7 会议管理
add_heading_with_number(doc, "会议管理系统", 2, [2, 7])
doc.add_paragraph("会议室预约、会议通知、会议纪要全流程管理。")
features7 = doc.add_paragraph()
features7.add_run("【主要功能】\n").bold = True
features7.add_run("• 会议室管理：会议室信息维护\n")
features7.add_run("• 会议预约：在线预约会议室\n")
features7.add_run("• 日历视图：会议日程可视化\n")
features7.add_run("• 会议通知：自动发送会议通知\n")
features7.add_run("• 会议纪要：纪要模板生成\n")
doc.add_paragraph()

# 2.8 工作日志
add_heading_with_number(doc, "工作日志系统", 2, [2, 8])
doc.add_paragraph("员工工作日志记录与领导点评。")
features8 = doc.add_paragraph()
features8.add_run("【主要功能】\n").bold = True
features8.add_run("• 日志填写：日报/周报/月报\n")
features8.add_run("• 日志审批：领导点评与审批\n")
features8.add_run("• 统计报表：工作量统计与分析\n")
doc.add_paragraph()

# 2.9 绩效考核
add_heading_with_number(doc, "绩效考核系统", 2, [2, 9])
doc.add_paragraph("绩效考核指标管理与评估。")
features9 = doc.add_paragraph()
features9.add_run("【主要功能】\n").bold = True
features9.add_run("• 考核周期：自定义考核周期\n")
features9.add_run("• 指标管理：考核指标体系\n")
features9.add_run("• 评估打分：在线评分\n")
features9.add_run("• 报表统计：考核结果分析\n")
doc.add_paragraph()

# 2.10 专报管理
add_heading_with_number(doc, "专报管理系统", 2, [2, 10])
doc.add_paragraph("专题报告的采集、编辑、报送全流程管理。")
features10 = doc.add_paragraph()
features10.add_run("【主要功能】\n").bold = True
features10.add_run("• 约稿任务：下发约稿通知\n")
features10.add_run("• 稿件提交：在线提交报告\n")
features10.add_run("• 编辑审核：编辑审核修改\n")
features10.add_run("• 采纳统计：采用情况统计\n")
features10.add_run("• AI辅助：AI辅助写作\n")
doc.add_paragraph()

# 2.11 督办管理
add_heading_with_number(doc, "督办管理系统", 2, [2, 11])
doc.add_paragraph("重要事项督办落实跟踪。")
features11 = doc.add_paragraph()
features11.add_run("【主要功能】\n").bold = True
features11.add_run("• 督办任务：任务创建与分配\n")
features11.add_run("• 进度跟踪：实时进度汇报\n")
features11.add_run("• 催办提醒：超时自动催办\n")
features11.add_run("• 完成确认：办结确认与评价\n")
doc.add_paragraph()

# 2.12 智能办公工具
add_heading_with_number(doc, "智能办公工具", 2, [2, 12])
doc.add_paragraph("AI辅助的智能办公工具集。")
features12 = doc.add_paragraph()
features12.add_run("【主要功能】\n").bold = True
features12.add_run("• 文档校对：语法、用词、格式校对\n")
features12.add_run("• 文档润色：措辞优化表达\n")
features12.add_run("• 会议纪要：自动生成纪要\n")
features12.add_run("• 辅助写作：各类文书智能生成\n")
features12.add_run("• 格式转换：Word/PDF互转\n")
features12.add_run("• 模板库：常用文书模板\n")
features12.add_run("• 意见建议：AI生成意见建议\n")
doc.add_paragraph()

# 2.13 系统管理
add_heading_with_number(doc, "系统管理", 2, [2, 13])
doc.add_paragraph("系统基础配置与用户权限管理。")
features13 = doc.add_paragraph()
features13.add_run("【主要功能】\n").bold = True
features13.add_run("• 用户管理：用户账号创建、角色分配\n")
features13.add_run("• 角色权限：RBAC权限管理\n")
features13.add_run("• 组织架构：部门、岗位管理\n")
features13.add_run("• 操作日志：系统操作审计\n")
features13.add_run("• AI配置：AI模型参数配置\n")
features13.add_run("• 系统配置：系统参数设置\n")
doc.add_paragraph()

# 第三章 优化升级建议
add_heading_with_number(doc, "优化升级建议", 1, [3])
doc.add_paragraph("基于对当前系统的全面评估，结合政府部门和企业智能化办公的最佳实践，提出以下优化升级建议：")

# 3.1 用户体验优化
add_heading_with_number(doc, "一、用户体验优化", 2, [3, 1])

doc.add_heading("1.1 界面现代化", 3)
ui_items = [
    "• 全面升级前端框架，采用Vue3+Element Plus打造现代化UI",
    "• 引入暗色模式切换，满足不同使用场景",
    "• 统一设计语言，组件风格一致性",
    "• 响应式布局，支持移动端访问",
    "• 动画交互优化，提升操作流畅度"
]
for item in ui_items:
    doc.add_paragraph(item)

doc.add_heading("1.2 操作简化", 3)
op_items = [
    "• 仪表盘重构：关键数据一目了然，支持自定义看板",
    "• 快捷操作：常用功能一键直达",
    "• 智能提示：操作引导与错误提示优化",
    "• 表单简化：减少填写步骤，智能填充",
    "• 批量操作：支持批量导入、批量审批"
]
for item in op_items:
    doc.add_paragraph(item)

# 3.2 智能化升级
add_heading_with_number(doc, "二、智能化功能升级", 2, [3, 2])

doc.add_heading("2.1 AI能力增强", 3)
ai_items = [
    "• 智能写作助手：接入国产大模型（文心一言/通义千问）",
    "• 智能审批：AI辅助审批决策建议",
    "• 智能分类：自动识别文档类型、部门",
    "• 智能推荐：基于使用习惯的智能推荐",
    "• 情感分析：舆情监测与情感分析"
]
for item in ai_items:
    doc.add_paragraph(item)

doc.add_heading("2.2 数据分析智能化", 3)
data_items = [
    "• 可视化大屏：数据驾驶舱实时展示",
    "• 预测分析：业务趋势预测",
    "• 关联分析：跨模块数据关联挖掘",
    "• 自定义报表：拖拽式报表设计器"
]
for item in data_items:
    doc.add_paragraph(item)

# 3.3 功能扩展
add_heading_with_number(doc, "三、功能模块扩展", 2, [3, 3])

doc.add_heading("3.1 政府场景专属功能", 3)
gov_items = [
    "• 政务公开：信息发布审核全流程",
    "• 民意征集：在线调查问卷",
    "• 建议提案：人大代表建议、政协提案管理",
    "• 信访管理：信访件全流程跟踪",
    "• 应急指挥：突发事件应急响应"
]
for item in gov_items:
    doc.add_paragraph(item)

doc.add_heading("3.2 企业场景专属功能", 3)
ent_items = [
    "• CRM客户管理：客户信息、跟进记录",
    "• 项目管理：项目立项、进度跟踪",
    "• 资产管理：固定资产管理",
    "• 采购管理：采购申请、审批、验收",
    "• 合同管理：合同起草、签署、存档"
]
for item in ent_items:
    doc.add_paragraph(item)

doc.add_heading("3.3 协同办公增强", 3)
collab_items = [
    "• 团队空间：项目团队协作空间",
    "• 在线文档：多人实时协同编辑",
    "• 任务看板：Kanban可视化任务管理",
    "• 日程共享：团队日程统一管理",
    "• 通讯录：组织架构与即时通讯"
]
for item in collab_items:
    doc.add_paragraph(item)

# 3.4 技术架构升级
add_heading_with_number(doc, "四、技术架构升级", 2, [3, 4])

doc.add_heading("4.1 性能优化", 3)
perf_items = [
    "• 数据库优化：索引优化、读写分离",
    "• 缓存策略：Redis缓存热点数据",
    "• 异步处理：任务队列解耦耗时操作",
    "• CDN加速：静态资源加速分发",
    "• 负载均衡：多实例部署支持"
]
for item in perf_items:
    doc.add_paragraph(item)

doc.add_heading("4.2 安全加固", 3)
sec_items = [
    "• 国密算法：支持SM2/SM3/SM4国产加密",
    "• 细粒度权限：数据级权限控制",
    "• 操作审计：完整操作日志追溯",
    "• 安全扫描：定期安全漏洞扫描",
    "• 数据备份：异地容灾备份"
]
for item in sec_items:
    doc.add_paragraph(item)

doc.add_heading("4.3 集成能力", 3)
int_items = [
    "• 微信集成：企业微信消息通知",
    "• 邮件集成：邮件收发集成",
    "• 钉钉/飞书：多平台对接",
    "• 打印集成：套打、连续打印",
    "• 数据接口：标准化API开放"
]
for item in int_items:
    doc.add_paragraph(item)

# 3.5 实施路线
add_heading_with_number(doc, "五、实施建议", 2, [3, 5])

doc.add_heading("5.1 分阶段实施计划", 3)
phases = [
    "【第一阶段·体验优化】(1-2个月)",
    "  • Vue3前端框架重构",
    "  • UI组件库升级",
    "  • 移动端适配",
    "",
    "【第二阶段·智能增强】(2-3个月)",
    "  • AI写作助手上线",
    "  • 智能推荐系统",
    "  • 数据分析大屏",
    "",
    "【第三阶段·功能扩展】(3-4个月)",
    "  • 企业/政府专属模块",
    "  • 协同办公增强",
    "  • 第三方集成",
    "",
    "【第四阶段·架构优化】(持续)",
    "  • 性能优化",
    "  • 安全加固",
    "  • 运维监控"
]
for phase in phases:
    doc.add_paragraph(phase)

doc.add_heading("5.2 预期效果", 3)
effects = [
    "• 用户满意度提升50%以上",
    "• 日常办公效率提升30%",
    "• 审批周期缩短40%",
    "• 数据利用价值提升60%",
    "• 系统可用性达到99.9%"
]
for effect in effects:
    doc.add_paragraph(effect)

# 第四章 总结
add_heading_with_number(doc, "总结", 1, [4])
doc.add_paragraph("OOA智能办公系统经过持续迭代，已形成较为完善的政府企业一体化办公解决方案。本次优化升级将以用户体验为中心，以AI智能为驱动，以业务协同为目标，打造新一代智慧办公平台。")
doc.add_paragraph()
doc.add_paragraph("通过本轮优化，系统将实现：")
summary_items = [
    "体验升级：从'能用'到'好用'的跨越",
    "智能升级：从'工具'到'助手'的转变",
    "协同升级：从'单点'到'全局'的整合",
    "价值升级：从'管理'到'赋能'的进化"
]
for item in summary_items:
    doc.add_paragraph(item)

# 保存
output_path = r"C:\Users\Administrator\Desktop\ooa\briefing_output\OOA智能服务系统用户手册.docx"
doc.save(output_path)
print(f"用户手册已生成: {output_path}")
