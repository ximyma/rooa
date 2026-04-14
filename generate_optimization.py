import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

doc = Document()
doc.add_heading('OOA智能办公系统综合优化建议报告', 0)

# 概述
doc.add_heading('一、现状评估', 1)
doc.add_heading('1.1 系统建设成果', 2)
doc.add_paragraph('OOA智能办公系统经过持续开发，已完成以下核心功能建设：')

table1 = doc.add_table(rows=6, cols=3)
table1.style = 'Table Grid'
headers = ['模块', '完成度', '数据规模']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '4472C4')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

data1 = [
    ['公文管理', '95%', '支持发文/收文/归档全流程'],
    ['智能简报', '90%', '61个数据源，84+历史简报'],
    ['网站监测', '85%', '5651个监测栏目'],
    ['档案管理', '80%', '全宗/案卷/文件三级管理'],
    ['智能问答', '75%', '支持RAG知识检索'],
]
for i, row_data in enumerate(data1, 1):
    for j, cell_data in enumerate(row_data):
        table1.rows[i].cells[j].text = cell_data

doc.add_paragraph()
doc.add_heading('1.2 存在的主要问题', 2)

problems = [
    ('体验层面', [
        '前端框架老旧(Bootstrap+jQuery)，界面不够现代化',
        '缺少深色模式，不同光照环境下使用体验不佳',
        '移动端适配不完善，无法满足移动办公需求',
        '操作路径较长，常见功能需要多次点击',
    ]),
    ('功能层面', [
        'AI能力有限，仅支持基础问答，缺乏智能写作等高级功能',
        '数据分析能力薄弱，缺少可视化大屏和智能分析',
        '跨系统集成不足，与企业微信/钉钉等未完全打通',
        '个性化程度低，无法根据用户角色定制工作台',
    ]),
    ('技术层面', [
        '后端架构为单体应用，扩展性受限',
        '数据库缺少读写分离，高并发性能不足',
        '缺少完善的缓存机制，重复查询效率低',
        '前端代码缺少组件化，迭代维护成本高',
    ]),
]

for cat, items in problems:
    doc.add_heading(f'{cat}', 3)
    for item in items:
        doc.add_paragraph(f'• {item}')

doc.add_heading('二、优化升级方案', 1)

doc.add_heading('2.1 第一阶段：体验焕新（预计2个月）', 2)

doc.add_heading('2.1.1 前端架构升级', 3)
p1 = doc.add_paragraph()
p1.add_run('目标：').bold = True
p1.add_run('采用Vue3 + Element Plus重构前端，打造现代化、组件化的用户界面')
doc.add_paragraph()
items1 = [
    '技术选型：Vue3 Composition API + Vite + Element Plus',
    '组件库：封装通用业务组件，形成组件市场',
    '状态管理：Pinia替换Vuex，更轻量易用',
    '路由管理：Vue Router 4，动态路由权限控制',
    '构建工具：Vite 3，开发体验提升10倍',
]
for item in items1:
    doc.add_paragraph(f'• {item}')

doc.add_heading('2.1.2 UI/UX优化', 3)
items2 = [
    '深色/浅色模式一键切换，主题色自定义',
    '全新首页工作台，重要事项一目了然',
    '侧边栏智能折叠，常用菜单优先展示',
    '操作引导动画，新用户快速上手',
    '全站搜索(Ctrl+K)，一键直达任何功能',
]
for item in items2:
    doc.add_paragraph(f'• {item}')

doc.add_heading('2.1.3 移动端适配', 3)
items3 = [
    '响应式布局，PC/平板/手机自适应',
    '手势操作支持，左滑返回、下拉刷新',
    '离线缓存，弱网环境正常访问',
    '消息推送，企业微信/短信双通道',
]
for item in items3:
    doc.add_paragraph(f'• {item}')

doc.add_heading('2.2 第二阶段：智能跃升（预计3个月）', 2)

doc.add_heading('2.2.1 AI能力升级', 3)
p2 = doc.add_paragraph()
p2.add_run('目标：').bold = True
p2.add_run('深度整合国产大模型，打造真正智能的办公助手')
doc.add_paragraph()

ai_features = [
    ('智能写作助手', [
        '公文智能起草，输入关键词自动生成初稿',
        '会议纪要生成，录音转写+智能摘要',
        '新闻简报自动生成，定时推送',
        '政策解读分析，自动提取要点',
    ]),
    ('智能审批', [
        '审批意见智能推荐，参考历史案例',
        '风险智能提醒，异常情况预警',
        '流程自动流转，条件触发智能跳转',
    ]),
    ('智能分析', [
        '数据驾驶舱，关键指标实时监控',
        '趋势预测分析，业务决策支持',
        '异常自动发现，数据问题预警',
    ]),
]
for title, items in ai_features:
    p = doc.add_paragraph()
    p.add_run(f'{title}：').bold = True
    doc.add_paragraph('  ' + '、'.join(items))

doc.add_heading('2.2.2 大模型集成方案', 3)
llm_table = doc.add_table(rows=5, cols=3)
llm_table.style = 'Table Grid'
llm_headers = ['模型', '适用场景', '部署方式']
for i, h in enumerate(llm_headers):
    cell = llm_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '4472C4')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

llm_data = [
    ['通义千问(Qwen)', '通用问答、写作辅助', '阿里云API'],
    ['文心一言(ERNIE)', '中文理解、政务问答', '百度云API'],
    ['智谱ChatGLM', '本地部署、数据安全', '私有化部署'],
    ['DeepSeek', '代码辅助、数据分析', '混合部署'],
]
for i, row_data in enumerate(llm_data, 1):
    for j, cell_data in enumerate(row_data):
        llm_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

doc.add_heading('2.3 第三阶段：业务扩展（预计3个月）', 2)

doc.add_heading('2.3.1 政府场景增强模块', 3)
gov_items = [
    ('政务公开管理', '信息发布审核流程、主动公开清单、依申请公开'),
    ('建议提案管理', '人大代表建议、政协提案的交办、督办、答复'),
    ('信访管理系统', '来信来访登记、转办、办理、回复全流程'),
    ('应急值班管理', '值班排班、突发事件上报、应急调度'),
    ('目标考核系统', '部门/个人目标制定、过程管理、年终考核'),
]
for title, desc in gov_items:
    p = doc.add_paragraph()
    p.add_run(f'• {title}：').bold = True
    p.add_run(desc)

doc.add_heading('2.3.2 企业场景增强模块', 3)
ent_items = [
    ('CRM客户管理', '客户信息、商机跟进、销售漏斗'),
    ('项目管理', '项目立项、任务分解、甘特图展示'),
    ('资产管理', '资产登记、盘点、维保提醒'),
    ('采购管理', '采购申请、供应商管理、合同执行'),
    ('预算管理', '年度预算编制、执行监控、差异分析'),
]
for title, desc in ent_items:
    p = doc.add_paragraph()
    p.add_run(f'• {title}：').bold = True
    p.add_run(desc)

doc.add_heading('2.4 第四阶段：架构优化（持续迭代）', 2)

doc.add_heading('2.4.1 性能优化', 3)
perf_items = [
    '数据库优化：慢查询分析、索引优化、分库分表',
    '缓存体系：Redis缓存热点数据、本地缓存减少请求',
    '异步处理：Celery任务队列、邮件/消息异步发送',
    'CDN加速：静态资源分离、全球加速节点',
    '负载均衡：Nginx负载分发、健康检查自动切换',
]
for item in perf_items:
    doc.add_paragraph(f'• {item}')

doc.add_heading('2.4.2 安全加固', 3)
sec_items = [
    '国密算法：支持SM2/SM3/SM4，满足合规要求',
    '细粒度权限：数据级权限控制、行级权限管控',
    '操作审计：完整日志记录、异常行为告警',
    '数据安全：敏感数据脱敏、传输加密、备份恢复',
    '渗透测试：定期安全扫描、漏洞修复',
]
for item in sec_items:
    doc.add_paragraph(f'• {item}')

doc.add_heading('2.4.3 集成扩展', 3)
int_items = [
    '企业微信深度集成：消息、小程序、审批流',
    '钉钉/飞书对接：多平台统一入口',
    '电子签章：CA数字证书、在线签署',
    '档案对接：与档案馆系统标准接口',
    '数据交换：与上级系统数据上报',
]
for item in int_items:
    doc.add_paragraph(f'• {item}')

doc.add_heading('三、实施保障', 1)

doc.add_heading('3.1 项目组织', 2)
org_table = doc.add_table(rows=5, cols=2)
org_table.style = 'Table Grid'
for i, (role, duty) in enumerate([
    ('项目总监', '整体把控、资源协调'),
    ('产品经理', '需求分析、方案设计'),
    ('技术负责人', '架构设计、技术选型'),
    ('开发团队', '2后端+2前端+1AI工程师'),
], 0):
    org_table.rows[i].cells[0].text = role
    org_table.rows[i].cells[1].text = duty

doc.add_paragraph()
doc.add_heading('3.2 风险控制', 2)
risks = [
    ('需求变更风险', '措施：敏捷开发，2周迭代，持续交付'),
    ('技术风险', '措施：技术预研，POC验证后再上线'),
    ('数据安全风险', '措施：分级授权，敏感数据加密'),
    ('用户体验风险', '措施：灰度发布，用户反馈快速响应'),
]
for risk, measure in risks:
    p = doc.add_paragraph()
    p.add_run(f'• {risk}：').bold = True
    p.add_run(measure)

doc.add_heading('3.3 预期收益', 2)
benefits = [
    ('效率提升', '日常操作时间减少40%，审批周期缩短50%'),
    ('体验改善', '用户满意度提升60%，培训成本降低50%'),
    ('管理升级', '数据驱动决策，管理工作量减少30%'),
    ('智能化水平', 'AI辅助覆盖80%日常场景，真正减负增效'),
]
for benefit, desc in benefits:
    p = doc.add_paragraph()
    p.add_run(f'• {benefit}：').bold = True
    p.add_run(desc)

doc.add_heading('四、总结', 1)
doc.add_paragraph('本优化建议从用户体验、智能能力、业务扩展、技术架构四个维度提出系统升级方案，通过分阶段实施，逐步实现OOA系统从"好用"到"智能"的跨越。')
doc.add_paragraph()
doc.add_paragraph('优化后的系统将成为政府企业智能化办公的标杆产品，为用户提供更高效、更智能、更安全的办公体验。')

# 保存
output_path = r"C:\Users\Administrator\Desktop\ooa\briefing_output\OOA系统优化升级建议报告.docx"
doc.save(output_path)
print(f"优化建议报告已生成: {output_path}")
