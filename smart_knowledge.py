"""
OA系统智能知识库模块
基于 knmchat4.py 改造，整合到 Flask OA 系统
使用本地嵌入模型和 TesseractOCR
"""

import os
import re
import json
import pickle
import numpy as np
from datetime import datetime
from collections import Counter

# --- 强制离线加载，防止任何网络访问导致启动卡顿 ---
os.environ['HF_HUB_OFFLINE'] = '1'
os.environ['TRANSFORMERS_OFFLINE'] = '1'
os.environ['HF_HOME'] = ''
os.environ['TRANSFORMERS_CACHE'] = ''
# 静默 torch 的 INFO/WARNING 日志，减少启动噪音
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

# 懒加载：不在模块导入时加载重型库，仅在使用时 import
jieba = None
jieba_analyse = None
docx_Document = None
PdfReader = None
openpyxl = None
PIL_Image = None
SentenceTransformer_cls = None
pytesseract = None

# 数据库
from models import db, KnowledgeBase, KnowledgeFile, KnowledgeBrowseLog

# 配置 - 使用配置管理器
from config_manager import config_manager

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 从配置获取路径（支持动态修改）
def get_embedding_model_path():
    return config_manager.get_embedding_model_path()

def get_tesseract_cmd():
    return config_manager.get_tesseract_cmd()

def get_tessdata_dir():
    return config_manager.get_tessdata_dir()


class SmartKnowledgeBase:
    """智能知识库管理器（懒加载版）"""

    def __init__(self):
        self._model = None   # 懒加载，不在构造时加载
        self._ocr_ready = False

    @property
    def model(self):
        """首次访问时再加载模型，避免拖慢应用启动"""
        if self._model is None:
            self._load_model()
        return self._model

    def _load_model(self):
        """加载嵌入模型，只在首次使用时触发"""
        global SentenceTransformer_cls
        if SentenceTransformer_cls is None:
            try:
                from sentence_transformers import SentenceTransformer as ST
                SentenceTransformer_cls = ST
            except ImportError:
                SentenceTransformer_cls = False
                print("[警告] sentence-transformers 未安装，向量检索不可用")
                return

        if not SentenceTransformer_cls:
            return

        model_path = get_embedding_model_path()
        if not os.path.exists(model_path):
            print(f"[警告] 模型路径不存在: {model_path}，使用关键词检索")
            return

        try:
            print(f"[信息] 懒加载嵌入模型: {model_path}")
            # 明确指定 device='cpu'，避免 GPU 初始化或CUDA探测耗时
            self._model = SentenceTransformer_cls(model_path, device='cpu', local_files_only=True)
            print(f"[成功] 嵌入模型加载完成，向量维度: {self._model.get_sentence_embedding_dimension()}")
        except Exception as e:
            print(f"[错误] 模型加载失败: {e}，使用关键词检索")
            self._model = None

    def _ensure_ocr(self):
        """懒加载 OCR"""
        global pytesseract
        if not self._ocr_ready:
            try:
                import pytesseract as pt
                pytesseract = pt
                tesseract_cmd = get_tesseract_cmd()
                if os.path.exists(tesseract_cmd):
                    import pytesseract as pt_module
                    pt_module.pytesseract.tesseract_cmd = tesseract_cmd
                    tessdata_dir = get_tessdata_dir()
                    if tessdata_dir and os.path.exists(tessdata_dir):
                        os.environ['TESSDATA_PREFIX'] = tessdata_dir
                self._ocr_ready = True
            except Exception as e:
                print(f"[警告] TesseractOCR 初始化失败: {e}")
                self._ocr_ready = True  # 只报一次

    def _ensure_jieba(self):
        global jieba, jieba_analyse
        if jieba is None:
            import jieba as j
            import jieba.analyse as ja
            jieba = j
            jieba_analyse = ja

    def _ensure_doclibs(self):
        global docx_Document, PdfReader, openpyxl, PIL_Image
        if docx_Document is None:
            from docx import Document as D
            from PyPDF2 import PdfReader as PR
            import openpyxl as ox
            from PIL import Image as PI
            docx_Document = D
            PdfReader = PR
            openpyxl = ox
            PIL_Image = PI
    
    # ==================== 文档内容提取 ====================
    
    def extract_content(self, file_path):
        """
        根据文件类型提取文本内容
        支持: PDF, Word, Excel, 图片(OCR), 文本文件
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.pdf':
                return self._extract_pdf(file_path)
            elif ext in ('.doc', '.docx'):
                return self._extract_word(file_path)
            elif ext in ('.xls', '.xlsx'):
                return self._extract_excel(file_path)
            elif ext in ('.jpg', '.jpeg', '.png', '.bmp', '.tiff'):
                return self._extract_image(file_path)
            elif ext in ('.txt', '.md', '.csv'):
                return self._extract_text(file_path)
            else:
                return f"[不支持的文件类型: {ext}]"
        except Exception as e:
            return f"[提取失败: {str(e)}]"
    
    def _extract_pdf(self, file_path):
        """提取PDF文本"""
        self._ensure_doclibs()
        text = []
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n'.join(text)
    
    def _extract_word(self, file_path):
        """提取Word文本，支持 .docx 和旧式 .doc (OLE) 格式"""
        self._ensure_doclibs()
        ext = os.path.splitext(file_path)[1].lower()

        # ── .docx：直接用 python-docx ──────────────────────────
        if ext == '.docx':
            doc = docx_Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(paragraphs)

        # ── .doc：先尝试 python-docx（有些.doc其实是docx改后缀）──
        try:
            doc = docx_Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                return '\n'.join(paragraphs)
        except Exception:
            pass

        # ── .doc：降级到 olefile + UTF-16LE 解析 ────────────────
        try:
            import olefile as olefile_lib
            import re as re_mod
            if not olefile_lib.isOleFile(file_path):
                return '[不支持的文件格式，无法提取内容]'
            ole = olefile_lib.OleFileIO(file_path)
            if not ole.exists('WordDocument'):
                ole.close()
                return '[未找到 WordDocument 流]'
            raw = ole.openstream('WordDocument').read()
            ole.close()

            # UTF-16LE 解码
            text_raw = raw.decode('utf-16-le', errors='ignore')

            # 过滤：保留连续的可读中文和ASCII段落（去除乱码短片段）
            # Word OLE 流中正文夹杂大量二进制控制字节，只取"纯文本"部分
            chunks = re_mod.split(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]{2,}', text_raw)
            lines = []
            for chunk in chunks:
                chunk = chunk.strip()
                # 要求长度 >= 4，且中文比例或可读字符比例合理
                if len(chunk) < 4:
                    continue
                chinese_count = len(re_mod.findall(r'[\u4e00-\u9fff]', chunk))
                ascii_word_count = len(re_mod.findall(r'[a-zA-Z0-9]', chunk))
                total = len(chunk)
                if (chinese_count + ascii_word_count) / max(total, 1) >= 0.3:
                    # 清理残余控制字符和常见乱码前缀
                    chunk = re_mod.sub(r'[\x00-\x1f\x7f]', '', chunk)
                    if chunk.strip():
                        lines.append(chunk.strip())
            return '\n'.join(lines) if lines else '[文档内容为空或全为图像]'
        except Exception as e:
            return f'[提取失败: {str(e)}]'
    
    def _extract_excel(self, file_path):
        """提取Excel文本"""
        self._ensure_doclibs()
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        texts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    texts.append(row_text)
        return '\n'.join(texts)
    
    def _extract_image(self, file_path):
        """图片OCR识别"""
        self._ensure_ocr()
        if pytesseract is None:
            return "[OCR不可用: 未安装pytesseract]"

        if not config_manager.is_ocr_enabled():
            return "[OCR已禁用]"

        try:
            self._ensure_doclibs()
            image = PIL_Image.open(file_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            lang = config_manager.get_ocr_languages()
            text = pytesseract.image_to_string(image, lang=lang)
            return text.strip()
        except Exception as e:
            return f"[OCR失败: {str(e)}]"

    
    def _extract_text(self, file_path):
        """提取纯文本"""
        # 检测编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return "[文本编码无法识别]"
    
    # ==================== 智能分析 ====================
    
    def generate_keywords(self, content, top_k=10):
        """
        使用jieba提取关键词
        """
        if not content or len(content) < 10:
            return []
        self._ensure_jieba()
        keywords = jieba_analyse.extract_tags(
            content,
            topK=top_k,
            withWeight=False,
            allowPOS=('ns', 'n', 'vn', 'v', 'nr', 'nt', 'nw')
        )
        return keywords
    
    def generate_summary(self, content, max_length=200):
        """
        生成文本摘要（取前几句）
        """
        if not content:
            return ""
        
        # 清理文本
        content = re.sub(r'\s+', ' ', content).strip()
        
        # 尝试提取第一段有意义的文字
        sentences = re.split(r'[。！？.!?]', content)
        summary = []
        length = 0
        
        for sent in sentences:
            sent = sent.strip()
            if len(sent) > 10:  # 过滤太短的句子
                summary.append(sent)
                length += len(sent)
                if length >= max_length:
                    break
        
        result = '。'.join(summary)
        if len(result) > max_length:
            result = result[:max_length] + '...'
        return result
    
    def generate_embedding(self, content):
        """
        生成文本向量嵌入
        """
        if not self.model or not content:
            return None
        
        try:
            # 限制长度，避免内存问题
            text = content[:5000] if len(content) > 5000 else content
            vector = self.model.encode(text)
            return vector
        except Exception as e:
            print(f"[错误] 向量生成失败: {e}")
            return None
    
    def suggest_title(self, content, original_name):
        """
        智能生成文档标题建议
        """
        if not content:
            return original_name
        
        # 提取关键词
        keywords = self.generate_keywords(content, top_k=3)
        
        # 尝试从内容中提取标题（通常在第一行）
        first_line = content.strip().split('\n')[0][:50]
        
        # 如果第一行看起来像标题（短且有关键词）
        if len(first_line) < 30 and len(first_line) > 5:
            return first_line
        
        # 否则使用关键词组合
        if keywords:
            return ' | '.join(keywords[:3])
        
        # 回退到原始文件名（去掉扩展名）
        return os.path.splitext(original_name)[0]
    
    def auto_tag(self, content, existing_tags=None):
        """
        自动生成标签
        """
        if not content:
            return []
        
        # 预定义标签类别
        tag_categories = {
            '政策': ['政策', '规定', '通知', '文件', '制度', '法规', '条例'],
            '技术': ['技术', '开发', '代码', '程序', '系统', '软件', '硬件'],
            '会议': ['会议', '纪要', '讨论', '决议', '议程', '报告'],
            '财务': ['财务', '预算', '报销', '费用', '资金', '会计', '审计'],
            '人事': ['人事', '招聘', '考核', '绩效', '薪酬', '员工', '培训'],
            '行政': ['行政', '办公', '后勤', '档案', '印章', '接待'],
            '项目': ['项目', '计划', '进度', '实施', '验收', '成果'],
        }
        
        auto_tags = []
        content_lower = content.lower()
        
        # 根据关键词匹配标签
        for tag, keywords in tag_categories.items():
            for kw in keywords:
                if kw in content_lower:
                    auto_tags.append(tag)
                    break
        
        # 提取高频词作为标签
        keywords = self.generate_keywords(content, top_k=5)
        auto_tags.extend(keywords[:3])
        
        # 去重
        auto_tags = list(set(auto_tags))
        
        return auto_tags
    
    # ==================== 批量处理 ====================
    
    def process_files_batch(self, file_paths, kb_id, uploader_id, 
                           progress_callback=None, status_callback=None):
        """
        批量处理文件上传
        
        Args:
            file_paths: 文件路径列表
            kb_id: 知识库ID
            uploader_id: 上传者ID
            progress_callback: 进度回调函数(current, total)
            status_callback: 状态回调函数(message)
        
        Returns:
            dict: {success: int, failed: int, errors: list}
        """
        results = {'success': 0, 'failed': 0, 'errors': []}
        total = len(file_paths)
        
        for i, file_path in enumerate(file_paths, 1):
            if progress_callback:
                progress_callback(i, total)
            
            if status_callback:
                status_callback(f"处理 {i}/{total}: {os.path.basename(file_path)}")
            
            try:
                # 提取内容
                content = self.extract_content(file_path)
                
                if not content or content.startswith('['):
                    results['failed'] += 1
                    results['errors'].append(f"{os.path.basename(file_path)}: 内容提取失败")
                    continue
                
                # 智能分析
                keywords = self.generate_keywords(content)
                summary = self.generate_summary(content)
                tags = self.auto_tag(content)
                vector = self.generate_embedding(content)
                suggested_title = self.suggest_title(content, os.path.basename(file_path))
                
                # 保存到数据库
                from oa_utils import save_upload_file
                saved_path = save_upload_file(
                    type('File', (), {
                        'filename': os.path.basename(file_path),
                        'save': lambda p: shutil.copy(file_path, p)
                    })(),
                    f'knowledge/{kb_id}'
                )
                
                if not saved_path:
                    raise Exception("文件保存失败")
                
                # 获取文件信息
                file_size = os.path.getsize(file_path)
                word_count = len(content)
                
                # 创建知识库文件记录
                kf = KnowledgeFile(
                    filename=os.path.basename(saved_path),
                    original_name=os.path.basename(file_path),
                    file_path=saved_path,
                    knowledge_base_id=kb_id,
                    uploaded_by=uploader_id,
                    status='approved',
                    content_text=content,
                    file_size=file_size,
                    file_type=os.path.splitext(file_path)[1][1:].lower(),
                    word_count=word_count,
                    category='',
                    tags=','.join(tags),
                    summary=summary,
                    # 存储向量和关键词（需要扩展模型）
                )
                
                db.session.add(kf)
                db.session.commit()
                
                # 更新FTS索引
                from app import update_fts_index
                update_fts_index(kf)
                
                results['success'] += 1
                
            except Exception as e:
                results['failed'] += 1
                results['errors'].append(f"{os.path.basename(file_path)}: {str(e)}")
        
        return results
    
    # ==================== 检索功能 ====================
    
    def search_similar(self, query, kb_ids=None, top_k=10):
        """
        向量/相似度检索

        当前实现以关键词相似度为主，模型可用时预留向量能力；
        即使模型未加载，也不应该直接返回空结果。
        """
        query_text = (query or '').strip()
        if not query_text:
            return []

        # 模型可用时尝试编码，但当前检索主逻辑仍然使用轻量相似度回退。
        if self.model:
            try:
                self.model.encode(query_text)
            except Exception:
                pass

        candidate_query = KnowledgeFile.query.filter(KnowledgeFile.status == 'approved')
        if kb_ids:
            candidate_query = candidate_query.filter(KnowledgeFile.knowledge_base_id.in_(kb_ids))

        results = []
        files = candidate_query.all()

        for f in files:
            score = self._calculate_similarity(query_text, f)
            if score > 0:
                snippet_source = f.content_text or f.summary or ''
                snippet = (snippet_source[:200] + '...') if len(snippet_source) > 200 else snippet_source
                results.append({
                    'file': f,
                    'score': score,
                    'snippet': snippet
                })

        results.sort(key=lambda x: x['score'], reverse=True)
        return results[:top_k]
    
    def _calculate_similarity(self, query, knowledge_file):
        """
        计算查询与文档的相似度（简化版）
        """
        self._ensure_jieba()
        query_keywords = set(jieba.lcut(query.lower()))
        doc_keywords = set(jieba.lcut((knowledge_file.content_text or '').lower()))
        
        if not query_keywords or not doc_keywords:
            return 0
        
        intersection = query_keywords & doc_keywords
        union = query_keywords | doc_keywords
        
        return len(intersection) / len(union) if union else 0


# 全局实例
smart_kb = SmartKnowledgeBase()
