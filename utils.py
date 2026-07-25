import os
import random
import logging
from werkzeug.utils import secure_filename
from flask import current_app
from datetime import datetime
import requests
import time
import json
from config import Config

# 配置日志
log_level = getattr(logging, os.environ.get('LOG_LEVEL', 'INFO').upper())
logging.basicConfig(
    level=log_level,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# utils.py 新增

def generate_ai_response(model_config, prompt, knowledge_context=''):
    """
    调用AI模型生成响应，支持知识库上下文
    """
    messages = []
    if knowledge_context:
        messages.append({
            'role': 'system',
            'content': f"以下是相关参考资料：\n{knowledge_context}\n\n请基于以上资料回答用户问题。"
        })
    messages.append({'role': 'user', 'content': prompt})
    return call_ai_model(model_config, messages)

def call_ai_model(config, messages):
    """调用AI模型，支持多种提供商"""
    # 模拟延时
    if config.delay > 0:
        time.sleep(config.delay / 1000.0)

    provider_map = {
        'openai': 'OpenAI',
        'deepseek': 'DeepSeek',
        'siliconflow': '硅基流动',
    }
    provider_name = provider_map.get(config.provider, config.provider)

    if config.provider in ('openai', 'deepseek', 'siliconflow'):
        return _call_openai_compatible(config, messages, provider_name)
    elif config.provider == 'local':
        return call_local(config, messages)
    else:
        return f"不支持的模型提供商: {config.provider}"


def _call_openai_compatible(config, messages, provider_name='AI'):
    """统一的 OpenAI 兼容 API 调用"""
    headers = {
        "Authorization": f"Bearer {config.api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": config.model_name,
        "messages": messages,
        "temperature": config.temperature,
        "max_tokens": config.max_tokens
    }
    try:
        url = _build_api_url(config.api_base, 'chat/completions')
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            logger.error(f"{provider_name} API错误: {response.status_code} - {response.text[:300]}")
            return f"{provider_name} API错误: {response.text[:200]}"
    except requests.exceptions.Timeout:
        logger.error(f"{provider_name} API请求超时")
        return f"{provider_name} API请求超时，请稍后重试"
    except requests.exceptions.RequestException as e:
        logger.error(f"{provider_name} API请求失败: {str(e)}")
        return f"请求失败: {str(e)}"

# 保留旧函数名作为别名，确保兼容性
call_openai = lambda config, messages: _call_openai_compatible(config, messages, 'OpenAI')
call_deepseek = lambda config, messages: _call_openai_compatible(config, messages, 'DeepSeek')
call_siliconflow = lambda config, messages: _call_openai_compatible(config, messages, '硅基流动')


def _build_api_url(api_base, path):
    """拼接 API URL，正确处理 /v1 路径"""
    from urllib.parse import urljoin
    base = api_base.rstrip('/')
    # 如果 base 已经以 /v1 结尾，直接追加 path
    if base.endswith('/v1'):
        return f"{base}/{path.lstrip('/')}"
    # 否则加上 /v1/
    return f"{base}/v1/{path.lstrip('/')}"

def call_local(config, messages):
    # 本地模型如Ollama/vLLM，可使用OpenAI格式
    headers = {"Content-Type": "application/json"}
    payload = {
        "model": config.model_name,
        "messages": messages,
        "temperature": config.temperature,
        "max_tokens": config.max_tokens
    }
    try:
        url = _build_api_url(config.api_base, 'chat/completions')
        logger.info(f"调用本地模型: {url}, model={config.model_name}")
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            logger.error(f"本地模型错误: {response.status_code} - {response.text}")
            return f"本地模型错误: {response.text}"
    except requests.exceptions.Timeout:
        logger.error("本地模型请求超时")
        return "请求超时，请稍后重试"
    except requests.exceptions.RequestException as e:
        logger.error(f"本地模型请求失败: {str(e)}")
        return f"请求失败: {str(e)}"
    except Exception as e:
        logger.error(f"本地模型未知错误: {str(e)}")
        return f"请求失败: {str(e)}"



def allowed_file(filename, allowed_extensions=None):
    if allowed_extensions is None:
        allowed_extensions = Config.ALLOWED_EXTENSIONS
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def save_upload_file(file, subfolder=''):
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], subfolder)
        os.makedirs(upload_dir, exist_ok=True)
        filepath = os.path.join(upload_dir, unique_filename)
        file.save(filepath)
        return filepath
    return None

def ai_document_writing(template, keywords):
    return f"根据模板【{template}】和关键词【{keywords}】生成的公文草稿：\n\n关于{keywords}的通知\n\n各相关部门：\n\n为了进一步推动{keywords}工作，现提出以下要求：\n1. 加强组织领导。\n2. 强化责任落实。\n3. 确保工作实效。\n\n请各单位认真执行。"

def ai_polish(text):
    return f"润色后的文本：\n{text}\n\n（此处为AI润色结果，实际可调用大模型API）"

def ai_proofread(text, scheme='standard'):
    errors = []
    if '错误' in text:
        errors.append('发现可能错别字："错误"应为"正确"')
    if '落马官员' in text and scheme == 'sensitive':
        errors.append('敏感信息：包含落马官员相关内容，请注意审查')
    if not errors:
        return "未发现明显错误。"
    return "\n".join(errors)

def ai_suggestion(content, department_duty):
    summary = f"摘要：{content[:100]}..."
    opinion = f"拟办意见：请{department_duty}根据上述情况，尽快办理。"
    return summary, opinion


# ===== 文件内容验证（魔数检测） =====
# 常见文件类型的魔数（文件头字节签名）
_FILE_SIGNATURES = {
    'pdf': [b'%PDF'],
    'docx': [b'PK\x03\x04'],  # ZIP-based (docx/xlsx/pptx)
    'doc': [b'\xd0\xcf\x11\xe0'],  # OLE2
    'xls': [b'\xd0\xcf\x11\xe0'],
    'xlsx': [b'PK\x03\x04'],
    'txt': None,  # 纯文本无固定魔数
    'md': None,
    'png': [b'\x89PNG\r\n\x1a\n'],
    'jpg': [b'\xff\xd8\xff'],
    'jpeg': [b'\xff\xd8\xff'],
    'gif': [b'GIF89a', b'GIF87a'],
    'bmp': [b'BM'],
    'zip': [b'PK\x03\x04'],
    'rar': [b'Rar!\x1a\x07'],
    '7z': [b"7z\xbc\xaf'\x1c"],
}

def validate_file_content(filepath, claimed_extension=None):
    """
    通过文件头魔数验证文件内容与其声称的扩展名是否一致。
    返回: (is_valid: bool, detected_type: str)
    """
    if claimed_extension is None:
        claimed_extension = filepath.rsplit('.', 1)[-1].lower() if '.' in filepath else ''
    
    if claimed_extension not in _FILE_SIGNATURES:
        return True, claimed_extension  # 未知类型，不阻止
    
    expected_sigs = _FILE_SIGNATURES[claimed_extension]
    if expected_sigs is None:
        return True, claimed_extension  # 纯文本无魔数
    
    try:
        with open(filepath, 'rb') as f:
            header = f.read(16)
        
        for sig in expected_sigs:
            if header.startswith(sig):
                return True, claimed_extension
        
        logger.warning(f"文件魔数不匹配: {filepath}, 声称扩展名={claimed_extension}, 文件头={header[:8].hex()}")
        return False, claimed_extension
    except Exception as e:
        logger.warning(f"文件内容验证失败: {filepath}, {e}")
        return True, claimed_extension  # 无法读取时不过度阻止


def pdf_to_text(pdf_path):
    import PyPDF2
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
        return text
    except Exception as e:
        logger.error(f"PDF解析失败: {str(e)}")
        return f"PDF解析失败: {str(e)}"


def docx_to_text(docx_path):
    """从 Word (.docx) 文档提取纯文本"""
    try:
        from docx import Document
        doc = Document(docx_path)
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        # 也读取表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return '\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX解析失败: {str(e)}")
        return f"DOCX解析失败: {str(e)}"


def docx_to_html(docx_path):
    """将 Word (.docx) 转换为带基本HTML格式的内容（用于预览）"""
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document(docx_path)
        html_parts = ['<div class="docx-preview">']
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                html_parts.append('<p>&nbsp;</p>')
                continue
            style_name = p.style.name if p.style else 'Normal'
            if 'Heading' in style_name or '标题' in style_name:
                # 提取标题级别
                level = style_name[-1] if style_name[-1].isdigit() else '2'
                html_parts.append(f'<h{level}>{text}</h{level}>')
            else:
                # 处理段落内的部分格式（粗体、斜体）
                runs_html = ''
                for run in p.runs:
                    r_text = run.text
                    if run.bold:
                        r_text = f'<strong>{r_text}</strong>'
                    if run.italic:
                        r_text = f'<em>{r_text}</em>'
                    runs_html += r_text
                html_parts.append(f'<p>{runs_html or text}</p>')
        # 处理表格
        for table in doc.tables:
            html_parts.append('<table class="table table-bordered table-sm">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text.strip()}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        html_parts.append('</div>')
        return '\n'.join(html_parts)
    except Exception as e:
        logger.error(f"DOCX转HTML失败: {docx_path}, {e}")
        return f'<p>预览失败: {str(e)}</p>'


def extract_file_content(file_path, apply_length_limit=True):
    """
    根据文件扩展名自动提取文本内容
    支持: txt, md, pdf, docx, doc
    
    参数:
        apply_length_limit (bool): 是否应用文本长度限制
    """
    import os
    if not os.path.exists(file_path):
        return ""

    ext = file_path.lower().rsplit('.', 1)[-1] if '.' in file_path else ''
    
    extracted_text = ""
    
    if ext in ('txt', 'md', 'markdown'):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
        except Exception as e:
            logger.error(f"文本文件读取失败: {file_path}, {e}")
            extracted_text = ""

    elif ext == 'pdf':
        extracted_text = pdf_to_text(file_path)

    elif ext in ('docx',):
        extracted_text = docx_to_text(file_path)

    elif ext == 'doc':
        # 老版 doc 格式，尝试用 python-docx 直接打开
        try:
            extracted_text = docx_to_text(file_path)
        except Exception as e:
            logger.error(f"DOC解析失败: {file_path}, {e}")
            extracted_text = "[老版doc格式暂无法解析，请另存为docx格式]"
    
    # 应用文本长度限制（如果启用）
    if apply_length_limit and extracted_text:
        try:
            # 从config_manager获取限制
            from config_manager import config_manager
            max_length = config_manager.get("document.max_extracted_length", -1)
        except ImportError:
            # 如果无法导入配置管理器，使用默认配置
            max_length = Config.MAX_EXTRACTED_TEXT_LENGTH
            
        if max_length == -1:
            # -1表示无限制，不进行截断
            return extracted_text
        elif max_length > 0 and len(extracted_text) > max_length:
            logger.info(f"文本内容过长，截断至 {max_length} 字符: {file_path}")
            extracted_text = extracted_text[:max_length]
    
    return extracted_text


def get_file_metadata(file_path, apply_length_limit=True):
    """
    获取文件元数据：(file_size, file_type, word_count)
    
    参数:
        apply_length_limit (bool): 是否应用文本长度限制
    """
    import os
    try:
        size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
    except Exception:
        size = 0

    ext = file_path.lower().rsplit('.', 1)[-1] if '.' in file_path else ''
    type_map = {
        'pdf': 'pdf', 'docx': 'docx', 'doc': 'doc',
        'txt': 'txt', 'md': 'md', 'markdown': 'md'
    }
    file_type = type_map.get(ext, ext if ext else 'unknown')
    text = extract_file_content(file_path, apply_length_limit)
    word_count = len(text)

    return size, file_type, word_count

def audio_to_text(audio_path):
    return "这是录音转写的文字内容。"


# ==================== 简报系统工具类 ====================

import re
import hashlib
import unicodedata
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, NavigableString, Tag, Comment
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders


class BriefingWebScraper:
    """简报系统 - 通用网页抓取器"""
    
    def __init__(self, timeout=15, retry=3):
        self.timeout = timeout
        self.retry = retry
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
        })

    def close(self):
        """关闭会话，释放连接资源"""
        if self.session:
            self.session.close()
            self.session = None

    def __del__(self):
        try:
            self.close()
        except Exception:
            pass

    def __enter__(self):
        return self

    def __exit__(self, *args):
        self.close()
    
    def fetch(self, url, encoding=None):
        """获取网页内容"""
        for attempt in range(self.retry):
            try:
                response = self.session.get(url, timeout=self.timeout)
                response.raise_for_status()
                
                if encoding:
                    response.encoding = encoding
                else:
                    response.encoding = response.apparent_encoding
                
                return BeautifulSoup(response.text, 'html.parser')
            
            except Exception as e:
                logger.warning(f"抓取失败 (尝试 {attempt+1}/{self.retry}): {url} - {str(e)}")
                if attempt < self.retry - 1:
                    time.sleep(2)
        
        return None
    
    def extract_articles_from_index(self, index_url, config=None):
        """从索引页提取文章列表"""
        soup = self.fetch(index_url)
        if not soup:
            return []
        
        articles = []
        config = config or {}
        
        # 查找所有链接
        link_selector = config.get('link_selector', 'a[href]')
        links = soup.select(link_selector)
        
        seen_urls = set()
        
        for link in links:
            href = link.get('href', '')
            if not href:
                continue
            
            # 构建完整URL
            full_url = urljoin(index_url, href)
            
            # URL去重
            if full_url in seen_urls:
                continue
            seen_urls.add(full_url)
            
            # 过滤条件
            if config.get('url_pattern'):
                if not re.search(config['url_pattern'], full_url):
                    continue
            
            if config.get('exclude_pattern'):
                if re.search(config['exclude_pattern'], full_url):
                    continue
            
            # 提取标题
            title = link.get_text(strip=True)
            if not title:
                continue
            
            # 标题长度过滤
            min_len = config.get('min_title_length', 5)
            if len(title) < min_len:
                continue
            
            articles.append({
                'title': title,
                'url': full_url
            })
        
        return articles
    
    def extract_article_content(self, url, config=None):
        """提取文章正文"""
        soup = self.fetch(url)
        if not soup:
            return None, None
        
        config = config or {}
        
        # 提取标题
        title = None
        title_selector = config.get('title_selector', 'h1, .title, .article-title')
        title_elem = soup.select_one(title_selector)
        if title_elem:
            title = title_elem.get_text(strip=True)
        
        # 提取正文
        content = None
        content_selectors = config.get('content_selector', [
            'article', '.article-content', '.content', '#content',
            '.post-content', '.entry-content', '.text'
        ])
        
        if isinstance(content_selectors, str):
            content_selectors = [content_selectors]
        
        for selector in content_selectors:
            content_elem = soup.select_one(selector)
            if content_elem:
                # 清理无关标签
                for tag in content_elem.find_all(['script', 'style', 'nav', 'aside', 'iframe']):
                    tag.decompose()
                
                # 提取段落
                paragraphs = []
                for p in content_elem.find_all(['p', 'div']):
                    text = p.get_text(strip=True)
                    if text and len(text) > 10:  # 过滤短文本
                        paragraphs.append(text)
                
                if paragraphs:
                    content = '\n\n'.join(paragraphs)
                    break
        
        # 发布日期
        publish_date = None
        date_selector = config.get('date_selector', '.date, .time, .publish-time')
        date_elem = soup.select_one(date_selector)
        if date_elem:
            publish_date = date_elem.get_text(strip=True)
        
        return {
            'title': title,
            'content': content,
            'publish_date': publish_date
        }


class BriefingDocumentGenerator:
    """简报系统 - 文档生成器"""
    
    @staticmethod
    def normalize_text(text):
        """规范化文本"""
        if not text:
            return ""
        text = unicodedata.normalize('NFKC', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    @staticmethod
    def create_word_document(articles, output_path, title="学习资料汇编", date=None):
        """生成Word文档"""
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 封面
        cover_title = doc.add_paragraph(title)
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover_title.runs[0]
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(28)
        run.font.bold = True
        
        if date:
            cover_date = doc.add_paragraph(date)
            cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cover_date.runs[0]
            run.font.size = Pt(16)
            run.font.name = '楷体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # 目录
        toc_title = doc.add_paragraph('目  录')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_title.runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        
        for idx, article in enumerate(articles, 1):
            toc_line = f"{idx}. {article.get('title', '无标题')}"
            p = doc.add_paragraph(toc_line)
            p.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_page_break()
        
        # 正文
        for idx, article in enumerate(articles, 1):
            # 标题
            article_title = article.get('title', '无标题')
            title_para = doc.add_paragraph(article_title)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.runs[0]
            run.font.name = '方正小标宋简体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            run.font.size = Pt(16)
            run.font.bold = True
            
            # 来源信息
            source = article.get('source_name', '')
            publish_date = article.get('publish_date', '')
            if source or publish_date:
                source_line = f"{publish_date} {source}"
                source_para = doc.add_paragraph(source_line)
                source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = source_para.runs[0]
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 正文内容
            content = article.get('content', '')
            if content:
                paragraphs = re.split(r'\n+', content)
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.first_line_indent = Inches(0.3)
                        p.paragraph_format.line_spacing = 1.5
                        run = p.runs[0]
                        run.font.size = Pt(12)
            
            # 文章间分隔
            doc.add_paragraph()
            doc.add_paragraph()
        
        # 保存文档
        doc.save(output_path)
        return output_path


class BriefingEmailSender:
    """简报系统 - 邮件发送器"""
    
    def __init__(self, app=None):
        self.app = app
        self.enabled = False
        if app:
            self.init_app(app)
    
    def init_app(self, app):
        self.app = app
        self.enabled = app.config.get('MAIL_SERVER') is not None
    
    def send_briefing(self, recipients, subject, body, attachments=None):
        """发送简报邮件"""
        if not self.enabled:
            logger.warning("邮件服务未配置")
            return False
        
        try:
            msg = MIMEMultipart()
            msg['From'] = self.app.config['MAIL_DEFAULT_SENDER']
            msg['To'] = ', '.join(recipients)
            msg['Subject'] = subject
            
            # 邮件正文
            msg.attach(MIMEText(body, 'html', 'utf-8'))
            
            # 附件
            if attachments:
                for file_path in attachments:
                    if os.path.exists(file_path):
                        with open(file_path, 'rb') as f:
                            part = MIMEBase('application', 'octet-stream')
                            part.set_payload(f.read())
                            encoders.encode_base64(part)
                            filename = os.path.basename(file_path)
                            part.add_header('Content-Disposition', 
                                          f'attachment; filename="{filename}"')
                            msg.attach(part)
            
            # 发送邮件
            with smtplib.SMTP_SSL(self.app.config['MAIL_SERVER'], 
                                 self.app.config['MAIL_PORT']) as server:
                server.login(self.app.config['MAIL_USERNAME'],
                           self.app.config['MAIL_PASSWORD'])
                server.send_message(msg)
            
            logger.info(f"邮件发送成功: {recipients}")
            return True
        
        except Exception as e:
            logger.error(f"邮件发送失败: {str(e)}")
            return False


class BriefingStatisticsCalculator:
    """简报系统 - 统计计算器"""
    
    @staticmethod
    def calculate_daily_stats(date=None):
        """计算每日统计"""
        from models import Briefing, BriefingArticle, db
        
        if date is None:
            date = datetime.now().strftime('%Y-%m-%d')
        
        # 查询当日简报
        briefings = Briefing.query.filter(
            db.func.date(Briefing.start_time) == date
        ).all()
        
        briefing_ids = [b.id for b in briefings]
        
        # 查询当日文章
        articles = BriefingArticle.query.filter(BriefingArticle.briefing_id.in_(briefing_ids)).all()
        
        # 统计
        stats = {
            'briefings_count': len(briefings),
            'articles_count': len(articles),
            'words_count': sum(a.word_count for a in articles),
            'top_keywords': {},
            'top_sources': {}
        }
        
        # 关键词统计
        keyword_count = {}
        for b in briefings:
            for kw in b.get_keywords():
                keyword_count[kw] = keyword_count.get(kw, 0) + 1
        
        stats['top_keywords'] = sorted(keyword_count.items(), 
                                       key=lambda x: x[1], reverse=True)[:10]
        
        # 数据源统计
        source_count = {}
        for a in articles:
            source = a.source_name
            if source:
                source_count[source] = source_count.get(source, 0) + 1
        
        stats['top_sources'] = sorted(source_count.items(),
                                      key=lambda x: x[1], reverse=True)[:10]
        
        return stats
    
    @staticmethod
    def get_trend_data(days=30):
        """获取趋势数据"""
        from models import Briefing, db
        
        end_date = datetime.now()
        start_date = end_date - timedelta(days=days) if 'timedelta' in dir() else datetime.now()
        
        # 按日期分组统计
        result = db.session.query(
            db.func.date(Briefing.start_time).label('date'),
            db.func.count(Briefing.id).label('count')
        ).filter(
            Briefing.start_time >= start_date
        ).group_by(
            db.func.date(Briefing.start_time)
        ).all()
        
        dates = [str(r.date) for r in result]
        counts = [r.count for r in result]
        
        return {'dates': dates, 'counts': counts}


class BriefingSearchAPI:
    """简报系统 - 搜索API封装"""
    
    def __init__(self, api_type='baidu'):
        self.api_type = api_type
    
    def search(self, keyword, max_results=10):
        """执行搜索"""
        if self.api_type == 'baidu':
            return self._search_baidu(keyword, max_results)
        elif self.api_type == 'bing':
            return self._search_bing(keyword, max_results)
        else:
            return []
    
    def _search_baidu(self, keyword, max_results):
        """百度搜索"""
        url = f"https://www.baidu.com/s?wd={keyword}&rn={max_results}"
        scraper = BriefingWebScraper()
        soup = scraper.fetch(url, encoding='utf-8')
        
        if not soup:
            return []
        
        results = []
        for item in soup.select('.result'):
            title_elem = item.select_one('h3 a')
            if not title_elem:
                continue
            
            title = title_elem.get_text(strip=True)
            link = title_elem.get('href', '')
            
            snippet_elem = item.select_one('.c-abstract')
            snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
            
            results.append({
                'title': title,
                'url': link,
                'snippet': snippet,
                'source': '百度搜索'
            })
        
        return results[:max_results]
    
    def _search_bing(self, keyword, max_results):
        """必应搜索"""
        url = f"https://www.bing.com/search?q={keyword}&count={max_results}"
        scraper = BriefingWebScraper()
        soup = scraper.fetch(url, encoding='utf-8')
        
        if not soup:
            return []
        
        results = []
        for item in soup.select('.b_algo'):
            title_elem = item.select_one('h2 a')
            if not title_elem:
                continue
            
            title = title_elem.get_text(strip=True)
            link = title_elem.get('href', '')
            
            snippet_elem = item.select_one('.b_caption p')
            snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
            
            results.append({
                'title': title,
                'url': link,
                'snippet': snippet,
                'source': '必应搜索'
            })
        
        return results[:max_results]